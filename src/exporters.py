import html
import hashlib
import logging
import os
import re
import sys
import warnings
from dataclasses import dataclass
from io import BytesIO

from markdown_it import MarkdownIt
from markdown_it.tree import SyntaxTreeNode

from src.errors import ExportError
from src.logging_utils import get_logger, log_event
from src.schemas import CoverLetterArtifact, ProjectEntry, TailoredResumeArtifact


_MARKDOWN = MarkdownIt("commonmark", {"html": False})
LOGGER = get_logger(__name__)
_WEASYPRINT_DLL_HANDLE = None


def _configure_weasyprint_windows_runtime():
    global _WEASYPRINT_DLL_HANDLE

    if sys.platform != "win32" or _WEASYPRINT_DLL_HANDLE is not None:
        return

    candidate_dirs = []
    env_dirs = os.getenv("WEASYPRINT_DLL_DIRECTORIES", "")
    if env_dirs:
        candidate_dirs.extend(path for path in env_dirs.split(os.pathsep) if path)
    candidate_dirs.append(r"C:\msys64\mingw64\bin")

    for dll_dir in candidate_dirs:
        if not os.path.isdir(dll_dir):
            continue
        try:
            os.environ["WEASYPRINT_DLL_DIRECTORIES"] = dll_dir
            _WEASYPRINT_DLL_HANDLE = os.add_dll_directory(dll_dir)
            log_event(
                LOGGER,
                logging.INFO,
                "weasyprint_windows_runtime_configured",
                "Configured the WeasyPrint Windows DLL search path.",
                dll_directory=dll_dir,
            )
            return
        except (AttributeError, FileNotFoundError, OSError):
            continue


def export_text_bytes(report: CoverLetterArtifact | TailoredResumeArtifact) -> bytes:
    return report.plain_text.encode("utf-8")


# ---------------------------------------------------------------------------
# ThemeSpec — single source of truth for theming (ADR-015 follow-up).
#
# Historically there were THREE hand-synced palette maps
# (`_RESUME_THEME_PALETTES`, `_COVER_LETTER_THEME_PALETTES`,
# `_DOCX_THEME_PALETTES`). ADR-015's follow-up note called for collapsing
# them into one typed spec before a third theme lands so a theme can't
# drift across the PDF resume / PDF cover letter / DOCX renderers — a
# user picking a theme MUST get the same look in every artifact + format.
#
# One `ThemeSpec` per theme now derives all three palettes. Adding a
# theme = one entry in `_THEME_SPECS`. The three module-level maps below
# are kept (derived from the registry) so the resolvers + any external
# reference stay byte-for-byte identical — this refactor is provably
# output-neutral for the existing themes.
#
# `layout` is reserved for the gated non-ATS two-column presentation
# theme (Phase 3); only the *resume* renderer will branch on it — a
# cover letter is prose and always renders single-column.
# ---------------------------------------------------------------------------


@dataclass(frozen=True)
class ThemeSpec:
    key: str
    label: str
    # Canonical colors ('#'-prefixed; the DOCX adapter strips '#'
    # and upper-cases for OOXML).
    ink: str
    muted: str
    accent: str
    line: str
    paper: str
    surface: str
    # Resume uses this for BOTH the `accent_soft` blockquote wash and
    # the inline-`code` background (they were always equal).
    accent_soft: str
    # Cover-letter <strong> color (independent of `ink`: classic_ats
    # darkens it; neutral collapses it to ink).
    cover_strong_color: str
    # Type — CSS font *stacks* for the PDF surfaces.
    body_font_family: str
    h1_font_family: str
    prose_font_family: str
    prose_line_height: str
    # DOCX needs single OOXML family names (no stack / no quotes).
    docx_body_font: str
    docx_heading_font: str
    docx_prose_font: str
    # Header underline rule width, shared by resume + cover letter.
    header_border_px: int
    # Color of the header divider rule (resume name underline + cover
    # letter greeting break). Defaults to the literal CSS token
    # `var(--accent)` so EVERY pre-existing theme renders the rule
    # exactly as before (byte-identical — the divider was always the
    # accent). A theme can override it to a separate hex (e.g.
    # creative_warm deepens just this structural line while keeping
    # the brighter accent on section headers).
    header_rule_color: str = "var(--accent)"
    # Opt-in header BAND background (ADR-029 v2 / report.md designer
    # expansion). "" → today's plain rule-only header (the resume
    # header markup is byte-IDENTICAL for any theme that doesn't set
    # this — the 3 ATS-simple themes stay unchanged). Non-empty → a
    # full-bleed coloured masthead behind name/role/contact. fg is the
    # on-band text colour (set it when the band is dark).
    header_band_bg: str = ""
    header_band_fg: str = ""
    # "single_column" | "two_column". Only the resume renderer branches
    # on this; the cover-letter + DOCX paths read it to stay binary
    # (a two-column theme renders single-column there).
    layout: str = "single_column"
    # ADR-032. For a two_column theme, names the bespoke designer
    # layout whose renderer builds the page ("timeline_tech",
    # "editorial_minimal", "classic_slate", "monochrome_black",
    # "plum_berry", "burgundy_champagne"). "" for single-column themes.
    # `_build_resume_html` dispatches on this via `_TWOCOL_RENDERERS`.
    # Kept SEPARATE from `layout` so `layout` stays the binary
    # single/two discriminator every non-resume reader relies on.
    twocol_layout: str = ""

    @staticmethod
    def _ooxml(color: str) -> str:
        # "#8f6845" -> "8F6845"; only valid for hex colors. All theme
        # color fields that reach DOCX are hex (rgba never does).
        return color.lstrip("#").upper()

    def resume_palette(self) -> dict:
        return {
            "ink": self.ink,
            "muted": self.muted,
            "accent": self.accent,
            "accent_soft": self.accent_soft,
            "line": self.line,
            "paper": self.paper,
            "surface": self.surface,
            "body_font_family": self.body_font_family,
            "h1_font_family": self.h1_font_family,
            "prose_font_family": self.prose_font_family,
            "prose_line_height": self.prose_line_height,
            "header_border_width": f"{self.header_border_px}px",
            "header_rule_color": self.header_rule_color,
            # Safe CSS defaults so the (class-gated, never-matched for
            # non-opting themes) band rules are always valid.
            "header_band_bg": self.header_band_bg or "transparent",
            "header_band_fg": self.header_band_fg or "inherit",
            "code_bg": self.accent_soft,
        }

    def cover_letter_palette(self) -> dict:
        return {
            "ink": self.ink,
            "muted": self.muted,
            "accent": self.accent,
            "line": self.line,
            "paper": self.paper,
            "surface": self.surface,
            "strong_color": self.cover_strong_color,
            "header_border_width": f"{self.header_border_px}px",
            "header_rule_color": self.header_rule_color,
            # The letter is all prose, so it follows the theme's PROSE
            # font (not body). Since the 2026-05-21 typography
            # unification every theme's prose font is the shared Arial
            # sans family, so the cover letter always matches its
            # résumé — the "matched set" guarantee.
            "body_font_family": self.prose_font_family,
        }

    def docx_palette(self) -> dict:
        return {
            "ink": self._ooxml(self.ink),
            "muted": self._ooxml(self.muted),
            "accent": self._ooxml(self.accent),
            "line": self._ooxml(self.line),
            "body_font": self.docx_body_font,
            "heading_font": self.docx_heading_font,
            "prose_font": self.docx_prose_font,
        }


# The registry. NEW themes are added here only — the three derived
# maps + every renderer pick them up automatically, so resume + cover
# letter + DOCX can never drift.
#
# TYPOGRAPHY (operator decision 2026-05-21): all five themes share ONE
# font family — Arial / Helvetica sans-serif (h1, body, prose; Arial
# for the DOCX OOXML names). Themes used to mix serif/sans for
# identity (professional_neutral was all-Georgia, classic_ats +
# creative_warm had serif headings); that was unified to a single sans
# family per the operator's request. Themes now differentiate by
# COLOR, PAPER, and HEADER TREATMENT only — not by typeface. Any new
# theme MUST use the same Arial sans family unless that decision is
# revisited.
_THEME_SPECS: dict[str, "ThemeSpec"] = {
    "classic_ats": ThemeSpec(
        key="classic_ats",
        label="Classic ATS",
        ink="#221912",
        muted="#6b5648",
        accent="#8f6845",
        line="#d7c2af",
        paper="#fffdf9",
        surface="#fffdfa",
        accent_soft="rgba(143, 104, 69, 0.10)",
        cover_strong_color="#17100b",
        # Unified Arial/Helvetica sans family across all themes
        # (operator decision 2026-05-21) — see the registry note above.
        body_font_family="Arial, Helvetica, sans-serif",
        h1_font_family="Arial, Helvetica, sans-serif",
        prose_font_family="Arial, Helvetica, sans-serif",
        prose_line_height="1.55",
        docx_body_font="Arial",
        docx_heading_font="Arial",
        docx_prose_font="Arial",
        header_border_px=3,
        layout="single_column",
    ),
    "professional_neutral": ThemeSpec(
        key="professional_neutral",
        label="Professional",
        ink="#0a0a0a",
        muted="#555555",
        accent="#0a0a0a",
        line="#bfbfbf",
        paper="#ffffff",
        surface="#ffffff",
        accent_soft="rgba(0, 0, 0, 0.04)",
        cover_strong_color="#0a0a0a",
        # Unified Arial/Helvetica sans family across all themes
        # (operator decision 2026-05-21) — see the registry note above.
        body_font_family="Arial, Helvetica, sans-serif",
        h1_font_family="Arial, Helvetica, sans-serif",
        prose_font_family="Arial, Helvetica, sans-serif",
        prose_line_height="1.55",
        docx_body_font="Arial",
        docx_heading_font="Arial",
        docx_prose_font="Arial",
        header_border_px=2,
        layout="single_column",
    ),
    # NEW (Phase 2a) — modern, one deep professional blue.
    # Single-column → fully ATS-safe. Differentiated by a cool slate
    # ink + blue accent on a faint cool paper. Accent #1A56DB clears
    # ~5.9:1 on white (safe for any text). Fonts are the shared
    # Arial/Helvetica sans family (see registry note) — web-safe so
    # WeasyPrint renders identically on the Linux render host (no
    # missing-font fallback). Audience: tech / product / data / ops.
    "modern_blue": ThemeSpec(
        key="modern_blue",
        label="Modern Blue",
        ink="#16202e",
        muted="#5a6b7b",
        accent="#1a56db",
        line="#dfe5ec",
        # Faint COOL off-white (operator-picked): the classic_ats
        # "designed, not stark" trick in a cool key so the paper sits
        # right under the blue accent. ~9 units off white max →
        # print/photocopy-safe, and ATS-irrelevant (paint, not text).
        paper="#f6f8fd",
        surface="#f8fafe",
        accent_soft="rgba(26, 86, 219, 0.07)",
        cover_strong_color="#16202e",
        body_font_family="Arial, Helvetica, sans-serif",
        h1_font_family="Arial, Helvetica, sans-serif",
        prose_font_family="Arial, Helvetica, sans-serif",
        prose_line_height="1.55",
        docx_body_font="Arial",
        docx_heading_font="Arial",
        docx_prose_font="Arial",
        header_border_px=2,
        layout="single_column",
    ),
    # NEW (Phase 2b) — modern editorial. Emerald accent for creative
    # energy without leaving professional; a soft warm "sand" header
    # band. Distinct from classic_ats (warm brown + cream paper) and
    # from modern_blue (cool blue). Faint near-neutral warm paper.
    # Emerald #00a388 is the proven Awesome-CV accent; it only colors
    # headings/rules/labels (body stays ink) so contrast is fine.
    # Single-column → ATS-safe. Audience: marketing / comms /
    # design-adjacent that still needs to pass ATS.
    "creative_warm": ThemeSpec(
        key="creative_warm",
        label="Creative Warm",
        ink="#232524",
        muted="#5f6b63",
        accent="#00a388",
        line="#d8e2dc",
        paper="#fcfcf6",
        surface="#fdfdf8",
        accent_soft="rgba(0, 163, 136, 0.08)",
        cover_strong_color="#232524",
        # Designed header: soft warm "sand" band (deeper than the
        # #fcfcf6 paper) with dark ink text — elegant/editorial, warm,
        # not loud. Operator chose this over the emerald-band variant.
        header_band_bg="#efe7d8",
        header_band_fg="#232524",
        # Deeper, greener than the #00a388 section accent so the
        # name/body divider reads as a deliberate anchor line rather
        # than the same bright emerald (operator request).
        header_rule_color="#0b7c5e",
        # Unified Arial/Helvetica sans family across all themes
        # (operator decision 2026-05-21) — see the registry note above.
        h1_font_family="Arial, Helvetica, sans-serif",
        body_font_family="Arial, Helvetica, sans-serif",
        prose_font_family="Arial, Helvetica, sans-serif",
        prose_line_height="1.55",
        docx_body_font="Arial",
        docx_heading_font="Arial",
        docx_prose_font="Arial",
        header_border_px=2,
        layout="single_column",
    ),
    # NEW (Phase 2c) — architectural minimal. Near-monochrome: deep
    # cool near-black ink AND accent (no colour — the "design" is
    # typographic), one HAIRLINE rule (header_border_px=1), airier
    # prose line-height for generous whitespace, and a solid ink
    # masthead band. Crisp pure white on purpose (deliberate contrast
    # to modern_blue's tint — minimalism reads as stark, not "designed
    # paper"). Single-column → ATS-safe. Audience: architecture /
    # design / senior-eng "confident minimal". Distinct from
    # professional_neutral (same monochrome idea, but architect_mono
    # adds the hairline rule + ink masthead + airier spacing).
    "architect_mono": ThemeSpec(
        key="architect_mono",
        label="Architect Mono",
        ink="#1a1f29",
        muted="#6b7280",
        accent="#131a28",
        line="#e4e7ec",
        paper="#ffffff",
        surface="#ffffff",
        accent_soft="rgba(19, 26, 40, 0.05)",
        cover_strong_color="#131a28",
        # Designed header: a confident solid INK masthead with white
        # text — an architectural/minimal statement (vs. the plain
        # rule the ATS-simple themes keep). Operator will tweak.
        header_band_bg="#131a28",
        header_band_fg="#ffffff",
        body_font_family="Arial, Helvetica, sans-serif",
        h1_font_family="Arial, Helvetica, sans-serif",
        prose_font_family="Arial, Helvetica, sans-serif",
        prose_line_height="1.6",
        docx_body_font="Arial",
        docx_heading_font="Arial",
        docx_prose_font="Arial",
        header_border_px=1,
        layout="single_column",
    ),
    # NEW — single-column "Noir Cream": a PURE-BLACK full-bleed masthead
    # band on warm cream paper. Colours lifted exactly from the
    # two-column "Monochrome Black" Claude-design template
    # (resume_builder/05-monochrome-black.html): the band is its
    # #000000 sidebar, the page is its #f3eee3 cream. True monochrome
    # (accent = ink). Distinct from professional_neutral (stark white,
    # no band) and architect_mono (cool blue-ink band #131a28 on white)
    # — this is the only theme pairing a true-black band with warm
    # cream. Single-column → ATS-safe. Audience: design / creative /
    # editorial roles wanting a bolder, warmer monochrome.
    "noir_cream": ThemeSpec(
        key="noir_cream",
        label="Noir Cream",
        ink="#0a0a0a",
        muted="#6e6a63",
        accent="#0a0a0a",
        line="#d8d2c6",
        paper="#f3eee3",
        surface="#f3eee3",
        accent_soft="#ece6d9",
        cover_strong_color="#0a0a0a",
        # Unified Arial/Helvetica sans family across all themes
        # (operator decision 2026-05-21) — see the registry note above.
        body_font_family="Arial, Helvetica, sans-serif",
        h1_font_family="Arial, Helvetica, sans-serif",
        prose_font_family="Arial, Helvetica, sans-serif",
        prose_line_height="1.55",
        docx_body_font="Arial",
        docx_heading_font="Arial",
        docx_prose_font="Arial",
        header_border_px=2,
        # Pure-black full-bleed masthead band (the template's #000000
        # sidebar) with cream text mirroring the paper.
        header_band_bg="#000000",
        header_band_fg="#f3eee3",
        layout="single_column",
    ),
    # ===================================================================
    # ADR-032 — six bespoke, NON-ATS, two-column résumé themes. Each is a
    # finished designer template (resume_builder/*.html) with its own
    # `twocol_layout` renderer (see `_TWOCOL_RENDERERS`). They REPLACE the
    # old `presentation_twocol` placeholder.
    #
    # The two-column renderers inline the design's own CSS (the design IS
    # the colour), so the colour fields below do NOT drive the résumé
    # look — they exist only so each theme's COVER LETTER + DOCX (which
    # always render single-column) come out in a matching palette. They
    # are lifted from each template's `:root` tokens.
    #
    # Gating: non-`professional_neutral` → Pro/Business by the existing
    # by-exclusion gate (ADR-027), no tiers.py change. NON-ATS: the
    # frontend picker hint warns explicitly. DOM is authored
    # header→main→sidebar so the PDF text layer extracts linearly.
    # ===================================================================
    # 01 — Timeline / Tech. Dark navy sidebar, blue accent, dot-and-rail
    # experience timeline. Audience: software / data / engineering.
    "timeline_tech": ThemeSpec(
        key="timeline_tech",
        label="Timeline Tech (2-col)",
        ink="#111418",
        muted="#5b6470",
        accent="#4f6bed",
        line="#e4e7ec",
        paper="#ffffff",
        surface="#ffffff",
        accent_soft="#eef1ff",
        cover_strong_color="#111418",
        body_font_family="Arial, Helvetica, sans-serif",
        h1_font_family="Arial, Helvetica, sans-serif",
        prose_font_family="Arial, Helvetica, sans-serif",
        prose_line_height="1.55",
        docx_body_font="Arial",
        docx_heading_font="Arial",
        docx_prose_font="Arial",
        header_border_px=2,
        layout="two_column",
        twocol_layout="timeline_tech",
    ),
    # 02 — Editorial Minimal. Light sand sidebar, terracotta accent,
    # hairline-rule section headers, head-row experience. Audience:
    # design / editorial / communications / brand / product.
    "editorial_minimal": ThemeSpec(
        key="editorial_minimal",
        label="Editorial Minimal (2-col)",
        ink="#1c1a17",
        muted="#6b6760",
        accent="#9a4a2a",
        line="#d9d4cb",
        paper="#fbf7f0",
        surface="#fbf7f0",
        accent_soft="#f1e7df",
        cover_strong_color="#1c1a17",
        body_font_family="Arial, Helvetica, sans-serif",
        h1_font_family="Arial, Helvetica, sans-serif",
        prose_font_family="Arial, Helvetica, sans-serif",
        prose_line_height="1.6",
        docx_body_font="Arial",
        docx_heading_font="Arial",
        docx_prose_font="Arial",
        header_border_px=1,
        layout="two_column",
        twocol_layout="editorial_minimal",
    ),
    # 04 — Classic Slate. Pale slate sidebar, deep emerald accent, left
    # date-gutter experience. Audience: consulting / finance / research /
    # academic / policy / legal / clinical.
    "classic_slate": ThemeSpec(
        key="classic_slate",
        label="Classic Slate (2-col)",
        ink="#0f1722",
        muted="#586374",
        accent="#1f5d4c",
        line="#d5dae2",
        paper="#ffffff",
        surface="#ffffff",
        accent_soft="#e6efeb",
        cover_strong_color="#0f1722",
        body_font_family="Arial, Helvetica, sans-serif",
        h1_font_family="Arial, Helvetica, sans-serif",
        prose_font_family="Arial, Helvetica, sans-serif",
        prose_line_height="1.55",
        docx_body_font="Arial",
        docx_heading_font="Arial",
        docx_prose_font="Arial",
        header_border_px=2,
        layout="two_column",
        twocol_layout="classic_slate",
    ),
    # 05 — Monochrome Black. True-black sidebar, monochrome (accent =
    # ink), heavy-rule head-row experience. Audience: creative direction /
    # architecture / fashion / senior product.
    "monochrome_black": ThemeSpec(
        key="monochrome_black",
        label="Monochrome Black (2-col)",
        ink="#0a0a0a",
        muted="#6e6a63",
        accent="#0a0a0a",
        line="#d8d2c6",
        paper="#f3eee3",
        surface="#f3eee3",
        accent_soft="#ece6d9",
        cover_strong_color="#0a0a0a",
        body_font_family="Arial, Helvetica, sans-serif",
        h1_font_family="Arial, Helvetica, sans-serif",
        prose_font_family="Arial, Helvetica, sans-serif",
        prose_line_height="1.6",
        docx_body_font="Arial",
        docx_heading_font="Arial",
        docx_prose_font="Arial",
        header_border_px=2,
        layout="two_column",
        twocol_layout="monochrome_black",
    ),
    # 08 — Plum Berry. Deep plum sidebar, berry-rose accent, dusty-pink
    # paper, head-row experience. Audience: PR / fashion comms / beauty /
    # lifestyle media / senior marketing.
    "plum_berry": ThemeSpec(
        key="plum_berry",
        label="Plum Berry (2-col)",
        ink="#2a1d2a",
        muted="#7a6877",
        accent="#b94f7a",
        line="#e3d4dd",
        paper="#f6ecef",
        surface="#f6ecef",
        accent_soft="#f2dce6",
        cover_strong_color="#2a1d2a",
        body_font_family="Arial, Helvetica, sans-serif",
        h1_font_family="Arial, Helvetica, sans-serif",
        prose_font_family="Arial, Helvetica, sans-serif",
        prose_line_height="1.65",
        docx_body_font="Arial",
        docx_heading_font="Arial",
        docx_prose_font="Arial",
        header_border_px=1,
        layout="two_column",
        twocol_layout="plum_berry",
    ),
    # 10 — Burgundy & Champagne. Deep wine sidebar, champagne-gold accent,
    # warm ivory paper, left date-gutter experience. Audience: law /
    # banking / private wealth / hospitality / senior advisory.
    "burgundy_champagne": ThemeSpec(
        key="burgundy_champagne",
        label="Burgundy Champagne (2-col)",
        ink="#1c1311",
        muted="#6f5e54",
        accent="#a17a3a",
        line="#e3d8c5",
        paper="#faf3e3",
        surface="#faf3e3",
        accent_soft="#f1e5cd",
        cover_strong_color="#1c1311",
        body_font_family="Arial, Helvetica, sans-serif",
        h1_font_family="Arial, Helvetica, sans-serif",
        prose_font_family="Arial, Helvetica, sans-serif",
        prose_line_height="1.6",
        docx_body_font="Arial",
        docx_heading_font="Arial",
        docx_prose_font="Arial",
        header_border_px=3,
        layout="two_column",
        twocol_layout="burgundy_champagne",
    ),
}


def resolve_theme(theme: str | None) -> "ThemeSpec":
    """Theme name -> ThemeSpec. Unknown / blank -> classic_ats, the
    same fallback the three legacy resolvers used."""
    return _THEME_SPECS.get(
        str(theme or "").strip(), _THEME_SPECS["classic_ats"]
    )


# Public: the canonical set of renderable theme keys. The backend
# export gates (artifact_export_service / resume_builder_service)
# import THIS instead of hand-maintaining their own theme sets —
# adding a ThemeSpec to the registry is the single edit a new theme
# needs server-side (entitlement is separate, by-exclusion in
# backend.tiers). ADR-015 follow-up: no theme set may drift.
SUPPORTED_THEMES: frozenset[str] = frozenset(_THEME_SPECS)


def export_pdf_bytes(report: CoverLetterArtifact | TailoredResumeArtifact) -> bytes:
    try:
        theme = getattr(report, "theme", None)
        document_kind = (
            "tailored_resume"
            if isinstance(report, TailoredResumeArtifact)
            else "cover_letter"
        )
        artifact = report if isinstance(report, TailoredResumeArtifact) else None
        return generate_pdf(
            report.markdown,
            title=report.title,
            theme=theme,
            document_kind=document_kind,
            artifact=artifact,
        ).getvalue()
    except ExportError as error:
        log_event(
            LOGGER,
            logging.ERROR,
            "pdf_export_failed",
            "PDF export failed.",
            report_title=report.title,
            filename_stem=report.filename_stem,
            error_type=type(error).__name__,
            details=error.details,
        )
        raise


def build_resume_preview_html(artifact: TailoredResumeArtifact) -> str:
    return _build_resume_html(
        artifact.markdown,
        title=artifact.title,
        theme=artifact.theme,
        artifact=artifact,
    )


def build_cover_letter_preview_html(artifact: CoverLetterArtifact) -> str:
    return _build_cover_letter_html(
        artifact.markdown,
        title=artifact.title,
        theme=getattr(artifact, "theme", "classic_ats"),
    )


def _paragraph_text(text):
    return html.escape(text or "")


def _style_cover_letter_body(body_html: str) -> str:
    styled_body = re.sub(
        r"\s*<h1>.*?</h1>",
        "",
        body_html or "",
        count=1,
        flags=re.DOTALL,
    ).strip()
    greeting_match = re.search(r"<p>.*?</p>", styled_body, flags=re.DOTALL)
    if greeting_match:
        greeting_html = '<div class="cover-letter-greeting-break"></div>{content}'.format(
            content=greeting_match.group(0)
        )
        styled_body = styled_body.replace(greeting_match.group(0), greeting_html, 1)
    signoff_match = re.search(r"<p>([^<]+),\s*([^<]+)</p>\s*$", styled_body, flags=re.DOTALL)
    if signoff_match:
        signoff_html = '<div class="cover-letter-signoff"><p>{line_one},</p><p>{line_two}</p></div>'.format(
            line_one=html.escape(signoff_match.group(1).strip()),
            line_two=html.escape(signoff_match.group(2).strip()),
        )
        styled_body = re.sub(r"<p>([^<]+),\s*([^<]+)</p>\s*$", signoff_html, styled_body, count=1, flags=re.DOTALL)
    return styled_body


def _split_cover_letter_title(title: str) -> tuple[str, str]:
    normalized_title = str(title or "Cover Letter").strip() or "Cover Letter"
    if " - " in normalized_title and normalized_title.lower().endswith(" cover letter"):
        name, role_with_suffix = normalized_title.split(" - ", 1)
        role = role_with_suffix[: -len(" Cover Letter")].strip()
        return name.strip() or normalized_title, role or "Cover Letter"
    if normalized_title.lower().endswith(" cover letter"):
        headline = normalized_title[: -len(" Cover Letter")].strip()
        return headline or normalized_title, "Cover Letter"
    return normalized_title, ""


# Derived from `_THEME_SPECS` (ADR-015 follow-up). Do NOT hand-edit —
# add a ThemeSpec to the registry instead so resume + cover letter +
# DOCX stay byte-for-byte in sync.
_COVER_LETTER_THEME_PALETTES = {
    key: spec.cover_letter_palette() for key, spec in _THEME_SPECS.items()
}


def _resolve_cover_letter_palette(theme: str | None) -> dict:
    return _COVER_LETTER_THEME_PALETTES.get(
        theme or "classic_ats", _COVER_LETTER_THEME_PALETTES["classic_ats"]
    )


def _build_cover_letter_html(text, title="Cover Letter", theme="classic_ats"):
    safe_title = html.escape(title or "Cover Letter")
    header_title, header_subtitle = _split_cover_letter_title(title)
    body_html = _style_cover_letter_body(_MARKDOWN.render(text or ""))
    palette = _resolve_cover_letter_palette(theme)

    return """
<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="utf-8" />
    <title>{title}</title>
    <style>
        /* Theme-keyed palette: classic_ats keeps the warm-brown letter
           feel; professional_neutral drops to pure black/white/gray for
           conservative recipients. Prose font is the shared Arial sans
           family across every theme (2026-05-21 unification). */
        :root {{
            --ink: {ink};
            --muted: {muted};
            --accent: {accent};
            --line: {line};
            --paper: {paper};
            --surface: {surface};
        }}

        @page {{ size: A4; margin: 0; }}

        * {{ box-sizing: border-box; }}

        body {{
            margin: 0;
            font-family: {body_font_family};
            color: var(--ink);
            background: var(--paper);
            line-height: 1.72;
            font-size: 11.4pt;
        }}

        .cover-letter-shell {{
            width: 100%;
            min-height: 100vh;
            background: var(--surface);
            padding: 18mm 16mm 18mm;
        }}

        .cover-letter-header {{ margin: 0 0 18px; }}

        .cover-letter-title {{
            margin: 0;
            font-size: 17.5pt;
            font-weight: 700;
            letter-spacing: 0.02em;
            color: var(--ink);
        }}

        .cover-letter-role {{
            margin: 4px 0 0;
            font-size: 10.8pt;
            letter-spacing: 0.08em;
            text-transform: uppercase;
            color: var(--muted);
        }}

        .cover-letter-greeting-break {{
            width: auto;
            border-top: {header_border_width} solid {header_rule_color};
            margin: 8px -16mm 14px;
        }}

        .cover-letter-signoff p {{ margin: 0 0 6px; }}
        .cover-letter-signoff p:last-child {{ margin-bottom: 0; }}

        p {{ margin: 0 0 12px; }}
        strong {{ color: {strong_color}; }}
        em {{ color: var(--muted); }}
        ul, ol {{ margin: 0 0 12px 1.25rem; padding: 0; }}
        li {{ margin: 0 0 6px; }}
    </style>
</head>
<body>
    <main class="cover-letter-shell">
        <header class="cover-letter-header">
            <h1 class="cover-letter-title">{header_title}</h1>
            {header_subtitle}
        </header>
        {body}
    </main>
</body>
</html>
""".format(
        title=safe_title,
        header_title=html.escape(header_title),
        header_subtitle=(
            '<p class="cover-letter-role">{subtitle}</p>'.format(
                subtitle=html.escape(header_subtitle)
            )
            if header_subtitle
            else ""
        ),
        body=body_html,
        **palette,
    )


def _build_resume_section_list(items, empty_state, ordered=False, class_name=""):
    values = [html.escape(str(item or "").strip()) for item in items if str(item or "").strip()]
    if not values:
        return '<p class="resume-empty">{text}</p>'.format(text=html.escape(empty_state))

    tag_name = "ol" if ordered else "ul"
    class_attr = ' class="{name}"'.format(name=class_name) if class_name else ""
    content = "".join("<li>{item}</li>".format(item=value) for value in values)
    return "<{tag}{class_attr}>{content}</{tag}>".format(
        tag=tag_name,
        class_attr=class_attr,
        content=content,
    )


def _build_resume_experience_html(experience_entries):
    if not experience_entries:
        return '<p class="resume-empty">No structured experience entries were available.</p>'

    cards = []
    for entry in experience_entries:
        title = html.escape(entry.title or "Relevant Experience")
        organization = html.escape(entry.organization or "")
        location = html.escape(entry.location or "")
        date_parts = [part for part in [entry.start, entry.end] if part]
        date_text = html.escape(" - ".join(date_parts)) if date_parts else ""
        meta_parts = [part for part in [organization, location] if part]
        bullets_html = _build_resume_section_list(
            entry.bullets,
            "No grounded bullet points were generated for this role.",
            class_name="resume-bullet-list",
        )
        cards.append(
            """
            <article class="resume-experience-card">
                <div class="resume-role-row">
                    <div>
                        <h3>{title}</h3>
                        {meta}
                    </div>
                    {dates}
                </div>
                {bullets}
            </article>
            """.format(
                title=title,
                # Each meta_part was already individually escaped above
                # (organization, location). Joining and re-escaping
                # would double-escape — '<acme>' becomes
                # '&amp;lt;acme&amp;gt;' on the page. Keep the join
                # plain.
                meta=(
                    '<p class="resume-role-meta">{meta}</p>'.format(meta=" | ".join(meta_parts))
                    if meta_parts
                    else ""
                ),
                dates=(
                    '<p class="resume-role-dates">{dates}</p>'.format(dates=date_text)
                    if date_text
                    else ""
                ),
                bullets=bullets_html,
            )
        )
    return "".join(cards)


def _build_resume_projects_html(project_entries: list[ProjectEntry]) -> str:
    if not project_entries:
        return ""
    cards: list[str] = []
    for project in project_entries:
        name = html.escape(project.name or "Project")
        date_parts = [part for part in [project.start, project.end] if part]
        date_text = html.escape(" - ".join(date_parts)) if date_parts else ""
        description = html.escape(project.description or "")
        bullets_html = ""
        if project.bullets:
            bullets_html = _build_resume_section_list(
                project.bullets,
                "",
                class_name="resume-bullet-list",
            )
        meta_parts: list[str] = []
        if project.technologies:
            meta_parts.append(
                "Tech: " + html.escape(", ".join(project.technologies))
            )
        if project.link:
            meta_parts.append("Link: " + html.escape(project.link))
        meta_html = (
            '<p class="resume-role-meta">{meta}</p>'.format(meta=" | ".join(meta_parts))
            if meta_parts
            else ""
        )
        cards.append(
            """
            <article class="resume-project-card">
                <div class="resume-role-row">
                    <div>
                        <h3>{name}</h3>
                        {description}
                    </div>
                    {dates}
                </div>
                {bullets}
                {meta}
            </article>
            """.format(
                name=name,
                description=(
                    '<p class="resume-role-meta">{description}</p>'.format(description=description)
                    if description
                    else ""
                ),
                dates=(
                    '<p class="resume-role-dates">{dates}</p>'.format(dates=date_text)
                    if date_text
                    else ""
                ),
                bullets=bullets_html,
                meta=meta_html,
            )
        )
    return "".join(cards)


def _build_resume_education_html(education_entries):
    if not education_entries:
        return '<p class="resume-empty">No education entries were available.</p>'

    blocks = []
    for entry in education_entries:
        institution = html.escape(entry.institution or "Education")
        degree_parts = [part for part in [entry.degree, entry.field_of_study] if part]
        date_parts = [part for part in [entry.start, entry.end] if part]
        blocks.append(
            """
            <article class="resume-education-card">
                <h3>{institution}</h3>
                {degree}
                {dates}
            </article>
            """.format(
                institution=institution,
                degree=(
                    '<p class="resume-education-meta">{degree}</p>'.format(
                        degree=html.escape(" - ".join(degree_parts))
                    )
                    if degree_parts
                    else ""
                ),
                dates=(
                    '<p class="resume-education-dates">{dates}</p>'.format(
                        dates=html.escape(" - ".join(date_parts))
                    )
                    if date_parts
                    else ""
                ),
            )
        )
    return "".join(blocks)


def _looks_like_contact_link(value: str) -> bool:
    """A contact value is a *link* (github/linkedin/portfolio/site) vs a
    plain detail (location / phone / email). Emails are details, not
    links, even though they contain a domain."""
    v = str(value or "").strip()
    if not v:
        return False
    if v.lower().startswith(("http://", "https://", "www.")):
        return True
    # A real email is local@domain.tld with NO scheme and NO slash → a
    # contact DETAIL, not a link. This is deliberately strict so a URL
    # that merely contains '@' (e.g. medium.com/@handle) is still
    # treated as a link.
    if re.match(r"^[^@\s/]+@[^@\s/]+\.[^@\s/]+$", v):
        return False
    head = v.split("/", 1)[0]  # domain part before any path
    return " " not in head and bool(
        re.match(r"^[A-Za-z0-9.-]+\.[A-Za-z]{2,}$", head)
    )


def _build_resume_contact_inline_html(contact_lines):
    raw = [str(item or "").strip() for item in contact_lines if str(item or "").strip()]
    if not raw:
        return ""
    details = [v for v in raw if not _looks_like_contact_link(v)]
    links = [v for v in raw if _looks_like_contact_link(v)]
    n = len(links)
    # Count-aware packing: AT MOST two lines, a URL never splits
    # mid-string (each item is white-space:nowrap), and neither line is
    # left lopsided —
    #   0-1 links : everything on one line with the details
    #   2 links   : details on line 1; BOTH links share line 2 (no lone
    #               short link stranded on its own line)
    #   3+ links  : details + first link on line 1; the rest on line 2
    #               (e.g. github on line 1; linkedin + portfolio on 2)
    if n <= 1:
        line1_links, line2_links = links, []
    elif n == 2:
        line1_links, line2_links = [], links
    else:
        line1_links, line2_links = links[:1], links[1:]
    line1 = details + line1_links
    line2 = line2_links
    if not line1:  # pathological (no details + the n==2 case): never emit an empty first line
        line1, line2 = line2, []

    def _row(items, links_row=False):
        if not items:
            return ""
        cells = " | ".join(
            '<span class="rc-item">{0}</span>'.format(html.escape(i))
            for i in items
        )
        cls = "resume-contact-inline resume-contact-links" if links_row else "resume-contact-inline"
        return '<p class="{0}">{1}</p>'.format(cls, cells)

    return _row(line1) + _row(line2, links_row=True)


def _build_resume_skills_inline_html(skills, skill_categories=None):
    """Render skills either as a flat pipe-list or, when categories are
    present, as one row per category ('Languages & Tools: Python, SQL').

    Categories take precedence — they're emitted only when the
    structuring pass found 8+ skills clustering naturally, so when
    they're present the flat list would be redundant. Falls back to
    flat for sparse skill sets, the JD-driven path, and any artifact
    built before this field existed.
    """
    if skill_categories:
        rows: list[str] = []
        for label, items in skill_categories.items():
            label_clean = str(label or "").strip()
            cleaned = [
                html.escape(str(item or "").strip())
                for item in items
                if str(item or "").strip()
            ]
            if not label_clean or not cleaned:
                continue
            rows.append(
                '<p class="resume-skill-category">'
                '<strong>{label}:</strong> {items}'
                '</p>'.format(label=html.escape(label_clean), items=", ".join(cleaned))
            )
        if rows:
            return "".join(rows)

    values = [html.escape(str(item or "").strip()) for item in skills if str(item or "").strip()]
    if not values:
        return '<p class="resume-empty">No highlighted skills were generated.</p>'
    return '<p class="resume-skill-inline">{items}</p>'.format(items=" | ".join(values))


_DEFAULT_RESUME_SECTION_ORDER = (
    "summary",
    "skills",
    "experience",
    "projects",
    "education",
    "publications",
    "certifications",
)


def _build_structured_resume_body_classic(
    artifact: TailoredResumeArtifact, header_banded: bool = False
):
    # Summary, Core Skills, Education always render even when sparse —
    # Summary because the workflow always generates one (placeholder
    # surfaces a failure), Skills and Education because both are
    # user-supplied content the resume requires.
    #
    # Experience, Projects, Publications, Certifications drop entirely
    # when empty: students / early-career candidates may legitimately
    # have any combination of these missing, and a placeholder reads as
    # awkward filler rather than as a useful diagnostic. Internships
    # are modelled as regular Experience entries (no separate
    # Internships section).
    #
    # Section order is honored from artifact.section_order when set
    # (resume_builder picks it from the agent's resume_generation
    # output or the deterministic per-profile heuristic). Falls back
    # to the standard professional ordering for legacy callers that
    # build artifacts directly without going through resume_builder.
    name = html.escape(artifact.header.full_name or artifact.title or "Candidate")
    contact_values = []
    if artifact.header.location:
        contact_values.append(artifact.header.location)
    contact_values.extend(item for item in artifact.header.contact_lines if str(item or "").strip())
    contact_html = _build_resume_contact_inline_html(contact_values)

    summary_block = """
    <section class="resume-classic-section resume-classic-section--plain-head">
        <h2>Summary</h2>
        <p class="resume-summary">{summary}</p>
    </section>
    """.format(
        summary=html.escape(artifact.professional_summary or "No professional summary generated."),
    )

    skills_block = """
    <section class="resume-classic-section resume-classic-section--plain-head">
        <h2>Core Skills</h2>
        {skills}
    </section>
    """.format(
        skills=_build_resume_skills_inline_html(
            artifact.highlighted_skills,
            skill_categories=getattr(artifact, "skill_categories", None) or None,
        )
    )

    experience_entries = list(artifact.experience_entries or [])
    experience_block = ""
    if experience_entries:
        experience_block = """
    <section class="resume-classic-section">
        <h2>Experience</h2>
        {experience}
    </section>
        """.format(experience=_build_resume_experience_html(experience_entries))

    project_entries = list(artifact.project_entries or [])
    projects_block = ""
    if project_entries:
        projects_block = """
    <section class="resume-classic-section">
        <h2>Projects</h2>
        {projects}
    </section>
        """.format(projects=_build_resume_projects_html(project_entries))

    education_block = """
    <section class="resume-classic-section">
        <h2>Education</h2>
        {education}
    </section>
    """.format(education=_build_resume_education_html(artifact.education_entries))

    publication_entries = [item for item in (artifact.publication_entries or []) if str(item or "").strip()]
    publications_block = ""
    if publication_entries:
        publications_block = """
    <section class="resume-classic-section">
        <h2>Publications</h2>
        {publications}
    </section>
        """.format(
            publications=_build_resume_section_list(
                publication_entries,
                "",
                class_name="resume-plain-list",
            )
        )

    certifications = [item for item in artifact.certifications if str(item or "").strip()]
    certifications_block = ""
    if certifications:
        certifications_block = """
    <section class="resume-classic-section">
        <h2>Certifications</h2>
        {certifications}
    </section>
        """.format(
            certifications=_build_resume_section_list(
                certifications,
                "",
                class_name="resume-plain-list",
            )
        )

    section_blocks = {
        "summary": summary_block,
        "skills": skills_block,
        "experience": experience_block,
        "projects": projects_block,
        "education": education_block,
        "publications": publications_block,
        "certifications": certifications_block,
    }

    order = list(artifact.section_order) if artifact.section_order else list(_DEFAULT_RESUME_SECTION_ORDER)
    seen: set[str] = set()
    ordered_blocks: list[str] = []
    for section_name in order:
        if section_name in seen:
            continue
        seen.add(section_name)
        block = section_blocks.get(section_name, "")
        if block:
            ordered_blocks.append(block)
    # Append any sections the agent forgot to mention so we never lose
    # rendered content when the agent emits a partial order.
    for section_name in _DEFAULT_RESUME_SECTION_ORDER:
        if section_name in seen:
            continue
        block = section_blocks.get(section_name, "")
        if block:
            ordered_blocks.append(block)

    # Mode-aware headline line. JD-tailored runs set
    # `artifact.target_role` (= the JD title) → render it as the
    # role line. The no-JD / resume-builder path leaves it "" →
    # the line is OMITTED entirely (a name-only header is standard;
    # we never fabricate a headline). Byte-identical to the prior
    # output whenever target_role is empty. Reuses the dormant
    # `.resume-classic-role` style already in the template.
    headline = html.escape(str(getattr(artifact, "target_role", "") or "").strip())
    # When empty, role_block is "" and the line below collapses to
    # exactly the original `        {contact_block}` — byte-identical
    # to pre-feature output for the no-JD / fixture path.
    role_block = (
        '<p class="resume-classic-role">{0}</p>\n        '.format(headline)
        if headline
        else ""
    )
    # Non-banded themes get the EXACT original markup
    # (`class="resume-classic-header"`) — byte-identical, so the 3
    # ATS-simple themes + every non-opting theme are unchanged.
    band_cls = " resume-classic-header--band" if header_banded else ""
    header_html = """
    <section class="resume-classic-header{band_cls}">
        <h1>{name}</h1>
        {role_block}{contact_block}
    </section>
    """.format(
        band_cls=band_cls,
        name=name,
        role_block=role_block,
        contact_block=contact_html,
    )

    return header_html + "\n".join(ordered_blocks)


# ===========================================================================
# TWO-COLUMN RÉSUMÉ THEMES (ADR-032)
#
# Six bespoke designer templates (resume_builder/*.html), each rendered by its
# own function and selected from `_TWOCOL_RENDERERS` by `ThemeSpec.twocol_layout`.
# `_build_resume_html` dispatches here for any `layout == "two_column"` theme.
#
# Each renderer emits a full self-contained HTML document: the design's own
# `<style>` block VERBATIM (the design is the colour — these are not
# re-palettised through ThemeSpec colour fields) plus the template DOM bound
# to artifact data. DOM order is ALWAYS header → main → sidebar so the PDF
# text layer extracts as a coherent linear read despite the visual columns
# (the realistic tolerance ceiling for a deliberately non-ATS layout). Every
# user value is `html.escape`-d. Empty sections drop entirely — no orphan
# headers, no crashes — and the templates keep `page-break-inside: avoid` so
# long résumés paginate without splitting an entry.
# ===========================================================================

# Section split shared by every two-column theme. Wide MAIN = the substance a
# reader works through; SIDEBAR = scannable reference.
_TWOCOL_MAIN_SECTIONS = ("summary", "experience", "projects", "publications")
_TWOCOL_SIDEBAR_SECTIONS = ("skills", "education", "certifications")


def _esc(value) -> str:
    """html.escape on a stringified, stripped value."""
    return html.escape(str(value or "").strip())


def _twocol_name_parts(artifact: TailoredResumeArtifact) -> tuple[str, str]:
    """Split the candidate name into (first-words, last-word) for the
    mixed-weight mastheads. A single-token name returns (name, "")."""
    full = str(artifact.header.full_name or artifact.title or "Candidate").strip()
    tokens = full.split()
    if len(tokens) <= 1:
        return _esc(full), ""
    return _esc(" ".join(tokens[:-1])), _esc(tokens[-1])


def _twocol_contact_items(artifact: TailoredResumeArtifact) -> list[tuple[str, str]]:
    """(label, value) contact rows for the sidebar. Location, then every
    contact line. The label is a coarse type guess used by the templates
    that show an uppercase tag above each value; templates that don't
    show tags just use the value."""
    items: list[tuple[str, str]] = []
    location = str(artifact.header.location or "").strip()
    if location:
        items.append(("Location", location))
    for raw in artifact.header.contact_lines:
        value = str(raw or "").strip()
        if not value:
            continue
        low = value.lower()
        if "@" in value and "/" not in value:
            label = "Email"
        elif low.startswith(("http://", "https://", "www.")) or _looks_like_contact_link(value):
            if "linkedin" in low:
                label = "LinkedIn"
            elif "github" in low:
                label = "GitHub"
            else:
                label = "Profile"
        elif any(ch.isdigit() for ch in value) and sum(ch.isdigit() for ch in value) >= 6:
            label = "Phone"
        else:
            label = "Detail"
        items.append((label, value))
    return items


def _twocol_skill_groups(artifact: TailoredResumeArtifact) -> list[tuple[str, list[str]]]:
    """Skills as (group-label, items) pairs. Uses skill_categories when
    present (one group per category); otherwise a single 'Skills' group
    from highlighted_skills. Empty → []."""
    categories = getattr(artifact, "skill_categories", None) or None
    groups: list[tuple[str, list[str]]] = []
    if categories:
        for label, items in categories.items():
            label_clean = str(label or "").strip()
            cleaned = [str(i or "").strip() for i in items if str(i or "").strip()]
            if label_clean and cleaned:
                groups.append((label_clean, cleaned))
        if groups:
            return groups
    flat = [str(i or "").strip() for i in artifact.highlighted_skills if str(i or "").strip()]
    if flat:
        groups.append(("Skills", flat))
    return groups


def _twocol_section_present(artifact: TailoredResumeArtifact) -> dict[str, bool]:
    """Which optional sections have content. Summary, skills, education
    are treated as always-present (the workflow always fills them);
    experience / projects / publications / certifications drop when
    empty."""
    return {
        "summary": True,
        "skills": bool(_twocol_skill_groups(artifact)),
        "experience": bool(artifact.experience_entries),
        "projects": bool(artifact.project_entries),
        "education": bool(artifact.education_entries),
        "publications": bool(
            [p for p in (artifact.publication_entries or []) if str(p or "").strip()]
        ),
        "certifications": bool(
            [c for c in (artifact.certifications or []) if str(c or "").strip()]
        ),
    }


def _twocol_entry_dates(entry) -> str:
    parts = [p for p in [getattr(entry, "start", ""), getattr(entry, "end", "")] if p]
    return _esc(" — ".join(parts)) if parts else ""


def _twocol_role_headline(artifact: TailoredResumeArtifact) -> str:
    """The JD-tailored role line (target_role). "" on the no-JD path —
    a name-only masthead is standard; the role line is never fabricated."""
    return _esc(getattr(artifact, "target_role", "") or "")


# --- shared EXPERIENCE structure builders -----------------------------------
# Three structures across the six templates; each is built once and the
# templates that share a structure reuse it. Every builder escapes user data
# and keeps each entry in a `.entry` block (the templates' CSS pins
# `page-break-inside: avoid` on `.entry`).


def _twocol_bullets(entry) -> str:
    bullets = [_esc(b) for b in (getattr(entry, "bullets", None) or []) if str(b or "").strip()]
    if not bullets:
        return ""
    return "<ul>{0}</ul>".format("".join("<li>{0}</li>".format(b) for b in bullets))


def _twocol_experience_timeline(entries) -> str:
    """Dot-and-rail timeline (template 01). Entries wrapped in `.timeline`."""
    rows = []
    for entry in entries:
        title = _esc(getattr(entry, "title", "") or "Experience")
        org_parts = [
            p for p in [getattr(entry, "organization", ""), getattr(entry, "location", "")] if p
        ]
        org = _esc(" · ".join(org_parts))
        dates = _twocol_entry_dates(entry)
        rows.append(
            '<div class="entry">'
            '<div class="entry-head"><div>'
            '<p class="entry-title">{title}</p>'
            '{org}'
            '</div>{dates}</div>'
            '{bullets}'
            '</div>'.format(
                title=title,
                org='<p class="entry-org">{0}</p>'.format(org) if org else "",
                dates='<span class="entry-dates">{0}</span>'.format(dates) if dates else "",
                bullets=_twocol_bullets(entry),
            )
        )
    return '<div class="timeline">{0}</div>'.format("".join(rows))


def _twocol_experience_headrow(entries) -> str:
    """Head-row experience (templates 02, 05, 08): title + org on the
    left, dates on the right, bullets below. No wrapping container."""
    rows = []
    for entry in entries:
        title = _esc(getattr(entry, "title", "") or "Experience")
        org_parts = [
            p for p in [getattr(entry, "organization", ""), getattr(entry, "location", "")] if p
        ]
        org = _esc(" · ".join(org_parts))
        dates = _twocol_entry_dates(entry)
        rows.append(
            '<div class="entry">'
            '<div class="entry-head"><div>'
            '<p class="entry-title">{title}</p>'
            '{org}'
            '</div>{dates}</div>'
            '{bullets}'
            '</div>'.format(
                title=title,
                org='<p class="entry-org">{0}</p>'.format(org) if org else "",
                dates='<span class="entry-dates">{0}</span>'.format(dates) if dates else "",
                bullets=_twocol_bullets(entry),
            )
        )
    return "".join(rows)


def _twocol_experience_gutter(entries) -> str:
    """Left date-gutter experience (templates 04, 10): dates in a fixed
    gutter, title/org/bullets in the body."""
    rows = []
    for entry in entries:
        title = _esc(getattr(entry, "title", "") or "Experience")
        org_parts = [
            p for p in [getattr(entry, "organization", ""), getattr(entry, "location", "")] if p
        ]
        org = _esc(" · ".join(org_parts))
        dates = _twocol_entry_dates(entry)
        rows.append(
            '<div class="entry">'
            '<div class="entry-gutter"><div class="entry-dates">{dates}</div></div>'
            '<div class="entry-body">'
            '<p class="entry-title">{title}</p>'
            '{org}'
            '{bullets}'
            '</div>'
            '</div>'.format(
                dates=dates,
                title=title,
                org='<p class="entry-org">{0}</p>'.format(org) if org else "",
                bullets=_twocol_bullets(entry),
            )
        )
    return "".join(rows)


# --- shared SIDEBAR block builders ------------------------------------------


def _twocol_education_blocks(artifact: TailoredResumeArtifact) -> str:
    """`.edu-entry` blocks for the sidebar."""
    blocks = []
    for entry in artifact.education_entries or []:
        degree_parts = [p for p in [getattr(entry, "degree", ""), getattr(entry, "field_of_study", "")] if p]
        degree = _esc(" — ".join(degree_parts)) or _esc(getattr(entry, "institution", "") or "Education")
        institution = _esc(getattr(entry, "institution", ""))
        dates = _twocol_entry_dates(entry)
        # If the degree line fell back to the institution, don't repeat it.
        inst_line = (
            '<p class="edu-inst">{0}</p>'.format(institution)
            if institution and degree_parts
            else ""
        )
        blocks.append(
            '<div class="edu-entry">'
            '<p class="edu-degree">{degree}</p>'
            '{inst}'
            '{dates}'
            '</div>'.format(
                degree=degree,
                inst=inst_line,
                dates='<p class="edu-dates">{0}</p>'.format(dates) if dates else "",
            )
        )
    return "".join(blocks)


def _twocol_skill_groups_html(artifact: TailoredResumeArtifact) -> str:
    """`.skill-group` blocks (group heading + `.skill-list`) for the sidebar.

    When skills are uncategorized there is exactly one group with the
    generic "Skills" label — its heading is suppressed (it would just
    repeat the section header). Real category labels always show."""
    groups = _twocol_skill_groups(artifact)
    suppress_heading = len(groups) == 1 and groups[0][0] == "Skills"
    out = []
    for label, items in groups:
        lis = "".join("<li>{0}</li>".format(_esc(i)) for i in items)
        heading = (
            ""
            if suppress_heading
            else '<h3 class="skill-group-h">{0}</h3>'.format(_esc(label))
        )
        out.append(
            '<div class="skill-group">'
            '{heading}'
            '<ul class="skill-list">{lis}</ul>'
            '</div>'.format(heading=heading, lis=lis)
        )
    return "".join(out)


def _twocol_cert_list_html(artifact: TailoredResumeArtifact) -> str:
    """`.cert-list` items for the sidebar."""
    certs = [c for c in (artifact.certifications or []) if str(c or "").strip()]
    if not certs:
        return ""
    lis = "".join("<li>{0}</li>".format(_esc(c)) for c in certs)
    return '<ul class="cert-list">{0}</ul>'.format(lis)


def _twocol_pubs_html(artifact: TailoredResumeArtifact) -> str:
    """`.pubs` list for the main column."""
    pubs = [p for p in (artifact.publication_entries or []) if str(p or "").strip()]
    if not pubs:
        return ""
    lis = "".join("<li>{0}</li>".format(_esc(p)) for p in pubs)
    return '<ul class="pubs">{0}</ul>'.format(lis)


def _twocol_summary_html(artifact: TailoredResumeArtifact) -> str:
    return "<p>{0}</p>".format(
        _esc(artifact.professional_summary or "No professional summary generated.")
    )


def _twocol_projects_chips(artifact: TailoredResumeArtifact) -> str:
    """Projects rendered with a `.tech-list` of `.tech` chips (templates
    01, 05, 08)."""
    out = []
    for project in artifact.project_entries or []:
        name = _esc(getattr(project, "name", "") or "Project")
        link = _esc(getattr(project, "link", ""))
        description = _esc(getattr(project, "description", ""))
        techs = [_esc(t) for t in (getattr(project, "technologies", None) or []) if str(t or "").strip()]
        tech_html = (
            '<div class="tech-list">{0}</div>'.format(
                "".join('<span class="tech">{0}</span>'.format(t) for t in techs)
            )
            if techs
            else ""
        )
        out.append(
            '<div class="project">'
            '<div class="project-head">'
            '<span class="project-name">{name}</span>'
            '{link}'
            '</div>'
            '{description}'
            '{tech}'
            '</div>'.format(
                name=name,
                link='<a class="project-link" href="#">{0}</a>'.format(link) if link else "",
                description='<p class="project-desc">{0}</p>'.format(description) if description else "",
                tech=tech_html,
            )
        )
    return "".join(out)


def _twocol_projects_split(artifact: TailoredResumeArtifact) -> str:
    """Projects with a two-column `.project-side` (name + link) /
    `.project-body` (desc + inline `.tech-line`) split (template 02)."""
    out = []
    for project in artifact.project_entries or []:
        name = _esc(getattr(project, "name", "") or "Project")
        link = _esc(getattr(project, "link", ""))
        description = _esc(getattr(project, "description", ""))
        techs = [_esc(t) for t in (getattr(project, "technologies", None) or []) if str(t or "").strip()]
        tech_html = (
            '<div class="tech-line"><span class="tech-label">Stack</span>{0}</div>'.format(
                " · ".join(techs)
            )
            if techs
            else ""
        )
        out.append(
            '<div class="project">'
            '<div class="project-side">'
            '<p class="project-name">{name}</p>'
            '{link}'
            '</div>'
            '<div class="project-body">'
            '{description}'
            '{tech}'
            '</div>'
            '</div>'.format(
                name=name,
                link='<a class="project-link" href="#">{0}</a>'.format(link) if link else "",
                description='<p class="project-desc">{0}</p>'.format(description) if description else "",
                tech=tech_html,
            )
        )
    return "".join(out)


def _twocol_projects_gutter(artifact: TailoredResumeArtifact) -> str:
    """Projects with a left `.project-gutter` label + `.project-body`
    (name with inline link, desc, inline `.tech-line`) — templates 04, 10.
    The gutter label is the first technology, or 'Project' when none."""
    out = []
    for project in artifact.project_entries or []:
        name = _esc(getattr(project, "name", "") or "Project")
        link = _esc(getattr(project, "link", ""))
        description = _esc(getattr(project, "description", ""))
        techs = [_esc(t) for t in (getattr(project, "technologies", None) or []) if str(t or "").strip()]
        gutter_label = techs[0] if techs else "Project"
        tech_html = (
            '<div class="tech-line"><span class="tech-label">Stack</span>{0}</div>'.format(
                " · ".join(techs)
            )
            if techs
            else ""
        )
        out.append(
            '<div class="project">'
            '<div class="project-gutter"><span class="gutter-label">{label}</span></div>'
            '<div class="project-body">'
            '<p class="project-name">{name}{link}</p>'
            '{description}'
            '{tech}'
            '</div>'
            '</div>'.format(
                label=gutter_label,
                name=name,
                link=' <a class="project-link" href="#">{0}</a>'.format(link) if link else "",
                description='<p class="project-desc">{0}</p>'.format(description) if description else "",
                tech=tech_html,
            )
        )
    return "".join(out)


# --- shared section-ordering helper -----------------------------------------


def _twocol_ordered(artifact: TailoredResumeArtifact, allowed: tuple) -> list[str]:
    """Section keys from the artifact's order, filtered to `allowed` and
    to sections that have content. Honors artifact.section_order, then
    appends any forgotten sections in the default order."""
    present = _twocol_section_present(artifact)
    order = list(artifact.section_order) if artifact.section_order else list(_DEFAULT_RESUME_SECTION_ORDER)
    seen: set[str] = set()
    result: list[str] = []
    for name in list(order) + list(_DEFAULT_RESUME_SECTION_ORDER):
        if name in seen:
            continue
        seen.add(name)
        if name in allowed and present.get(name):
            result.append(name)
    return result


def _twocol_document(title: str, style: str, body: str) -> str:
    """Wrap a two-column theme's `<style>` + body in an A4 HTML document."""
    return (
        '<!DOCTYPE html>\n<html lang="en">\n<head>\n'
        '<meta charset="utf-8" />\n<title>{title}</title>\n'
        '<style>{style}</style>\n</head>\n<body>\n{body}\n</body>\n</html>\n'
    ).format(title=_esc(title or "Tailored Resume"), style=style, body=body)


def _twocol_main_sections_html(
    artifact: TailoredResumeArtifact,
    section_header,
    experience_builder,
    projects_builder,
) -> str:
    """Assemble the MAIN column (summary / experience / projects /
    publications) in artifact order, dropping empty sections.

    `section_header(label)` returns the theme's `<h2 class="section-h">`
    markup; the experience / projects builders are the structure
    functions the template uses. Each section is a `<section
    class="section">` so the templates' `.section` spacing applies."""
    parts = []
    for name in _twocol_ordered(artifact, _TWOCOL_MAIN_SECTIONS):
        if name == "summary":
            inner = '<div class="summary">{0}</div>'.format(_twocol_summary_html(artifact))
            label = "Summary"
        elif name == "experience":
            inner = experience_builder(artifact.experience_entries or [])
            label = "Experience"
        elif name == "projects":
            inner = projects_builder(artifact)
            label = "Projects"
        elif name == "publications":
            inner = _twocol_pubs_html(artifact)
            label = "Publications"
        else:  # pragma: no cover - _twocol_ordered filters to the four
            continue
        parts.append(
            '<section class="section">{header}{inner}</section>'.format(
                header=section_header(label), inner=inner
            )
        )
    return "".join(parts)


def _twocol_sidebar_sections_html(
    artifact: TailoredResumeArtifact, section_header
) -> str:
    """Assemble the SIDEBAR's content sections (skills / education /
    certifications) in artifact order, dropping empty sections. Contact
    is rendered separately by each template (its markup differs)."""
    parts = []
    for name in _twocol_ordered(artifact, _TWOCOL_SIDEBAR_SECTIONS):
        if name == "skills":
            inner = _twocol_skill_groups_html(artifact)
            label = "Skills"
        elif name == "education":
            inner = _twocol_education_blocks(artifact)
            label = "Education"
        elif name == "certifications":
            inner = _twocol_cert_list_html(artifact)
            label = "Certifications"
        else:  # pragma: no cover
            continue
        parts.append(
            '<section class="sidebar-section">{header}{inner}</section>'.format(
                header=section_header(label), inner=inner
            )
        )
    return "".join(parts)


# --- TEMPLATE 01 — Timeline / Tech ------------------------------------------
_TWOCOL_STYLE_TIMELINE_TECH = """
:root{--font-sans:Arial,"Helvetica Neue",Helvetica,sans-serif;--ink:#111418;--muted:#5b6470;--line:#e4e7ec;--accent:#4f6bed;--accent-soft:#eef1ff;--paper:#ffffff;--surface:#ffffff;--sidebar-bg:#0d1220;--sidebar-fg:#e7ebf3;--sidebar-mute:#8e96a8;--sidebar-rule:#1f273b;--sidebar-w:34%;--main-pad:16mm;--side-pad:11mm;}
@page{size:A4;margin:0;}
*{box-sizing:border-box;}
html,body{margin:0;padding:0;background:var(--paper);}
body{font-family:var(--font-sans);color:var(--ink);font-size:10pt;line-height:1.45;-webkit-font-smoothing:antialiased;}
.sheet{display:flex;width:210mm;min-height:297mm;margin:0 auto;background:var(--paper);}
.sidebar{width:var(--sidebar-w);background:var(--sidebar-bg);color:var(--sidebar-fg);padding:var(--side-pad);display:flex;flex-direction:column;order:-1;}
.main{flex:1;background:var(--surface);padding:var(--main-pad);display:flex;flex-direction:column;}
.masthead{padding-bottom:9mm;margin-bottom:8mm;border-bottom:1px solid var(--line);}
.masthead .name{font-size:30pt;line-height:1.02;letter-spacing:-0.5px;margin:0 0 4pt 0;}
.masthead .name .first{font-weight:300;color:var(--ink);}
.masthead .name .last{font-weight:800;color:var(--ink);}
.masthead .role{font-size:10pt;font-weight:700;letter-spacing:3px;text-transform:uppercase;color:var(--accent);margin-top:6pt;}
.section{margin-bottom:7mm;}
.section:last-child{margin-bottom:0;}
.section-h{display:flex;align-items:center;font-size:9.5pt;font-weight:800;letter-spacing:2.4px;text-transform:uppercase;color:var(--ink);margin:0 0 4mm 0;}
.section-h::before{content:"";display:inline-block;width:14px;height:2px;background:var(--accent);margin-right:8px;}
.sidebar .section-h{color:var(--sidebar-fg);}
.sidebar .section-h::before{background:var(--accent);}
.summary p{margin:0;font-size:10pt;line-height:1.55;color:var(--ink);}
.timeline{position:relative;padding-left:14px;border-left:1px solid var(--line);}
.entry{position:relative;padding:0 0 5mm 10px;page-break-inside:avoid;}
.entry:last-child{padding-bottom:0;}
.entry::before{content:"";position:absolute;left:-19px;top:5px;width:9px;height:9px;border-radius:50%;background:var(--accent);box-shadow:0 0 0 3px var(--paper);}
.entry-head{display:flex;justify-content:space-between;align-items:baseline;gap:10px;margin-bottom:1mm;}.entry-head>div{flex:1;min-width:0;}.entry-head .entry-dates{flex-shrink:0;}
.entry-title{font-size:11pt;font-weight:700;color:var(--ink);margin:0;}
.entry-org{font-size:10pt;color:var(--muted);margin:0;}
.entry-dates{font-size:9pt;font-weight:600;color:var(--accent);white-space:nowrap;letter-spacing:0.3px;}
.entry ul{margin:2mm 0 0 0;padding-left:14px;font-size:9.7pt;color:var(--ink);}
.entry ul li{margin-bottom:1.2mm;line-height:1.45;}
.project{page-break-inside:avoid;margin-bottom:4mm;}
.project:last-child{margin-bottom:0;}
.project-head{display:flex;align-items:baseline;gap:8px;margin-bottom:1mm;}
.project-name{font-size:10.5pt;font-weight:700;color:var(--ink);}
.project-link{font-size:9pt;color:var(--accent);text-decoration:none;}
.project-desc{font-size:9.7pt;margin:0 0 1mm 0;color:var(--ink);}
.tech-list{display:flex;flex-wrap:wrap;gap:5px;margin-top:1mm;}
.tech{font-size:8.5pt;font-weight:600;color:var(--accent);background:var(--accent-soft);padding:2px 7px;border-radius:3px;}
.pubs{margin:0;padding:0;list-style:none;}
.pubs li{font-size:9.5pt;color:var(--ink);line-height:1.45;padding-left:12px;position:relative;margin-bottom:1.8mm;}
.pubs li::before{content:"\\2014";position:absolute;left:0;color:var(--accent);}
.brand{margin-bottom:9mm;padding-bottom:7mm;border-bottom:1px solid var(--sidebar-rule);}
.brand .monogram{width:38px;height:38px;border-radius:50%;border:1.5px solid var(--accent);display:flex;align-items:center;justify-content:center;font-size:14pt;font-weight:700;color:var(--accent);letter-spacing:0.5px;}
.contact-list{list-style:none;padding:0;margin:0;}
.contact-list li{font-size:9pt;color:var(--sidebar-fg);margin-bottom:2.2mm;word-break:break-word;line-height:1.4;}
.contact-list .label{display:block;font-size:7.5pt;letter-spacing:1.6px;text-transform:uppercase;color:var(--sidebar-mute);margin-bottom:1px;}
.skill-group{margin-bottom:4mm;}
.skill-group:last-child{margin-bottom:0;}
.skill-group-h{font-size:7.8pt;font-weight:700;letter-spacing:1.8px;text-transform:uppercase;color:var(--accent);margin:0 0 2mm 0;}
.skill-list{list-style:none;padding:0;margin:0;}
.skill-list li{font-size:9.3pt;color:var(--sidebar-fg);padding:1.3mm 0;border-bottom:1px solid var(--sidebar-rule);}
.skill-list li:last-child{border-bottom:none;}
.edu-entry{margin-bottom:4mm;}
.edu-entry:last-child{margin-bottom:0;}
.edu-degree{font-size:9.7pt;font-weight:700;color:var(--sidebar-fg);margin:0;line-height:1.3;}
.edu-inst{font-size:9pt;color:var(--sidebar-mute);margin:1px 0 0 0;}
.edu-dates{font-size:8.4pt;color:var(--accent);font-weight:600;margin-top:2px;letter-spacing:0.3px;}
.cert-list{list-style:none;padding:0;margin:0;}
.cert-list li{font-size:9pt;color:var(--sidebar-fg);margin-bottom:2.5mm;line-height:1.35;padding-left:10px;position:relative;}
.cert-list li::before{content:"";position:absolute;left:0;top:7px;width:4px;height:4px;background:var(--accent);border-radius:50%;}
.sidebar-section{margin-bottom:8mm;}
.sidebar-section:last-child{margin-bottom:0;}
"""


def _twocol_monogram(artifact: TailoredResumeArtifact) -> str:
    """Two-letter initials for the templates with a monogram badge."""
    first, last = _twocol_name_parts(artifact)
    # Use raw (unescaped) name for letter extraction, then escape the result.
    full = str(artifact.header.full_name or artifact.title or "Candidate").strip()
    tokens = [t for t in full.split() if t]
    if len(tokens) >= 2:
        return _esc((tokens[0][0] + tokens[-1][0]))
    if tokens and tokens[0]:
        return _esc(tokens[0][:2])
    return "CV"


def _render_twocol_timeline_tech(
    artifact: TailoredResumeArtifact, title: str
) -> str:
    """Template 01 — dark navy sidebar, dot-and-rail experience timeline."""

    def section_h(label):
        return '<h2 class="section-h">{0}</h2>'.format(_esc(label))

    role = _twocol_role_headline(artifact)
    first, last = _twocol_name_parts(artifact)
    name_html = '<span class="first">{0}</span>'.format(first) + (
        ' <span class="last">{0}</span>'.format(last) if last else ""
    )
    contact_lis = "".join(
        '<li><span class="label">{label}</span>{value}</li>'.format(
            label=_esc(label), value=_esc(value)
        )
        for label, value in _twocol_contact_items(artifact)
    )
    sidebar_contact = (
        '<section class="sidebar-section">'
        '<h2 class="section-h">Contact</h2>'
        '<ul class="contact-list">{0}</ul>'
        '</section>'.format(contact_lis)
        if contact_lis
        else ""
    )
    main = _twocol_main_sections_html(
        artifact, section_h, _twocol_experience_timeline, _twocol_projects_chips
    )
    sidebar_body = _twocol_sidebar_sections_html(artifact, section_h)
    body = (
        '<div class="sheet">'
        '<main class="main">'
        '<header class="masthead">'
        '<h1 class="name">{name}</h1>'
        '{role}'
        '</header>'
        '{main}'
        '</main>'
        '<aside class="sidebar">'
        '<div class="brand"><div class="monogram">{monogram}</div></div>'
        '{contact}{sidebar}'
        '</aside>'
        '</div>'
    ).format(
        monogram=_twocol_monogram(artifact),
        contact=sidebar_contact,
        sidebar=sidebar_body,
        name=name_html,
        role='<div class="role">{0}</div>'.format(role) if role else "",
        main=main,
    )
    return _twocol_document(title, _TWOCOL_STYLE_TIMELINE_TECH, body)


# --- TEMPLATE 02 — Editorial Minimal ----------------------------------------
_TWOCOL_STYLE_EDITORIAL_MINIMAL = """
:root{--font-sans:Arial,"Helvetica Neue",Helvetica,sans-serif;--ink:#1c1a17;--muted:#6b6760;--line:#d9d4cb;--accent:#9a4a2a;--accent-soft:#f1e7df;--paper:#fbf7f0;--surface:#fbf7f0;--sidebar-bg:#efe7d8;--sidebar-fg:#1c1a17;--sidebar-mute:#7a7468;--sidebar-rule:#d4ccba;--sidebar-w:35%;--main-pad:18mm;--side-pad:13mm;}
@page{size:A4;margin:0;}
*{box-sizing:border-box;}
html,body{margin:0;padding:0;background:var(--paper);}
body{font-family:var(--font-sans);color:var(--ink);font-size:10pt;line-height:1.5;-webkit-font-smoothing:antialiased;}
.sheet{display:flex;width:210mm;min-height:297mm;margin:0 auto;background:var(--paper);}
.sidebar{width:var(--sidebar-w);background:var(--sidebar-bg);color:var(--sidebar-fg);padding:var(--side-pad);display:flex;flex-direction:column;order:-1;}
.main{flex:1;background:var(--surface);padding:var(--main-pad);display:flex;flex-direction:column;}
.masthead{margin-bottom:11mm;}
.masthead .eyebrow{font-size:8.5pt;font-weight:700;letter-spacing:3px;text-transform:uppercase;color:var(--accent);margin-bottom:6mm;}
.masthead .name{font-size:34pt;line-height:0.98;letter-spacing:-1px;margin:0 0 6mm 0;font-weight:300;color:var(--ink);}
.masthead .name .last{display:block;font-weight:800;}
.masthead .role{font-size:11pt;color:var(--ink);border-top:1px solid var(--line);padding-top:5mm;font-weight:400;}
.section{margin-bottom:9mm;}
.section:last-child{margin-bottom:0;}
.section-h{font-size:9pt;font-weight:700;letter-spacing:4px;text-transform:uppercase;color:var(--ink);margin:0 0 4mm 0;padding-bottom:2mm;border-bottom:1px solid var(--line);}
.sidebar .section-h{color:var(--ink);border-bottom-color:var(--sidebar-rule);}
.summary p{margin:0;font-size:10.5pt;line-height:1.6;color:var(--ink);}
.entry{padding:0 0 5mm 0;page-break-inside:avoid;}
.entry:last-child{padding-bottom:0;}
.entry-head{display:flex;justify-content:space-between;align-items:baseline;gap:10px;margin-bottom:1mm;}.entry-head>div{flex:1;min-width:0;}.entry-head .entry-dates{flex-shrink:0;}
.entry-title{font-size:11pt;font-weight:700;color:var(--ink);margin:0;}
.entry-org{font-size:10pt;color:var(--accent);margin:0;font-weight:600;}
.entry-dates{font-size:9pt;color:var(--muted);white-space:nowrap;font-style:italic;}
.entry ul{margin:2mm 0 0 0;padding-left:16px;font-size:10pt;color:var(--ink);}
.entry ul li{margin-bottom:1.4mm;line-height:1.5;}
.project{page-break-inside:avoid;margin-bottom:4mm;display:flex;gap:6mm;}
.project:last-child{margin-bottom:0;}
.project-side{width:32%;flex-shrink:0;}
.project-name{font-size:10.5pt;font-weight:700;color:var(--ink);margin:0 0 1mm 0;}
.project-link{font-size:9pt;color:var(--accent);text-decoration:none;word-break:break-word;}
.project-body{flex:1;}
.project-desc{font-size:9.8pt;margin:0 0 1.5mm 0;color:var(--ink);line-height:1.45;}
.tech-line{font-size:8.8pt;color:var(--muted);letter-spacing:0.3px;}
.tech-line .tech-label{font-weight:700;text-transform:uppercase;letter-spacing:1.5px;margin-right:6px;color:var(--accent);}
.pubs{margin:0;padding:0;list-style:none;}
.pubs li{font-size:9.7pt;color:var(--ink);line-height:1.5;margin-bottom:2.5mm;padding-bottom:2.5mm;border-bottom:1px solid var(--line);}
.pubs li:last-child{border-bottom:none;padding-bottom:0;}
.sidebar-section{margin-bottom:9mm;}
.sidebar-section:last-child{margin-bottom:0;}
.contact-list{list-style:none;padding:0;margin:0;}
.contact-list li{font-size:9.3pt;color:var(--sidebar-fg);margin-bottom:2.5mm;word-break:break-word;line-height:1.4;}
.skill-group{margin-bottom:4mm;}
.skill-group:last-child{margin-bottom:0;}
.skill-group-h{font-size:8pt;font-weight:700;letter-spacing:2px;text-transform:uppercase;color:var(--accent);margin:0 0 2mm 0;}
.skill-list{list-style:none;padding:0;margin:0;}
.skill-list li{font-size:9.5pt;color:var(--sidebar-fg);padding:1.1mm 0;line-height:1.35;}
.edu-entry{margin-bottom:4mm;padding-bottom:4mm;border-bottom:1px solid var(--sidebar-rule);}
.edu-entry:last-child{border-bottom:none;padding-bottom:0;margin-bottom:0;}
.edu-degree{font-size:10pt;font-weight:700;color:var(--sidebar-fg);margin:0;line-height:1.3;}
.edu-inst{font-size:9.3pt;color:var(--sidebar-fg);margin:1px 0 0 0;}
.edu-dates{font-size:8.7pt;color:var(--sidebar-mute);margin-top:2px;font-style:italic;}
.cert-list{list-style:none;padding:0;margin:0;}
.cert-list li{font-size:9.3pt;color:var(--sidebar-fg);margin-bottom:2.5mm;line-height:1.4;}
"""


def _render_twocol_editorial_minimal(
    artifact: TailoredResumeArtifact, title: str
) -> str:
    """Template 02 - light sand sidebar, terracotta, hairline-rule headers."""

    def section_h(label):
        return '<h2 class="section-h">{0}</h2>'.format(_esc(label))

    role = _twocol_role_headline(artifact)
    first, last = _twocol_name_parts(artifact)
    name_html = first + (' <span class="last">{0}</span>'.format(last) if last else "")
    contact_lis = "".join(
        "<li>{0}</li>".format(_esc(value)) for _, value in _twocol_contact_items(artifact)
    )
    sidebar_contact = (
        '<section class="sidebar-section">'
        '<h2 class="section-h">Contact</h2>'
        '<ul class="contact-list">{0}</ul>'
        '</section>'.format(contact_lis)
        if contact_lis
        else ""
    )
    main = _twocol_main_sections_html(
        artifact, section_h, _twocol_experience_headrow, _twocol_projects_split
    )
    sidebar_body = _twocol_sidebar_sections_html(artifact, section_h)
    body = (
        '<div class="sheet">'
        '<main class="main">'
        '<header class="masthead">'
        '<div class="eyebrow">Curriculum Vitae</div>'
        '<h1 class="name">{name}</h1>'
        '{role}'
        '</header>'
        '{main}'
        '</main>'
        '<aside class="sidebar">{contact}{sidebar}</aside>'
        '</div>'
    ).format(
        contact=sidebar_contact,
        sidebar=sidebar_body,
        name=name_html,
        role='<div class="role">{0}</div>'.format(role) if role else "",
        main=main,
    )
    return _twocol_document(title, _TWOCOL_STYLE_EDITORIAL_MINIMAL, body)


# --- TEMPLATE 04 — Classic Slate --------------------------------------------
_TWOCOL_STYLE_CLASSIC_SLATE = """
:root{--font-sans:Arial,"Helvetica Neue",Helvetica,sans-serif;--ink:#0f1722;--muted:#586374;--line:#d5dae2;--accent:#1f5d4c;--accent-soft:#e6efeb;--paper:#ffffff;--surface:#ffffff;--sidebar-bg:#eef1f5;--sidebar-fg:#0f1722;--sidebar-mute:#4c5666;--sidebar-rule:#c9d0db;--sidebar-w:33%;--main-pad:17mm;--side-pad:12mm;}
@page{size:A4;margin:0;}
*{box-sizing:border-box;}
html,body{margin:0;padding:0;background:var(--paper);}
body{font-family:var(--font-sans);color:var(--ink);font-size:10pt;line-height:1.5;-webkit-font-smoothing:antialiased;}
.sheet{display:flex;width:210mm;min-height:297mm;margin:0 auto;background:var(--paper);}
.sidebar{width:var(--sidebar-w);background:var(--sidebar-bg);color:var(--sidebar-fg);padding:var(--side-pad);display:flex;flex-direction:column;border-right:2px solid var(--accent);order:-1;}
.main{flex:1;background:var(--surface);padding:var(--main-pad);display:flex;flex-direction:column;}
.masthead{margin-bottom:9mm;padding-bottom:6mm;border-bottom:2px solid var(--ink);display:flex;justify-content:space-between;align-items:flex-end;gap:10mm;}
.masthead-l{flex:1;min-width:0;}
.masthead .name{margin:0;font-size:28pt;line-height:1;font-weight:800;letter-spacing:-0.5px;color:var(--ink);text-transform:uppercase;}
.masthead .name .last{display:block;font-weight:300;letter-spacing:4px;font-size:14pt;margin-top:2mm;color:var(--muted);}
.masthead .role{font-size:10pt;font-weight:700;letter-spacing:2.5px;text-transform:uppercase;color:var(--accent);text-align:right;max-width:50mm;line-height:1.4;overflow-wrap:anywhere;}
.section{margin-bottom:6.5mm;}
.section:last-child{margin-bottom:0;}
.section-h{display:flex;align-items:center;font-size:9.3pt;font-weight:800;letter-spacing:3px;text-transform:uppercase;color:var(--ink);margin:0 0 4mm 0;}
.section-h .label{white-space:nowrap;}
.section-h::after{content:"";flex:1;height:1px;background:var(--ink);margin-left:4mm;}
.sidebar .section-h{color:var(--accent);}
.sidebar .section-h::after{background:var(--accent);}
.summary p{margin:0;font-size:10.2pt;line-height:1.55;}
.entry{display:flex;gap:5mm;page-break-inside:avoid;margin-bottom:5mm;}
.entry:last-child{margin-bottom:0;}
.entry-gutter{width:22mm;flex-shrink:0;padding-top:1mm;}
.entry-dates{font-size:9pt;color:var(--accent);font-weight:700;letter-spacing:0.3px;line-height:1.3;}
.entry-body{flex:1;}
.entry-title{font-size:11pt;font-weight:800;color:var(--ink);margin:0;line-height:1.25;}
.entry-org{font-size:10pt;color:var(--muted);margin:1px 0 0 0;font-style:italic;}
.entry ul{margin:2mm 0 0 0;padding-left:16px;font-size:9.8pt;}
.entry ul li{margin-bottom:1.2mm;line-height:1.45;}
.project{display:flex;gap:5mm;page-break-inside:avoid;margin-bottom:4mm;}
.project:last-child{margin-bottom:0;}
.project-gutter{width:22mm;flex-shrink:0;padding-top:1mm;}
.project-gutter .gutter-label{font-size:8.4pt;color:var(--accent);font-weight:700;letter-spacing:1.6px;text-transform:uppercase;}
.project-body{flex:1;}
.project-name{font-size:10.5pt;font-weight:800;color:var(--ink);margin:0 0 1mm 0;}
.project-link{display:inline-block;font-size:9pt;color:var(--accent);text-decoration:none;margin-left:6px;font-weight:400;font-style:italic;}
.project-desc{font-size:9.8pt;margin:0 0 1.5mm 0;line-height:1.45;}
.tech-line{font-size:8.8pt;color:var(--muted);letter-spacing:0.2px;}
.tech-line .tech-label{font-weight:700;text-transform:uppercase;letter-spacing:1.5px;margin-right:6px;color:var(--ink);}
.pubs{margin:0;padding:0;list-style:none;counter-reset:pubcount;}
.pubs li{font-size:9.5pt;line-height:1.45;margin-bottom:2.5mm;padding-left:9mm;position:relative;counter-increment:pubcount;}
.pubs li::before{content:"[" counter(pubcount) "]";position:absolute;left:0;top:0;font-weight:700;color:var(--accent);font-size:9pt;}
.sidebar-section{margin-bottom:7mm;}
.sidebar-section:last-child{margin-bottom:0;}
.contact-list{list-style:none;padding:0;margin:0;}
.contact-list li{font-size:9.2pt;color:var(--sidebar-fg);margin-bottom:2mm;word-break:break-word;line-height:1.4;}
.contact-list .tag{display:block;font-size:7.5pt;letter-spacing:1.6px;text-transform:uppercase;color:var(--sidebar-mute);margin-bottom:1px;font-weight:700;}
.skill-group{margin-bottom:3.5mm;}
.skill-group:last-child{margin-bottom:0;}
.skill-group-h{font-size:8pt;font-weight:800;letter-spacing:2px;text-transform:uppercase;color:var(--sidebar-mute);margin:0 0 1.5mm 0;padding-bottom:1mm;border-bottom:1px solid var(--sidebar-rule);}
.skill-list{list-style:none;padding:0;margin:0;}
.skill-list li{font-size:9.3pt;color:var(--sidebar-fg);padding:1mm 0;line-height:1.3;}
.edu-entry{margin-bottom:4mm;}
.edu-entry:last-child{margin-bottom:0;}
.edu-degree{font-size:9.8pt;font-weight:800;color:var(--sidebar-fg);margin:0;line-height:1.3;}
.edu-inst{font-size:9.2pt;color:var(--sidebar-fg);margin:1px 0 0 0;font-style:italic;}
.edu-dates{font-size:8.6pt;color:var(--accent);font-weight:700;margin-top:1.5mm;letter-spacing:0.4px;}
.cert-list{list-style:none;padding:0;margin:0;}
.cert-list li{font-size:9.2pt;color:var(--sidebar-fg);margin-bottom:2.5mm;line-height:1.4;padding-left:8mm;position:relative;}
.cert-list li::before{content:"\\2713";position:absolute;left:0;top:0;color:var(--accent);font-weight:800;font-size:10pt;line-height:1.4;}
"""


def _render_twocol_classic_slate(
    artifact: TailoredResumeArtifact, title: str
) -> str:
    """Template 04 - pale slate sidebar, emerald, left date-gutter."""

    def section_h(label):
        return '<h2 class="section-h"><span class="label">{0}</span></h2>'.format(_esc(label))

    role = _twocol_role_headline(artifact)
    first, last = _twocol_name_parts(artifact)
    name_html = first + (' <span class="last">{0}</span>'.format(last) if last else "")
    contact_lis = "".join(
        '<li><span class="tag">{tag}</span>{value}</li>'.format(
            tag=_esc(tag), value=_esc(value)
        )
        for tag, value in _twocol_contact_items(artifact)
    )
    sidebar_contact = (
        '<section class="sidebar-section">'
        '<h2 class="section-h"><span class="label">Contact</span></h2>'
        '<ul class="contact-list">{0}</ul>'
        '</section>'.format(contact_lis)
        if contact_lis
        else ""
    )
    main = _twocol_main_sections_html(
        artifact, section_h, _twocol_experience_gutter, _twocol_projects_gutter
    )
    sidebar_body = _twocol_sidebar_sections_html(artifact, section_h)
    body = (
        '<div class="sheet">'
        '<main class="main">'
        '<header class="masthead">'
        '<div class="masthead-l"><h1 class="name">{name}</h1></div>'
        '{role}'
        '</header>'
        '{main}'
        '</main>'
        '<aside class="sidebar">{contact}{sidebar}</aside>'
        '</div>'
    ).format(
        contact=sidebar_contact,
        sidebar=sidebar_body,
        name=name_html,
        role='<div class="role">{0}</div>'.format(role) if role else "",
        main=main,
    )
    return _twocol_document(title, _TWOCOL_STYLE_CLASSIC_SLATE, body)


# --- TEMPLATE 05 — Monochrome Black -----------------------------------------
_TWOCOL_STYLE_MONOCHROME_BLACK = """
:root{--font-sans:Arial,"Helvetica Neue",Helvetica,sans-serif;--ink:#0a0a0a;--muted:#6e6a63;--line:#d8d2c6;--accent:#0a0a0a;--accent-soft:#ece6d9;--paper:#f3eee3;--surface:#f3eee3;--sidebar-bg:#000000;--sidebar-fg:#f3eee3;--sidebar-mute:#8c8884;--sidebar-rule:#1f1f1f;--sidebar-w:34%;--main-pad:17mm;--side-pad:12mm;}
@page{size:A4;margin:0;}
*{box-sizing:border-box;}
html,body{margin:0;padding:0;background:var(--paper);}
body{font-family:var(--font-sans);color:var(--ink);font-size:10pt;line-height:1.5;-webkit-font-smoothing:antialiased;}
.sheet{display:flex;width:210mm;min-height:297mm;margin:0 auto;background:var(--paper);}
.sidebar{width:var(--sidebar-w);background:var(--sidebar-bg);color:var(--sidebar-fg);padding:var(--side-pad);display:flex;flex-direction:column;order:-1;}
.main{flex:1;background:var(--surface);padding:var(--main-pad);display:flex;flex-direction:column;}
.masthead{margin-bottom:11mm;}
.masthead .index{font-size:8.5pt;font-weight:700;letter-spacing:4px;text-transform:uppercase;color:var(--muted);margin-bottom:6mm;display:flex;justify-content:space-between;align-items:center;}
.masthead .index::after{content:"";flex:1;height:1px;background:var(--ink);margin-left:4mm;}
.masthead .name{margin:0 0 5mm 0;font-size:42pt;line-height:0.92;letter-spacing:-1.5px;color:var(--ink);font-weight:200;}
.masthead .name .last{display:block;font-weight:900;}
.masthead .role{font-size:10.5pt;font-weight:700;letter-spacing:3px;text-transform:uppercase;color:var(--ink);padding-top:4mm;border-top:2px solid var(--ink);display:inline-block;padding-right:12mm;}
.section{margin-bottom:8mm;}
.section:last-child{margin-bottom:0;}
.section-h{font-size:9.5pt;font-weight:900;letter-spacing:3.5px;text-transform:uppercase;color:var(--ink);margin:0 0 4mm 0;padding-bottom:2.5mm;border-bottom:2px solid var(--ink);}
.sidebar .section-h{color:var(--sidebar-fg);border-bottom:2px solid var(--sidebar-fg);}
.summary p{margin:0;font-size:10.5pt;line-height:1.6;color:var(--ink);}
.entry{padding:0 0 5mm 6mm;page-break-inside:avoid;border-left:2px solid var(--ink);margin-left:1mm;position:relative;}
.entry::before{content:"";position:absolute;left:-6px;top:5px;width:10px;height:10px;background:var(--ink);}
.entry:last-child{padding-bottom:0;}
.entry-head{display:flex;justify-content:space-between;align-items:baseline;gap:10px;margin-bottom:0.5mm;}.entry-head>div{flex:1;min-width:0;}.entry-head .entry-dates{flex-shrink:0;}
.entry-title{font-size:11.5pt;font-weight:800;color:var(--ink);margin:0;line-height:1.25;}
.entry-org{font-size:10pt;color:var(--muted);margin:1px 0 0 0;}
.entry-dates{font-size:9pt;color:var(--ink);white-space:nowrap;font-weight:700;letter-spacing:0.5px;}
.entry ul{margin:2mm 0 0 0;padding-left:14px;font-size:10pt;}
.entry ul li{margin-bottom:1.3mm;line-height:1.5;}
.project{page-break-inside:avoid;margin-bottom:4mm;padding-bottom:4mm;border-bottom:1px solid var(--line);}
.project:last-child{margin-bottom:0;padding-bottom:0;border-bottom:none;}
.project-head{display:flex;justify-content:space-between;align-items:baseline;gap:8px;margin-bottom:1mm;}
.project-name{font-size:10.8pt;font-weight:800;color:var(--ink);text-transform:uppercase;letter-spacing:0.5px;}
.project-link{font-size:9pt;color:var(--muted);text-decoration:underline;text-underline-offset:2px;}
.project-desc{font-size:9.8pt;margin:0 0 1.5mm 0;line-height:1.5;}
.tech-list{display:flex;flex-wrap:wrap;gap:6px;margin-top:1.5mm;}
.tech{font-size:8.4pt;font-weight:700;color:var(--ink);border:1px solid var(--ink);padding:1.5px 7px;letter-spacing:0.4px;text-transform:uppercase;}
.pubs{margin:0;padding:0;list-style:none;counter-reset:pubcount;}
.pubs li{font-size:9.5pt;line-height:1.5;margin-bottom:2.5mm;padding-left:11mm;position:relative;counter-increment:pubcount;}
.pubs li::before{content:counter(pubcount,decimal-leading-zero);position:absolute;left:0;top:0;font-weight:900;color:var(--ink);font-size:10pt;letter-spacing:0.5px;}
.sidebar-section{margin-bottom:8mm;}
.sidebar-section:last-child{margin-bottom:0;}
.brand{margin-bottom:10mm;padding-bottom:8mm;border-bottom:2px solid var(--sidebar-fg);}
.brand .monogram{font-size:30pt;font-weight:900;letter-spacing:-1px;line-height:1;color:var(--sidebar-fg);margin-bottom:3mm;}
.brand .tag{font-size:7.8pt;font-weight:700;letter-spacing:2.5px;text-transform:uppercase;color:var(--sidebar-mute);}
.contact-list{list-style:none;padding:0;margin:0;}
.contact-list li{font-size:9.3pt;color:var(--sidebar-fg);margin-bottom:3mm;word-break:break-word;line-height:1.35;}
.contact-list .tag{display:block;font-size:7.5pt;letter-spacing:2px;text-transform:uppercase;color:var(--sidebar-mute);margin-bottom:1px;font-weight:700;}
.skill-group{margin-bottom:4mm;}
.skill-group:last-child{margin-bottom:0;}
.skill-group-h{font-size:8pt;font-weight:800;letter-spacing:2px;text-transform:uppercase;color:var(--sidebar-mute);margin:0 0 2mm 0;}
.skill-list{list-style:none;padding:0;margin:0;}
.skill-list li{font-size:9.4pt;color:var(--sidebar-fg);padding:1.2mm 0;line-height:1.3;border-bottom:1px solid var(--sidebar-rule);}
.skill-list li:last-child{border-bottom:none;}
.edu-entry{margin-bottom:4mm;}
.edu-entry:last-child{margin-bottom:0;}
.edu-degree{font-size:10pt;font-weight:800;color:var(--sidebar-fg);margin:0;line-height:1.3;}
.edu-inst{font-size:9.2pt;color:var(--sidebar-fg);margin:1px 0 0 0;}
.edu-dates{font-size:8.5pt;color:var(--sidebar-mute);margin-top:1.5mm;font-weight:700;letter-spacing:0.5px;}
.cert-list{list-style:none;padding:0;margin:0;counter-reset:certcount;}
.cert-list li{font-size:9.2pt;color:var(--sidebar-fg);margin-bottom:3mm;line-height:1.4;padding-left:9mm;position:relative;counter-increment:certcount;}
.cert-list li::before{content:counter(certcount,decimal-leading-zero);position:absolute;left:0;top:0;font-size:8.5pt;font-weight:800;color:var(--sidebar-mute);letter-spacing:0.5px;}
"""


def _render_twocol_monochrome_black(
    artifact: TailoredResumeArtifact, title: str
) -> str:
    """Template 05 - true-black sidebar, monochrome, head-row experience."""

    def section_h(label):
        return '<h2 class="section-h">{0}</h2>'.format(_esc(label))

    role = _twocol_role_headline(artifact)
    first, last = _twocol_name_parts(artifact)
    name_html = first + (' <span class="last">{0}</span>'.format(last) if last else "")
    contact_lis = "".join(
        '<li><span class="tag">{tag}</span>{value}</li>'.format(
            tag=_esc(tag), value=_esc(value)
        )
        for tag, value in _twocol_contact_items(artifact)
    )
    sidebar_contact = (
        '<section class="sidebar-section">'
        '<h2 class="section-h">Contact</h2>'
        '<ul class="contact-list">{0}</ul>'
        '</section>'.format(contact_lis)
        if contact_lis
        else ""
    )
    main = _twocol_main_sections_html(
        artifact, section_h, _twocol_experience_headrow, _twocol_projects_chips
    )
    sidebar_body = _twocol_sidebar_sections_html(artifact, section_h)
    body = (
        '<div class="sheet">'
        '<main class="main">'
        '<header class="masthead">'
        '<div class="index"><span>Resume</span></div>'
        '<h1 class="name">{name}</h1>'
        '{role}'
        '</header>'
        '{main}'
        '</main>'
        '<aside class="sidebar">'
        '<div class="brand"><div class="monogram">{monogram}</div>'
        '<div class="tag">Curriculum Vitae</div></div>'
        '{contact}{sidebar}'
        '</aside>'
        '</div>'
    ).format(
        monogram=_twocol_monogram(artifact),
        contact=sidebar_contact,
        sidebar=sidebar_body,
        name=name_html,
        role='<div class="role">{0}</div>'.format(role) if role else "",
        main=main,
    )
    return _twocol_document(title, _TWOCOL_STYLE_MONOCHROME_BLACK, body)


# --- TEMPLATE 08 — Plum Berry -----------------------------------------------
_TWOCOL_STYLE_PLUM_BERRY = """
:root{--font-sans:Arial,"Helvetica Neue",Helvetica,sans-serif;--ink:#2a1d2a;--muted:#7a6877;--line:#e3d4dd;--accent:#b94f7a;--accent-soft:#f2dce6;--paper:#f6ecef;--surface:#f6ecef;--sidebar-bg:#3b1f3a;--sidebar-fg:#f6e8ee;--sidebar-mute:#c9a8bc;--sidebar-rule:#4f2c4d;--sidebar-w:35%;--main-pad:17mm;--side-pad:12mm;}
@page{size:A4;margin:0;}
*{box-sizing:border-box;}
html,body{margin:0;padding:0;background:var(--paper);}
body{font-family:var(--font-sans);color:var(--ink);font-size:10pt;line-height:1.55;}
.sheet{display:flex;width:210mm;min-height:297mm;margin:0 auto;background:var(--paper);}
.sidebar{width:var(--sidebar-w);background:var(--sidebar-bg);color:var(--sidebar-fg);padding:var(--side-pad);display:flex;flex-direction:column;order:-1;}
.main{flex:1;background:var(--surface);padding:var(--main-pad);display:flex;flex-direction:column;}
.masthead{margin-bottom:10mm;}
.masthead .eyebrow{font-size:8.5pt;letter-spacing:4px;text-transform:uppercase;color:var(--accent);font-weight:700;margin-bottom:6mm;}
.masthead .name{margin:0;font-size:38pt;line-height:0.98;letter-spacing:-1px;font-weight:200;color:var(--ink);}
.masthead .name .last{display:block;font-weight:900;color:var(--accent);}
.masthead .role{margin-top:5mm;padding-top:4mm;border-top:1px solid var(--ink);font-size:11pt;font-style:italic;color:var(--ink);}
.section{margin-bottom:7mm;}
.section:last-child{margin-bottom:0;}
.section-h{display:flex;align-items:center;gap:4mm;margin:0 0 4mm 0;font-size:9pt;font-weight:800;letter-spacing:4px;text-transform:uppercase;color:var(--ink);}
.section-h::before,.section-h::after{content:"";flex:1;height:1px;background:var(--ink);}
.section-h::before{flex:0 0 4mm;}
.sidebar .section-h{color:var(--sidebar-fg);}
.sidebar .section-h::before,.sidebar .section-h::after{background:var(--sidebar-fg);}
.summary p{margin:0;font-size:10.4pt;line-height:1.65;}
.entry{padding:0 0 5mm 0;page-break-inside:avoid;}
.entry:last-child{padding-bottom:0;}
.entry-head{display:flex;justify-content:space-between;align-items:baseline;gap:10px;margin-bottom:0.5mm;}.entry-head>div{flex:1;min-width:0;}.entry-head .entry-dates{flex-shrink:0;}
.entry-title{font-size:11.2pt;font-weight:800;margin:0;color:var(--ink);}
.entry-org{font-size:10pt;color:var(--accent);margin:0;font-style:italic;font-weight:600;}
.entry-dates{font-size:9pt;color:var(--muted);white-space:nowrap;font-weight:600;letter-spacing:0.3px;}
.entry ul{margin:2mm 0 0 0;padding-left:14px;font-size:9.8pt;}
.entry ul li{margin-bottom:1.3mm;line-height:1.5;}
.entry ul li::marker{color:var(--accent);}
.project{page-break-inside:avoid;margin-bottom:4mm;}
.project:last-child{margin-bottom:0;}
.project-head{display:flex;align-items:baseline;gap:8px;margin-bottom:1mm;}
.project-name{font-size:10.5pt;font-weight:800;color:var(--ink);}
.project-link{font-size:9pt;color:var(--accent);text-decoration:underline;text-underline-offset:2px;}
.project-desc{font-size:9.8pt;margin:0;line-height:1.5;}
.tech-list{display:flex;flex-wrap:wrap;gap:5px;margin-top:1.5mm;}
.tech{font-size:8.4pt;font-weight:700;color:var(--accent);background:var(--accent-soft);padding:1.5px 7px;border-radius:999px;letter-spacing:0.3px;}
.pubs{margin:0;padding:0;list-style:none;}
.pubs li{font-size:9.5pt;line-height:1.5;margin-bottom:2.5mm;padding-bottom:2.5mm;border-bottom:1px solid var(--line);}
.pubs li:last-child{border-bottom:none;padding-bottom:0;}
.sidebar-section{margin-bottom:7mm;}
.sidebar-section:last-child{margin-bottom:0;}
.contact-list{list-style:none;padding:0;margin:0;}
.contact-list li{font-size:9.2pt;color:var(--sidebar-fg);margin-bottom:2.2mm;word-break:break-word;line-height:1.4;}
.contact-list .tag{display:block;font-size:7.5pt;letter-spacing:2px;text-transform:uppercase;color:var(--sidebar-mute);margin-bottom:1px;font-weight:700;}
.skill-group{margin-bottom:3.5mm;}
.skill-group:last-child{margin-bottom:0;}
.skill-group-h{font-size:7.8pt;font-weight:800;letter-spacing:1.8px;text-transform:uppercase;color:var(--sidebar-mute);margin:0 0 1.5mm 0;}
.skill-list{list-style:none;padding:0;margin:0;}
.skill-list li{font-size:9.3pt;color:var(--sidebar-fg);padding:1.1mm 0;line-height:1.3;}
.skill-list li::before{content:"\\2014";margin-right:8px;color:var(--sidebar-mute);}
.edu-entry{margin-bottom:4mm;padding-bottom:4mm;border-bottom:1px solid var(--sidebar-rule);}
.edu-entry:last-child{border-bottom:none;padding-bottom:0;margin-bottom:0;}
.edu-degree{font-size:9.8pt;font-weight:800;margin:0;line-height:1.3;}
.edu-inst{font-size:9pt;color:var(--sidebar-mute);margin:1px 0 0 0;font-style:italic;}
.edu-dates{font-size:8.5pt;color:var(--sidebar-fg);margin-top:1.5mm;font-weight:700;letter-spacing:0.4px;}
.cert-list{list-style:none;padding:0;margin:0;}
.cert-list li{font-size:9.1pt;color:var(--sidebar-fg);margin-bottom:2.5mm;line-height:1.4;}
"""


def _render_twocol_plum_berry(
    artifact: TailoredResumeArtifact, title: str
) -> str:
    """Template 08 - plum sidebar, berry accent, head-row experience."""

    def section_h(label):
        return '<h2 class="section-h">{0}</h2>'.format(_esc(label))

    role = _twocol_role_headline(artifact)
    first, last = _twocol_name_parts(artifact)
    name_html = first + ('<span class="last">{0}</span>'.format(last) if last else "")
    contact_lis = "".join(
        '<li><span class="tag">{tag}</span>{value}</li>'.format(
            tag=_esc(tag), value=_esc(value)
        )
        for tag, value in _twocol_contact_items(artifact)
    )
    sidebar_contact = (
        '<section class="sidebar-section">'
        '<h2 class="section-h">Contact</h2>'
        '<ul class="contact-list">{0}</ul>'
        '</section>'.format(contact_lis)
        if contact_lis
        else ""
    )
    main = _twocol_main_sections_html(
        artifact, section_h, _twocol_experience_headrow, _twocol_projects_chips
    )
    sidebar_body = _twocol_sidebar_sections_html(artifact, section_h)
    body = (
        '<div class="sheet">'
        '<main class="main">'
        '<header class="masthead">'
        '<div class="eyebrow">Curriculum Vitae</div>'
        '<h1 class="name">{name}</h1>'
        '{role}'
        '</header>'
        '{main}'
        '</main>'
        '<aside class="sidebar">{contact}{sidebar}</aside>'
        '</div>'
    ).format(
        contact=sidebar_contact,
        sidebar=sidebar_body,
        name=name_html,
        role='<div class="role">{0}</div>'.format(role) if role else "",
        main=main,
    )
    return _twocol_document(title, _TWOCOL_STYLE_PLUM_BERRY, body)


# --- TEMPLATE 10 — Burgundy & Champagne -------------------------------------
_TWOCOL_STYLE_BURGUNDY_CHAMPAGNE = """
:root{--font-sans:Arial,"Helvetica Neue",Helvetica,sans-serif;--ink:#1c1311;--muted:#6f5e54;--line:#e3d8c5;--accent:#a17a3a;--accent-soft:#f1e5cd;--paper:#faf3e3;--surface:#faf3e3;--sidebar-bg:#5a121f;--sidebar-fg:#f3e6cf;--sidebar-mute:#d3b78a;--sidebar-rule:#71202c;--sidebar-w:33%;--main-pad:17mm;--side-pad:12mm;}
@page{size:A4;margin:0;}
*{box-sizing:border-box;}
html,body{margin:0;padding:0;background:var(--paper);}
body{font-family:var(--font-sans);color:var(--ink);font-size:10pt;line-height:1.5;}
.sheet{display:flex;width:210mm;min-height:297mm;margin:0 auto;background:var(--paper);}
.sidebar{width:var(--sidebar-w);background:var(--sidebar-bg);color:var(--sidebar-fg);padding:var(--side-pad);display:flex;flex-direction:column;order:-1;}
.main{flex:1;background:var(--surface);padding:var(--main-pad);display:flex;flex-direction:column;}
.masthead{margin-bottom:9mm;padding-bottom:5mm;border-bottom:3px double var(--ink);}
.masthead .meta{font-size:8.5pt;letter-spacing:4px;text-transform:uppercase;color:var(--accent);font-weight:800;margin-bottom:6mm;}
.masthead .name{margin:0;font-size:30pt;line-height:1;letter-spacing:-0.5px;color:var(--ink);font-weight:800;}
.masthead .name .first{font-weight:200;letter-spacing:0;}
.masthead .role{margin-top:4mm;font-size:10.5pt;font-style:italic;color:var(--accent);letter-spacing:0.3px;}
.section{margin-bottom:7mm;}
.section:last-child{margin-bottom:0;}
.section-h{font-size:9.3pt;font-weight:800;letter-spacing:3.5px;text-transform:uppercase;color:var(--ink);margin:0 0 4mm 0;padding-bottom:2mm;border-bottom:1px solid var(--accent);}
.sidebar .section-h{color:var(--accent);border-bottom-color:var(--accent);}
.summary p{margin:0;font-size:10.3pt;line-height:1.6;}
.entry{display:flex;gap:5mm;page-break-inside:avoid;margin-bottom:5mm;}
.entry:last-child{margin-bottom:0;}
.entry-gutter{width:22mm;flex-shrink:0;padding-top:1mm;}
.entry-dates{font-size:9pt;color:var(--accent);font-weight:800;letter-spacing:0.5px;line-height:1.3;text-transform:uppercase;}
.entry-body{flex:1;}
.entry-title{font-size:11pt;font-weight:800;margin:0;line-height:1.25;color:var(--ink);}
.entry-org{font-size:10pt;color:var(--muted);margin:1px 0 0 0;font-style:italic;}
.entry ul{margin:2mm 0 0 0;padding-left:16px;font-size:9.8pt;}
.entry ul li{margin-bottom:1.3mm;line-height:1.5;}
.entry ul li::marker{color:var(--accent);}
.project{display:flex;gap:5mm;page-break-inside:avoid;margin-bottom:4mm;}
.project:last-child{margin-bottom:0;}
.project-gutter{width:22mm;flex-shrink:0;padding-top:1mm;}
.project-gutter .gutter-label{font-size:8.4pt;color:var(--accent);font-weight:800;letter-spacing:1.6px;text-transform:uppercase;}
.project-body{flex:1;}
.project-name{font-size:10.5pt;font-weight:800;color:var(--ink);margin:0 0 1mm 0;}
.project-link{display:inline-block;font-size:9pt;color:var(--accent);text-decoration:none;margin-left:6px;font-weight:400;font-style:italic;}
.project-desc{font-size:9.8pt;margin:0 0 1.5mm 0;line-height:1.5;}
.tech-line{font-size:8.8pt;color:var(--muted);}
.tech-line .tech-label{font-weight:800;text-transform:uppercase;letter-spacing:1.5px;margin-right:6px;color:var(--accent);}
.pubs{margin:0;padding:0;list-style:none;counter-reset:pubcount;}
.pubs li{font-size:9.5pt;line-height:1.5;margin-bottom:2.5mm;padding-left:9mm;position:relative;counter-increment:pubcount;}
.pubs li::before{content:"\\00a7 " counter(pubcount);position:absolute;left:0;top:0;font-weight:800;color:var(--accent);font-size:9pt;letter-spacing:0.4px;}
.sidebar-section{margin-bottom:7mm;}
.sidebar-section:last-child{margin-bottom:0;}
.contact-list{list-style:none;padding:0;margin:0;}
.contact-list li{font-size:9.2pt;color:var(--sidebar-fg);margin-bottom:2.2mm;word-break:break-word;line-height:1.4;}
.contact-list .tag{display:block;font-size:7.5pt;letter-spacing:2px;text-transform:uppercase;color:var(--accent);margin-bottom:1px;font-weight:800;}
.skill-group{margin-bottom:3.5mm;}
.skill-group:last-child{margin-bottom:0;}
.skill-group-h{font-size:7.8pt;font-weight:800;letter-spacing:2px;text-transform:uppercase;color:var(--accent);margin:0 0 1.5mm 0;}
.skill-list{list-style:none;padding:0;margin:0;}
.skill-list li{font-size:9.3pt;color:var(--sidebar-fg);padding:1.2mm 0;line-height:1.3;border-bottom:1px solid var(--sidebar-rule);}
.skill-list li:last-child{border-bottom:none;}
.edu-entry{margin-bottom:4mm;}
.edu-entry:last-child{margin-bottom:0;}
.edu-degree{font-size:9.8pt;font-weight:800;margin:0;line-height:1.3;}
.edu-inst{font-size:9pt;color:var(--sidebar-mute);margin:1px 0 0 0;font-style:italic;}
.edu-dates{font-size:8.5pt;color:var(--accent);margin-top:1.5mm;font-weight:800;letter-spacing:0.4px;}
.cert-list{list-style:none;padding:0;margin:0;}
.cert-list li{font-size:9.1pt;color:var(--sidebar-fg);margin-bottom:2.5mm;line-height:1.4;padding-left:9mm;position:relative;}
.cert-list li::before{content:"\\00b7";position:absolute;left:4mm;top:-2px;font-size:16pt;color:var(--accent);}
"""


def _render_twocol_burgundy_champagne(
    artifact: TailoredResumeArtifact, title: str
) -> str:
    """Template 10 - wine sidebar, champagne-gold accent, date-gutter."""

    def section_h(label):
        return '<h2 class="section-h">{0}</h2>'.format(_esc(label))

    role = _twocol_role_headline(artifact)
    first, last = _twocol_name_parts(artifact)
    name_html = (
        '<span class="first">{0}</span> {1}'.format(first, last)
        if last
        else first
    )
    contact_lis = "".join(
        '<li><span class="tag">{tag}</span>{value}</li>'.format(
            tag=_esc(tag), value=_esc(value)
        )
        for tag, value in _twocol_contact_items(artifact)
    )
    sidebar_contact = (
        '<section class="sidebar-section">'
        '<h2 class="section-h">Contact</h2>'
        '<ul class="contact-list">{0}</ul>'
        '</section>'.format(contact_lis)
        if contact_lis
        else ""
    )
    main = _twocol_main_sections_html(
        artifact, section_h, _twocol_experience_gutter, _twocol_projects_gutter
    )
    sidebar_body = _twocol_sidebar_sections_html(artifact, section_h)
    body = (
        '<div class="sheet">'
        '<main class="main">'
        '<header class="masthead">'
        '<div class="meta">Curriculum Vitae</div>'
        '<h1 class="name">{name}</h1>'
        '{role}'
        '</header>'
        '{main}'
        '</main>'
        '<aside class="sidebar">{contact}{sidebar}</aside>'
        '</div>'
    ).format(
        contact=sidebar_contact,
        sidebar=sidebar_body,
        name=name_html,
        role='<div class="role">{0}</div>'.format(role) if role else "",
        main=main,
    )
    return _twocol_document(title, _TWOCOL_STYLE_BURGUNDY_CHAMPAGNE, body)


# Dispatch table — `ThemeSpec.twocol_layout` → renderer. Populated by each
# template block below as it is defined; `_build_resume_html_twocol` reads it
# at call time so definition order does not matter.
_TWOCOL_RENDERERS = {
    "timeline_tech": _render_twocol_timeline_tech,
    "editorial_minimal": _render_twocol_editorial_minimal,
    "classic_slate": _render_twocol_classic_slate,
    "monochrome_black": _render_twocol_monochrome_black,
    "plum_berry": _render_twocol_plum_berry,
    "burgundy_champagne": _render_twocol_burgundy_champagne,
}


def _build_resume_html_twocol(
    artifact: TailoredResumeArtifact, title: str, spec: "ThemeSpec"
) -> str:
    """Dispatch to the bespoke renderer named by `spec.twocol_layout`
    (ADR-032). Falls back to the timeline_tech renderer if the layout
    name is somehow unknown (defensive - the registry always sets it)."""
    renderer = _TWOCOL_RENDERERS.get(spec.twocol_layout) or _render_twocol_timeline_tech
    return renderer(artifact, title)


# Derived from `_THEME_SPECS` (ADR-015 follow-up). Do NOT hand-edit —
# add a ThemeSpec to the registry instead.
_RESUME_THEME_PALETTES = {
    key: spec.resume_palette() for key, spec in _THEME_SPECS.items()
}


def _resolve_resume_palette(theme: str | None) -> dict:
    return _RESUME_THEME_PALETTES.get(
        theme or "classic_ats", _RESUME_THEME_PALETTES["classic_ats"]
    )


def _build_resume_html(text, title="Tailored Resume", theme="classic_ats", artifact: TailoredResumeArtifact | None = None):
    # Layout branch FIRST so the single-column path below is provably
    # untouched (byte-identical for every single_column theme — the
    # branch only fires for a two_column ThemeSpec WITH a structured
    # artifact; markdown-only callers always take the classic path).
    # ADR-032: the two-column branch dispatches to the bespoke designer
    # renderer named by `_spec.twocol_layout`.
    _spec = resolve_theme(theme)
    if artifact is not None and _spec.layout == "two_column":
        return _build_resume_html_twocol(artifact, title, _spec)

    body_html = _MARKDOWN.render(text or "")
    if artifact is not None:
        body_html = _build_structured_resume_body_classic(
            artifact, header_banded=bool(_spec.header_band_bg)
        )
    safe_title = html.escape(title or "Tailored Resume")
    palette = _resolve_resume_palette(theme)

    return """
<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="utf-8" />
    <title>{title}</title>
    <style>
        @page {{ size: A4; margin: 0; }}
        /* Theme-keyed palette: classic_ats ships warm-brown,
           professional_neutral collapses to pure black/white/gray for
           conservative recruiters / B&W printing. Typeface is the
           shared Arial sans family for every theme (2026-05-21). */
        :root {{
            --ink: {ink};
            --muted: {muted};
            --accent: {accent};
            --accent-soft: {accent_soft};
            --line: {line};
            --paper: {paper};
            --surface: {surface};
        }}
        * {{ box-sizing: border-box; }}
        html, body {{ margin: 0; padding: 0; }}
        body {{ font-family: {body_font_family}; color: var(--ink); background: var(--paper); font-size: 10.5pt; line-height: 1.5; }}
        /* Prose-y parts (summary, bullets) use the theme's prose font.
           Since the 2026-05-21 typography unification that is the same
           shared Arial sans family as the body for every theme; the
           rule is kept so a future theme could re-introduce a distinct
           prose face without a renderer change. */
        .resume-summary, .resume-bullet-list li {{ font-family: {prose_font_family}; line-height: {prose_line_height}; }}
        /* min-height keeps short resumes filling a single A4 page;
           overflow-x: hidden prevents the negative-margin h2 trick from
           leaking past the page edge horizontally while letting
           WeasyPrint paginate vertically when the resume runs long. */
        .resume-shell {{ position: relative; min-height: 297mm; background: var(--surface); padding: 13mm 15mm 13mm; overflow-x: hidden; }}
        .resume-experience-card, .resume-education-card, .resume-project-card {{ page-break-inside: avoid; break-inside: avoid; }}
        h2, h3 {{ page-break-after: avoid; break-after: avoid; }}
        .resume-shell::before {{ content: none; }}
        .resume-shell::after {{ content: none; }}
        /* Name is the resume's title — bold, mixed-case, light tracking.
           Font family is theme-keyed so neutral stays all-Arial. */
        h1 {{ font-family: {h1_font_family}; font-size: 22pt; font-weight: 700; margin: 0 0 4px; letter-spacing: -0.005em; color: var(--ink); text-transform: none; }}
        h2 {{ font-size: 10pt; margin: 0 -15mm 6px; text-transform: uppercase; letter-spacing: 0.18em; color: var(--accent); padding: 0 15mm 3px; border-bottom: 2px solid var(--line); }}
        h3 {{ font-size: 10.5pt; margin: 10px 0 4px; color: var(--accent); }}
        p {{ margin: 0 0 6px; }}
        ul {{ margin: 0 0 8px 1rem; padding: 0; }}
        li {{ margin: 0 0 3px; }}
        strong {{ color: var(--ink); }}
        em {{ color: var(--muted); }}
        code {{ background: {code_bg}; border: 1px solid var(--line); border-radius: 4px; padding: 0.08rem 0.28rem; }}
        hr {{ border: 0; border-top: 1px solid var(--line); margin: 14px 0; }}
        blockquote {{ margin: 0 0 10px; padding: 8px 12px; border-left: 4px solid var(--accent); background: var(--accent-soft); color: var(--muted); }}
        .resume-classic-header {{ position: relative; z-index: 1; padding: 0 15mm 10px; margin: 0 -15mm; border-bottom: {header_border_width} solid {header_rule_color}; }}
        /* Opt-in header band. Only applied when the body builder adds
           the --band class (themes that set ThemeSpec.header_band_bg);
           inert for every other theme. Full-bleed via the existing
           negative side-margins. */
        /* `margin-top: -13mm` cancels the shell's 13mm top padding so
           the band bleeds to the very top edge of the page (no white
           strip above it); the matching 13mm top padding keeps the
           name where it visually sat before. -15mm sides keep the
           existing full-width bleed. */
        .resume-classic-header--band {{ background: {header_band_bg}; margin: -13mm -15mm 0; padding: 13mm 15mm 12px; }}
        .resume-classic-header--band h1 {{ color: {header_band_fg}; }}
        .resume-classic-header--band .resume-classic-role,
        .resume-classic-header--band .resume-contact-inline,
        .resume-classic-header--band .resume-contact-inline .rc-item {{ color: {header_band_fg}; }}
        .resume-classic-role {{ font-size: 10.2pt; color: var(--muted); text-transform: uppercase; letter-spacing: 0.12em; margin-bottom: 4px; }}
        .resume-contact-inline {{ color: var(--muted); font-size: 9.6pt; line-height: 1.55; max-width: 88%; }}
        .rc-item {{ white-space: nowrap; }}
        .resume-contact-links {{ margin-top: 2px; }}
        .resume-skill-inline {{ color: var(--ink); font-size: 9.8pt; line-height: 1.7; }}
        .resume-skill-category {{ color: var(--ink); font-size: 9.8pt; line-height: 1.7; margin: 0 0 2px; }}
        .resume-skill-category strong {{ color: var(--accent); font-weight: 700; }}
        .resume-classic-section {{ position: relative; z-index: 1; margin-top: 12px; }}
        .resume-classic-section--plain-head h2 {{ border-bottom: 0; padding-bottom: 0; }}
        .resume-experience-card, .resume-project-card {{ background: transparent; border: 0; border-radius: 0; padding: 0; }}
        .resume-education-card {{ background: transparent; border: 0; border-radius: 0; padding: 0; }}
        .resume-role-row {{ display: flex; justify-content: space-between; gap: 12px; align-items: flex-start; }}
        .resume-role-row h3, .resume-education-card h3 {{ margin: 0 0 4px; }}
        .resume-role-meta, .resume-role-dates, .resume-education-meta, .resume-education-dates {{ margin: 0; color: var(--muted); font-size: 9.5pt; }}
        .resume-bullet-list, .resume-contact-list, .resume-plain-list {{ margin: 0; padding-left: 1rem; }}
        .resume-empty {{ color: var(--muted); font-style: italic; }}
        .resume-experience-card + .resume-experience-card,
        .resume-project-card + .resume-project-card {{ position: relative; margin-top: 10px; padding-top: 10px; }}
        .resume-experience-card + .resume-experience-card::before,
        .resume-project-card + .resume-project-card::before {{ content: ""; position: absolute; top: 0; left: 12px; right: 12px; border-top: 1px solid var(--line); }}
        .resume-education-card + .resume-education-card {{ margin-top: 10px; }}
        @media all and (max-width: 720px) {{ .resume-classic-header {{ padding: 0 15mm 10px; margin: 0 -15mm; }} .resume-classic-header--band {{ margin: -13mm -15mm 0; padding: 13mm 15mm 12px; }} .resume-contact-inline {{ max-width: 100%; }} .resume-role-row {{ display: block; }} .resume-role-dates {{ margin-top: 6px; }} }}
    </style>
</head>
<body>
    <main class="resume-shell resume-shell--classic">{body}</main>
</body>
</html>
""".format(title=safe_title, body=body_html, **palette)


def _inline_to_markup(inline_node):
    parts = []
    for child in inline_node.children or []:
        if child.type == "text":
            parts.append(_paragraph_text(child.content))
        elif child.type == "softbreak":
            parts.append(" ")
        elif child.type == "hardbreak":
            parts.append("<br/><br/>")
        elif child.type == "strong":
            parts.append("<b>{text}</b>".format(text=_inline_to_markup(child)))
        elif child.type == "em":
            parts.append("<i>{text}</i>".format(text=_inline_to_markup(child)))
        elif child.type == "code_inline":
            parts.append(
                '<font face="Courier">{text}</font>'.format(
                    text=_paragraph_text(child.content)
                )
            )
        elif child.type == "link":
            href = _paragraph_text((child.attrs or {}).get("href", ""))
            label = _inline_to_markup(child)
            if href:
                parts.append('<link href="{href}">{label}</link>'.format(href=href, label=label))
            else:
                parts.append(label)
        else:
            parts.append(_paragraph_text(child.content))
    return "".join(parts)


def _flatten_list_items(list_node, level=0):
    items = []
    for item_index, item_node in enumerate(list_node.children or [], start=1):
        bullet = "{index}.".format(index=item_index) if list_node.type == "ordered_list" else (
            "-" if level == 0 else "*"
        )
        first_paragraph_rendered = False

        for child in item_node.children or []:
            if child.type == "paragraph":
                items.append(
                    {
                        "kind": "list_paragraph",
                        "level": level,
                        "bullet": bullet if not first_paragraph_rendered else "",
                        "text": _inline_to_markup(child.children[0]) if child.children else "",
                        "continued": first_paragraph_rendered,
                    }
                )
                first_paragraph_rendered = True
            elif child.type in {"bullet_list", "ordered_list"}:
                items.extend(_flatten_list_items(child, level=level + 1))
            elif child.type in {"code_block", "fence"}:
                items.append(
                    {
                        "kind": "code_block",
                        "level": level + 1,
                        "text": child.content.rstrip(),
                    }
                )

    return items


def _parse_markdown_blocks(text):
    root = SyntaxTreeNode(_MARKDOWN.parse(text or ""))
    blocks = []

    for node in root.children or []:
        if node.type == "heading":
            level = int(node.tag[1]) if node.tag.startswith("h") else 2
            title = _inline_to_markup(node.children[0]) if node.children else ""
            block_type = "title" if level == 1 else "heading" if level == 2 else "subheading"
            blocks.append((block_type, title))
            continue

        if node.type == "paragraph":
            content = _inline_to_markup(node.children[0]) if node.children else ""
            blocks.append(("paragraph", content))
            continue

        if node.type in {"bullet_list", "ordered_list"}:
            blocks.append(("list", _flatten_list_items(node)))
            continue

        if node.type in {"code_block", "fence"}:
            blocks.append(("code_block", node.content.rstrip()))
            continue

        if node.type == "hr":
            blocks.append(("rule", ""))

    return blocks


def _generate_pdf_with_reportlab(text, title):
    with warnings.catch_warnings():
        warnings.filterwarnings(
            "ignore",
            message=r"the load_module\(\) method is deprecated and slated for removal in Python 3.12; use exec_module\(\) instead",
            category=DeprecationWarning,
        )
        from reportlab.lib import colors
        from reportlab.lib.enums import TA_LEFT
        from reportlab.lib.pagesizes import letter
        from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
        from reportlab.platypus import (
            HRFlowable,
            Paragraph,
            Preformatted,
            SimpleDocTemplate,
            Spacer,
        )
        from reportlab.pdfbase import pdfdoc

    def md5_compat(*args, **kwargs):
        kwargs.pop("usedforsecurity", None)
        return hashlib.md5(*args, **kwargs)

    # Older ReportLab builds still call into pdfdoc.md5 with a
    # usedforsecurity kwarg shape that hashlib.md5 does not accept consistently
    # across runtimes, so we normalize it here for the fallback backend.
    pdfdoc.md5 = md5_compat

    styles = getSampleStyleSheet()
    styles.add(
        ParagraphStyle(
            name="ReportTitle",
            parent=styles["Title"],
            fontName="Helvetica-Bold",
            fontSize=23,
            leading=29,
            textColor=colors.HexColor("#0f172a"),
            alignment=TA_LEFT,
            spaceAfter=16,
        )
    )
    styles.add(
        ParagraphStyle(
            name="SectionHeading",
            parent=styles["Heading1"],
            fontName="Helvetica-Bold",
            fontSize=15,
            leading=20,
            textColor=colors.HexColor("#142033"),
            spaceBefore=12,
            spaceAfter=7,
        )
    )
    styles.add(
        ParagraphStyle(
            name="SubHeading",
            parent=styles["Heading2"],
            fontName="Helvetica-Bold",
            fontSize=12,
            leading=16,
            textColor=colors.HexColor("#1d4ed8"),
            spaceBefore=8,
            spaceAfter=4,
        )
    )
    styles.add(
        ParagraphStyle(
            name="BodyCopy",
            parent=styles["BodyText"],
            fontName="Helvetica",
            fontSize=10.3,
            leading=15,
            textColor=colors.HexColor("#1f2937"),
            spaceAfter=8,
        )
    )
    styles.add(
        ParagraphStyle(
            name="ListCopy",
            parent=styles["BodyText"],
            fontName="Helvetica",
            fontSize=10.1,
            leading=14.6,
            textColor=colors.HexColor("#1f2937"),
            spaceAfter=4,
        )
    )
    styles.add(
        ParagraphStyle(
            name="CodeBlock",
            parent=styles["Code"],
            fontName="Courier",
            fontSize=8.8,
            leading=10.8,
            textColor=colors.HexColor("#1f2937"),
            backColor=colors.HexColor("#f3f6fb"),
            borderColor=colors.HexColor("#dce5f0"),
            borderWidth=0.6,
            borderPadding=8,
            leftIndent=10,
            rightIndent=10,
            spaceBefore=3,
            spaceAfter=8,
        )
    )

    def list_style(level, continued=False):
        text_indent = 22 + (level * 16)
        return ParagraphStyle(
            name="ListLevel{level}{continued}".format(
                level=level,
                continued="Continued" if continued else "Lead",
            ),
            parent=styles["ListCopy"],
            leftIndent=text_indent,
            bulletIndent=text_indent - 12,
            firstLineIndent=0,
            spaceBefore=0 if continued else 1,
        )

    def render_page(canvas, doc):
        canvas.saveState()
        width, height = letter

        canvas.setStrokeColor(colors.HexColor("#1d4ed8"))
        canvas.setLineWidth(2)
        canvas.line(doc.leftMargin, height - 28, width - doc.rightMargin, height - 28)

        canvas.setStrokeColor(colors.HexColor("#bfdbfe"))
        canvas.setLineWidth(1)
        canvas.line(doc.leftMargin, 30, width - doc.rightMargin, 30)

        canvas.setFont("Helvetica", 8.5)
        canvas.setFillColor(colors.HexColor("#64748b"))
        canvas.drawString(doc.leftMargin, 18, "Application Copilot")
        canvas.drawRightString(
            width - doc.rightMargin,
            18,
            "Page {page}".format(page=canvas.getPageNumber()),
        )
        canvas.restoreState()

    buffer = BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=letter,
        rightMargin=44,
        leftMargin=44,
        topMargin=48,
        bottomMargin=42,
        title=title or "Tailored Resume",
    )

    flowables = []
    blocks = _parse_markdown_blocks(text)

    for index, (block_type, value) in enumerate(blocks):
        if block_type == "title":
            flowables.append(Paragraph(value, styles["ReportTitle"]))
            flowables.append(
                HRFlowable(
                    width="100%",
                    thickness=1.2,
                    color=colors.HexColor("#cbd5e1"),
                    spaceBefore=0,
                    spaceAfter=10,
                )
            )
            continue

        if block_type == "heading":
            flowables.append(Paragraph(value, styles["SectionHeading"]))
            continue

        if block_type == "subheading":
            flowables.append(Paragraph(value, styles["SubHeading"]))
            continue

        if block_type == "paragraph":
            flowables.append(Paragraph(value, styles["BodyCopy"]))
            continue

        if block_type == "list":
            for item in value:
                if item["kind"] == "list_paragraph":
                    style = list_style(item["level"], continued=item["continued"])
                    if item["bullet"]:
                        flowables.append(
                            Paragraph(item["text"], style, bulletText=item["bullet"])
                        )
                    else:
                        flowables.append(Paragraph(item["text"], style))
                elif item["kind"] == "code_block":
                    code_style = ParagraphStyle(
                        name="NestedCodeLevel{level}".format(level=item["level"]),
                        parent=styles["CodeBlock"],
                        leftIndent=22 + (item["level"] * 16),
                        rightIndent=10,
                    )
                    flowables.append(Preformatted(item["text"], code_style))
            flowables.append(Spacer(1, 6))
            continue

        if block_type == "code_block":
            flowables.append(Preformatted(value, styles["CodeBlock"]))
            continue

        if block_type == "rule" and index != len(blocks) - 1:
            flowables.append(
                HRFlowable(
                    width="100%",
                    thickness=0.8,
                    color=colors.HexColor("#d6dde6"),
                    spaceBefore=2,
                    spaceAfter=8,
                )
            )

    doc.build(flowables, onFirstPage=render_page, onLaterPages=render_page)
    buffer.seek(0)
    return buffer


def _build_pdf_html(text, title, theme=None, document_kind="cover_letter"):
    return (
        _build_resume_html(text, title=title, theme=theme or "classic_ats")
        if document_kind == "tailored_resume"
        else _build_cover_letter_html(text, title=title, theme=theme or "classic_ats")
    )


def _generate_pdf_with_weasyprint(text, title, theme=None, document_kind="cover_letter", artifact: TailoredResumeArtifact | None = None):
    _configure_weasyprint_windows_runtime()
    from weasyprint import HTML

    html_document = (
        _build_resume_html(text, title=title, theme=theme or "classic_ats", artifact=artifact)
        if document_kind == "tailored_resume"
        else _build_cover_letter_html(text, title=title, theme=theme or "classic_ats")
    )

    return BytesIO(HTML(string=html_document).write_pdf())


def _generate_pdf_with_reportlab_fallback(text, title, renderer_error=None):
    if renderer_error is not None:
        log_event(
            LOGGER,
            logging.WARNING,
            "pdf_export_weasyprint_failed",
            "WeasyPrint PDF export failed; attempting ReportLab fallback.",
            title=title,
            error_type=type(renderer_error).__name__,
            details=str(renderer_error),
        )

    try:
        return _generate_pdf_with_reportlab(text, title)
    except Exception as error:
        log_event(
            LOGGER,
            logging.ERROR,
            "pdf_export_reportlab_failed",
            "ReportLab PDF export failed after Playwright fallback.",
            title=title,
            error_type=type(error).__name__,
        )
        raise ExportError(
            "PDF export failed. Try downloading the Markdown package instead.",
            details=str(error),
        ) from error


def generate_pdf(text, title="Tailored Resume", theme=None, document_kind="cover_letter", artifact: TailoredResumeArtifact | None = None):
    try:
        return _generate_pdf_with_weasyprint(text, title, theme=theme, document_kind=document_kind, artifact=artifact)
    except Exception as renderer_error:
        return _generate_pdf_with_reportlab_fallback(text, title, renderer_error=renderer_error)


# ---------------------------------------------------------------------------
# DOCX export
#
# Mirrors the structural decomposition used by
# `_build_structured_resume_body_classic` (resume) and
# `_build_cover_letter_html` (cover letter) so the DOCX reads as the
# same document, not a different format. We pull from the structured
# artifact fields (header, experience_entries, etc.) for the resume —
# NOT from the markdown — because Phase 2 of the DOCX export plan
# removes markdown export entirely.
#
# Phase 1 implements the `classic_ats` theme only; `professional_neutral`
# lands in Phase 4.
# ---------------------------------------------------------------------------


# Per-theme DOCX palettes — OOXML-safe hex (no '#' prefix), derived
# from `_THEME_SPECS` (ADR-015 follow-up) so they can't drift from the
# PDF resume/cover-letter palettes. `ThemeSpec.docx_palette()` strips
# '#' + upper-cases the canonical hex and carries the single OOXML
# font family names. Do NOT hand-edit — add a ThemeSpec instead.
_DOCX_THEME_PALETTES: dict[str, dict[str, str]] = {
    key: spec.docx_palette() for key, spec in _THEME_SPECS.items()
}


def _resolve_docx_palette(theme: str | None) -> dict[str, str]:
    """Pick the DOCX palette by name. Unknown / blank values fall back
    to classic_ats so the renderer never crashes on an unexpected
    artifact.theme — same fallback policy as `_resolve_resume_theme`
    in artifact_export_service."""
    return _DOCX_THEME_PALETTES.get(
        str(theme or "").strip(),
        _DOCX_THEME_PALETTES["classic_ats"],
    )

# Default page margins (in inches). Matches the ~18mm @page margin the
# WeasyPrint renderer uses for the classic_ats resume shell.
_DOCX_PAGE_MARGIN_INCHES = 0.7
# US Letter page width. The résumé DOCX is always Letter; the content
# column width = this minus both margins. Used to place the role-row
# RIGHT tab stop at the true right edge of the text column.
_DOCX_LETTER_WIDTH_INCHES = 8.5


def _docx_add_bottom_border(paragraph, *, color_hex: str, size_eighths_pt: int = 6):
    """Add a bottom border to a paragraph by injecting raw OOXML.

    python-docx doesn't expose paragraph borders directly. The widget
    we want is `<w:pBdr><w:bottom w:val="single" w:sz="6"
    w:color="8F6845"/></w:pBdr>` inside the paragraph properties.
    `size_eighths_pt` is in eighths of a point, so 6 = 0.75pt, 8 = 1pt.
    """
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    p_pr = paragraph._p.get_or_add_pPr()
    # Remove any existing bottom border so re-runs don't stack.
    for existing in p_pr.findall(qn("w:pBdr")):
        p_pr.remove(existing)
    bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size_eighths_pt))
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color_hex)
    bdr.append(bottom)
    p_pr.append(bdr)


def _docx_apply_run_font(run, *, family: str, size_pt: float, color_hex: str | None = None, bold: bool = False, italic: bool = False, small_caps: bool = False):
    """Set the common font attrs on a run.

    Called per-run so each piece of text picks up the right family /
    size / color, since python-docx doesn't inherit cleanly from
    style overrides when we add multiple runs to one paragraph.
    """
    from docx.shared import Pt, RGBColor

    run.font.name = family
    run.font.size = Pt(size_pt)
    run.bold = bold
    run.italic = italic
    if color_hex:
        run.font.color.rgb = RGBColor.from_string(color_hex)
    if small_caps:
        run.font.small_caps = True


def _docx_set_page_margins(document, *, inches: float):
    """Apply equal margins (top/bottom/left/right) on the active section."""
    from docx.shared import Inches

    for section in document.sections:
        section.top_margin = Inches(inches)
        section.bottom_margin = Inches(inches)
        section.left_margin = Inches(inches)
        section.right_margin = Inches(inches)


def _docx_resume_section_heading(document, label: str, *, palette: dict):
    """Add a section H2 with accent color + thin bottom border.

    Matches the visual weight of the `.resume-classic-section h2` HTML
    rule (small-caps letterspaced look, accent-colored underline)."""
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = _docx_pt(8)
    paragraph.paragraph_format.space_after = _docx_pt(2)
    run = paragraph.add_run(label.upper())
    _docx_apply_run_font(
        run,
        family=palette["heading_font"],
        size_pt=11.5,
        color_hex=palette["accent"],
        bold=True,
        small_caps=False,
    )
    _docx_add_bottom_border(
        paragraph,
        color_hex=palette["line"],
        size_eighths_pt=4,
    )
    return paragraph


def _docx_pt(value: float):
    """Pt(...) wrapper that imports lazily so module-level test discovery
    doesn't need to bring in docx.shared just to enumerate names."""
    from docx.shared import Pt

    return Pt(value)


def _docx_inches(value: float):
    from docx.shared import Inches

    return Inches(value)


def _docx_add_resume_header(document, artifact: TailoredResumeArtifact, *, palette: dict):
    name = (artifact.header.full_name or artifact.title or "Candidate").strip()

    name_paragraph = document.add_paragraph()
    name_paragraph.paragraph_format.space_after = _docx_pt(2)
    # Header alignment matches the HTML/PDF builder: left-aligned, same
    # margin as every section heading. The earlier centred header was
    # inconsistent with the body and with the PDF rendering of the same
    # artifact — both readers should see the same layout.
    name_paragraph.alignment = _docx_alignment("left")
    run = name_paragraph.add_run(name)
    _docx_apply_run_font(
        run,
        family=palette["heading_font"],
        size_pt=20,
        color_hex=palette["ink"],
        bold=True,
    )

    # Mode-aware headline / role line. The HTML+PDF builders render
    # `artifact.target_role` as an uppercase muted line between the
    # name and the contact block (`.resume-classic-role`, added
    # 2026-05-19). The DOCX header builder was missed in that change —
    # so a JD-tailored résumé showed its role on the PDF but not the
    # DOCX. Render it here too (omitted entirely when target_role is
    # "" — a name-only header stays standard; we never fabricate one).
    headline = str(getattr(artifact, "target_role", "") or "").strip()
    if headline:
        role_paragraph = document.add_paragraph()
        role_paragraph.alignment = _docx_alignment("left")
        role_paragraph.paragraph_format.space_after = _docx_pt(2)
        role_run = role_paragraph.add_run(headline.upper())
        _docx_apply_run_font(
            role_run,
            family=palette["body_font"],
            size_pt=10,
            color_hex=palette["muted"],
        )

    contact_values = []
    if artifact.header.location:
        contact_values.append(artifact.header.location.strip())
    contact_values.extend(
        item.strip()
        for item in (artifact.header.contact_lines or [])
        if str(item or "").strip()
    )
    if contact_values:
        contact_paragraph = document.add_paragraph()
        contact_paragraph.alignment = _docx_alignment("left")
        contact_paragraph.paragraph_format.space_after = _docx_pt(6)
        contact_run = contact_paragraph.add_run(" | ".join(contact_values))
        _docx_apply_run_font(
            contact_run,
            family=palette["body_font"],
            size_pt=10,
            color_hex=palette["muted"],
        )
    else:
        # Still want the underline below the name even when contact is
        # empty, so add it on the name paragraph itself.
        contact_paragraph = name_paragraph

    # Accent underline below the contact line (or name when contact is
    # missing) gives the resume header its editorial identity.
    _docx_add_bottom_border(
        contact_paragraph,
        color_hex=palette["accent"],
        size_eighths_pt=8,
    )


def _docx_alignment(name: str):
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    return {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
        "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
    }.get(name, WD_ALIGN_PARAGRAPH.LEFT)


def _docx_add_role_row(document, *, title: str, dates: str, palette: dict):
    """Title on the left, dates right-aligned via tab stop. Mirrors the
    `.resume-role-row` flex layout."""
    from docx.enum.text import WD_TAB_ALIGNMENT

    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = _docx_pt(4)
    paragraph.paragraph_format.space_after = _docx_pt(0)
    # RIGHT tab stop at the right edge of the text column so the date
    # sits flush-right (mirrors the PDF's flex `.resume-role-row`).
    # Tab-stop positions are measured from the LEFT margin, so the
    # right edge = content width = page width minus both margins.
    # BUGFIX: this previously read `7.1 - 2 * margin` — but 7.1 was
    # ALREADY the content width (8.5 - 2*0.7), so the margins were
    # subtracted twice and every date landed 1.4in short of the
    # margin instead of flush-right.
    paragraph.paragraph_format.tab_stops.add_tab_stop(
        _docx_inches(_DOCX_LETTER_WIDTH_INCHES - 2 * _DOCX_PAGE_MARGIN_INCHES),
        WD_TAB_ALIGNMENT.RIGHT,
    )

    title_run = paragraph.add_run(title)
    _docx_apply_run_font(
        title_run,
        family=palette["body_font"],
        size_pt=11.5,
        color_hex=palette["ink"],
        bold=True,
    )
    if dates:
        tab_run = paragraph.add_run("\t")
        _docx_apply_run_font(
            tab_run,
            family=palette["body_font"],
            size_pt=11,
            color_hex=palette["muted"],
        )
        dates_run = paragraph.add_run(dates)
        _docx_apply_run_font(
            dates_run,
            family=palette["body_font"],
            size_pt=10.5,
            color_hex=palette["muted"],
        )


def _docx_add_meta_line(document, text: str, *, palette: dict, italic: bool = True):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = _docx_pt(2)
    run = paragraph.add_run(text)
    _docx_apply_run_font(
        run,
        family=palette["body_font"],
        size_pt=10.5,
        color_hex=palette["muted"],
        italic=italic,
    )


def _docx_add_bullet(document, text: str, *, palette: dict):
    """Bulleted list item using Word's built-in 'List Bullet' style so
    the file opens with proper bullet formatting in Word + Google Docs.
    The style is part of the default template; no extra wiring needed."""
    paragraph = document.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.space_after = _docx_pt(2)
    run = paragraph.add_run(text)
    _docx_apply_run_font(
        run,
        family=palette["body_font"],
        size_pt=10.5,
        color_hex=palette["ink"],
    )


def _docx_add_paragraph_text(document, text: str, *, palette: dict, font_key: str = "body_font", size_pt: float = 11, italic: bool = False, color_key: str = "ink"):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = _docx_pt(4)
    run = paragraph.add_run(text)
    _docx_apply_run_font(
        run,
        family=palette[font_key],
        size_pt=size_pt,
        color_hex=palette[color_key],
        italic=italic,
    )
    return paragraph


def _docx_resume_summary_block(document, artifact: TailoredResumeArtifact, *, palette: dict):
    _docx_resume_section_heading(document, "Summary", palette=palette)
    # Prose summary uses the theme's prose font — the shared Arial
    # sans family across every theme since the 2026-05-21 typography
    # unification.
    _docx_add_paragraph_text(
        document,
        artifact.professional_summary or "No professional summary generated.",
        palette=palette,
        font_key="prose_font",
        size_pt=11,
        color_key="ink",
    )


def _docx_resume_skills_block(document, artifact: TailoredResumeArtifact, *, palette: dict):
    _docx_resume_section_heading(document, "Core Skills", palette=palette)
    categories = getattr(artifact, "skill_categories", None) or {}
    if categories:
        # One paragraph per category, with a bold accent-colored label
        # ("Languages & Tools: Python, SQL, ..."). Mirrors the HTML's
        # .resume-skill-category styling.
        for label, items in categories.items():
            label_clean = str(label or "").strip()
            cleaned = [str(item or "").strip() for item in items if str(item or "").strip()]
            if not label_clean or not cleaned:
                continue
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.space_before = _docx_pt(0)
            paragraph.paragraph_format.space_after = _docx_pt(2)
            label_run = paragraph.add_run(f"{label_clean}: ")
            _docx_apply_run_font(
                label_run,
                family=palette["body_font"],
                size_pt=11,
                color_hex=palette["accent"],
                bold=True,
            )
            value_run = paragraph.add_run(", ".join(cleaned))
            _docx_apply_run_font(
                value_run,
                family=palette["body_font"],
                size_pt=11,
                color_hex=palette["ink"],
            )
        return

    skills = [str(s).strip() for s in (artifact.highlighted_skills or []) if str(s or "").strip()]
    if skills:
        _docx_add_paragraph_text(
            document,
            " | ".join(skills),
            palette=palette,
            font_key="body_font",
            size_pt=11,
            color_key="ink",
        )
    else:
        _docx_add_paragraph_text(
            document,
            "No highlighted skills were generated.",
            palette=palette,
            font_key="body_font",
            size_pt=10.5,
            italic=True,
            color_key="muted",
        )


def _docx_resume_experience_block(document, artifact: TailoredResumeArtifact, *, palette: dict) -> bool:
    entries = list(artifact.experience_entries or [])
    if not entries:
        return False
    _docx_resume_section_heading(document, "Experience", palette=palette)
    for entry in entries:
        title = (entry.title or "Relevant Experience").strip()
        date_parts = [part for part in [entry.start, entry.end] if part]
        dates = " - ".join(date_parts) if date_parts else ""
        _docx_add_role_row(document, title=title, dates=dates, palette=palette)
        meta_parts = [part for part in [entry.organization, entry.location] if part]
        if meta_parts:
            _docx_add_meta_line(document, " | ".join(meta_parts), palette=palette)
        bullets = [str(b).strip() for b in (entry.bullets or []) if str(b or "").strip()]
        if bullets:
            for bullet in bullets:
                _docx_add_bullet(document, bullet, palette=palette)
        else:
            _docx_add_paragraph_text(
                document,
                "No grounded bullet points were generated for this role.",
                palette=palette,
                font_key="body_font",
                size_pt=10.5,
                italic=True,
                color_key="muted",
            )
    return True


def _docx_resume_projects_block(document, artifact: TailoredResumeArtifact, *, palette: dict) -> bool:
    projects = list(artifact.project_entries or [])
    if not projects:
        return False
    _docx_resume_section_heading(document, "Projects", palette=palette)
    for project in projects:
        name = (project.name or "Project").strip()
        date_parts = [part for part in [project.start, project.end] if part]
        dates = " - ".join(date_parts) if date_parts else ""
        _docx_add_role_row(document, title=name, dates=dates, palette=palette)
        if project.description:
            _docx_add_meta_line(document, project.description, palette=palette, italic=False)
        bullets = [str(b).strip() for b in (project.bullets or []) if str(b or "").strip()]
        for bullet in bullets:
            _docx_add_bullet(document, bullet, palette=palette)
        meta_parts = []
        if project.technologies:
            meta_parts.append("Tech: " + ", ".join(project.technologies))
        if project.link:
            meta_parts.append("Link: " + project.link)
        if meta_parts:
            _docx_add_meta_line(document, " | ".join(meta_parts), palette=palette, italic=True)
    return True


def _docx_resume_education_block(document, artifact: TailoredResumeArtifact, *, palette: dict):
    _docx_resume_section_heading(document, "Education", palette=palette)
    entries = list(artifact.education_entries or [])
    if not entries:
        _docx_add_paragraph_text(
            document,
            "No education entries were available.",
            palette=palette,
            font_key="body_font",
            size_pt=10.5,
            italic=True,
            color_key="muted",
        )
        return
    for entry in entries:
        institution = (entry.institution or "Education").strip()
        degree_parts = [part for part in [entry.degree, entry.field_of_study] if part]
        date_parts = [part for part in [entry.start, entry.end] if part]
        dates = " - ".join(date_parts) if date_parts else ""
        _docx_add_role_row(document, title=institution, dates=dates, palette=palette)
        if degree_parts:
            _docx_add_meta_line(document, " - ".join(degree_parts), palette=palette, italic=False)


def _docx_resume_publications_block(document, artifact: TailoredResumeArtifact, *, palette: dict) -> bool:
    items = [str(item).strip() for item in (artifact.publication_entries or []) if str(item or "").strip()]
    if not items:
        return False
    _docx_resume_section_heading(document, "Publications", palette=palette)
    for item in items:
        _docx_add_bullet(document, item, palette=palette)
    return True


def _docx_resume_certifications_block(document, artifact: TailoredResumeArtifact, *, palette: dict) -> bool:
    items = [str(item).strip() for item in (artifact.certifications or []) if str(item or "").strip()]
    if not items:
        return False
    _docx_resume_section_heading(document, "Certifications", palette=palette)
    for item in items:
        _docx_add_bullet(document, item, palette=palette)
    return True


def _build_resume_docx(artifact: TailoredResumeArtifact) -> bytes:
    """Render a structured TailoredResumeArtifact to DOCX bytes.

    Mirrors the section ordering / empty-section policy of
    `_build_structured_resume_body_classic`: Summary, Skills, Education
    always render even when sparse; Experience, Projects, Publications,
    Certifications drop entirely when empty. Section order honors
    `artifact.section_order` and falls back to
    `_DEFAULT_RESUME_SECTION_ORDER` for legacy callers.

    Theme is read from `artifact.theme`; supported values are the five
    user-facing themes (`professional_neutral`, `classic_ats`,
    `modern_blue`, `creative_warm`, `architect_mono`) — all sharing the
    one Arial sans family, differentiated by colour / paper / header.
    Unknown values fall back to `classic_ats`.
    """
    from docx import Document

    palette = _resolve_docx_palette(artifact.theme)
    document = Document()
    _docx_set_page_margins(document, inches=_DOCX_PAGE_MARGIN_INCHES)

    # Default style baseline so paragraphs without a per-run font fall
    # back cleanly when opened in Word's Style pane.
    normal_style = document.styles["Normal"]
    normal_style.font.name = palette["body_font"]
    from docx.shared import Pt as _Pt

    normal_style.font.size = _Pt(11)

    _docx_add_resume_header(document, artifact, palette=palette)

    section_renderers = {
        "summary": lambda: (_docx_resume_summary_block(document, artifact, palette=palette), True)[1],
        "skills": lambda: (_docx_resume_skills_block(document, artifact, palette=palette), True)[1],
        "experience": lambda: _docx_resume_experience_block(document, artifact, palette=palette),
        "projects": lambda: _docx_resume_projects_block(document, artifact, palette=palette),
        "education": lambda: (_docx_resume_education_block(document, artifact, palette=palette), True)[1],
        "publications": lambda: _docx_resume_publications_block(document, artifact, palette=palette),
        "certifications": lambda: _docx_resume_certifications_block(document, artifact, palette=palette),
    }

    order = list(artifact.section_order) if artifact.section_order else list(_DEFAULT_RESUME_SECTION_ORDER)
    seen: set[str] = set()
    for section_name in order:
        if section_name in seen:
            continue
        seen.add(section_name)
        renderer = section_renderers.get(section_name)
        if renderer is not None:
            renderer()
    # Append any sections the agent forgot to mention so we never lose
    # rendered content when the agent emits a partial order.
    for section_name in _DEFAULT_RESUME_SECTION_ORDER:
        if section_name in seen:
            continue
        renderer = section_renderers.get(section_name)
        if renderer is not None:
            renderer()

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _build_cover_letter_docx(artifact: CoverLetterArtifact) -> bytes:
    """Render a CoverLetterArtifact to DOCX bytes.

    The cover letter artifact only exposes a flat `markdown` field (the
    structured paragraphs live upstream in `CoverLetterAgentOutput` and
    aren't on the artifact). Parse the markdown into blocks via the
    existing `_parse_markdown_blocks` helper and emit each block as a
    matching DOCX paragraph / list. Title is split via
    `_split_cover_letter_title` so the heading + role-eyebrow read the
    same way as the HTML render.

    Theme is read from `artifact.theme`; every theme uses the shared
    Arial sans family (prose font) for the letter body — the cover
    letter is letter-shaped prose in any palette. The theme switch
    only changes ink / muted / accent / line colors.
    """
    from docx import Document

    palette = _resolve_docx_palette(artifact.theme)
    document = Document()
    _docx_set_page_margins(document, inches=_DOCX_PAGE_MARGIN_INCHES)

    normal_style = document.styles["Normal"]
    normal_style.font.name = palette["prose_font"]
    from docx.shared import Pt as _Pt

    normal_style.font.size = _Pt(11.4)

    header_title, header_subtitle = _split_cover_letter_title(artifact.title or "Cover Letter")

    title_paragraph = document.add_paragraph()
    title_paragraph.paragraph_format.space_after = _docx_pt(2)
    title_run = title_paragraph.add_run(header_title)
    _docx_apply_run_font(
        title_run,
        family=palette["heading_font"],
        size_pt=18,
        color_hex=palette["ink"],
        bold=True,
    )
    if header_subtitle:
        sub_paragraph = document.add_paragraph()
        sub_paragraph.paragraph_format.space_after = _docx_pt(8)
        sub_run = sub_paragraph.add_run(header_subtitle.upper())
        _docx_apply_run_font(
            sub_run,
            family=palette["body_font"],
            size_pt=10,
            color_hex=palette["muted"],
        )
    _docx_add_bottom_border(
        title_paragraph if not header_subtitle else sub_paragraph,
        color_hex=palette["accent"],
        size_eighths_pt=8,
    )

    blocks = _parse_markdown_blocks(artifact.markdown or "")
    # Drop the leading H1 (already rendered as the header) and any
    # leading rule, mirroring the HTML render's title strip.
    deferred_blocks = []
    for index, (kind, payload) in enumerate(blocks):
        if index == 0 and kind == "title":
            continue
        deferred_blocks.append((kind, payload))

    for kind, payload in deferred_blocks:
        if kind in {"heading", "subheading"}:
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.space_before = _docx_pt(8)
            paragraph.paragraph_format.space_after = _docx_pt(2)
            run = paragraph.add_run(_strip_inline_markup(str(payload or "")))
            _docx_apply_run_font(
                run,
                family=palette["heading_font"],
                size_pt=12.5,
                color_hex=palette["ink"],
                bold=True,
            )
            continue
        if kind == "paragraph":
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.space_after = _docx_pt(8)
            run = paragraph.add_run(_strip_inline_markup(str(payload or "")))
            _docx_apply_run_font(
                run,
                family=palette["prose_font"],
                size_pt=11.4,
                color_hex=palette["ink"],
            )
            continue
        if kind == "list":
            for item in payload or []:
                if not isinstance(item, dict):
                    continue
                if item.get("kind") != "list_paragraph":
                    continue
                paragraph = document.add_paragraph(style="List Bullet")
                paragraph.paragraph_format.space_after = _docx_pt(2)
                run = paragraph.add_run(_strip_inline_markup(str(item.get("text", "") or "")))
                _docx_apply_run_font(
                    run,
                    family=palette["prose_font"],
                    size_pt=11.4,
                    color_hex=palette["ink"],
                )
            continue
        if kind == "rule":
            divider = document.add_paragraph()
            _docx_add_bottom_border(
                divider,
                color_hex=palette["line"],
                size_eighths_pt=4,
            )
            continue
        # Unhandled kind (code_block, etc.) — skip silently. The cover
        # letter agent doesn't emit code blocks today, but if that
        # changes we can extend this.

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


_INLINE_MARKUP_TAG = re.compile(r"<[^>]+>")


def _strip_inline_markup(text: str) -> str:
    """The markdown-it tree → block parser emits inline children with
    HTML-style tags (`<b>...</b>`, `<i>...</i>`). DOCX runs don't take
    raw HTML, so flatten the markup to plain text for now. Phase 1
    accepts the loss of bold/italic styling inside paragraphs; if QA
    flags it as a problem we can teach the parser to emit per-run
    styling instead.
    """
    return _INLINE_MARKUP_TAG.sub("", text or "").strip()


def export_docx_bytes(report: CoverLetterArtifact | TailoredResumeArtifact) -> bytes:
    """Render an artifact to DOCX bytes.

    Phase 1 implements the `classic_ats` theme only; the second theme
    (`professional_neutral`) lands in Phase 4 with a palette switch.
    """
    try:
        if isinstance(report, TailoredResumeArtifact):
            return _build_resume_docx(report)
        if isinstance(report, CoverLetterArtifact):
            return _build_cover_letter_docx(report)
        raise ExportError(
            "Unsupported artifact type for DOCX export.",
            details=type(report).__name__,
        )
    except ExportError:
        raise
    except Exception as error:
        log_event(
            LOGGER,
            logging.ERROR,
            "docx_export_failed",
            "DOCX export failed.",
            report_title=getattr(report, "title", ""),
            filename_stem=getattr(report, "filename_stem", ""),
            error_type=type(error).__name__,
        )
        raise ExportError(
            "DOCX export failed. Try the PDF download instead.",
            details=str(error),
        ) from error
