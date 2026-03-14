from io import BytesIO
import logging

import docx
from pypdf import PdfReader

from src.errors import ParsingError
from src.logging_utils import get_logger, log_event
from src.parsers.common import decode_text, detect_file_type, read_file_bytes
from src.schemas import ResumeDocument


LOGGER = get_logger(__name__)


def _extract_text_from_pdf(file_bytes):
    try:
        reader = PdfReader(BytesIO(file_bytes))
    except Exception as exc:
        raise ParsingError("Failed to open the PDF resume.") from exc

    pages = [(page.extract_text() or "").strip() for page in reader.pages]
    return "\n".join(page for page in pages if page).strip()


def _extract_text_from_docx(file_bytes):
    try:
        document = docx.Document(BytesIO(file_bytes))
    except Exception as exc:
        raise ParsingError("Failed to open the DOCX resume.") from exc

    paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs]
    return "\n".join(paragraph for paragraph in paragraphs if paragraph).strip()


def _parse_resume_bytes(file_bytes, file_type, source):
    if file_type == "application/pdf":
        text = _extract_text_from_pdf(file_bytes)
        filetype = "PDF"
    elif file_type in {
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "application/msword",
    }:
        text = _extract_text_from_docx(file_bytes)
        filetype = "DOCX"
    elif file_type == "text/plain":
        text = decode_text(file_bytes).strip()
        filetype = "TXT"
    else:
        raise ParsingError("Unsupported resume file type. Use PDF, DOCX, or TXT.")

    if not text:
        raise ParsingError("The resume was parsed, but no readable text was extracted.")

    return ResumeDocument(text=text, filetype=filetype, source=source)


def parse_resume_document(file, source="uploaded"):
    file_name = getattr(file, "name", source)
    try:
        file_type = detect_file_type(file)
        file_bytes = read_file_bytes(file)
        return _parse_resume_bytes(file_bytes, file_type, source)
    except ParsingError as error:
        log_event(
            LOGGER,
            logging.ERROR,
            "resume_parse_failed",
            "Resume parsing failed.",
            source=source,
            file_name=file_name,
            file_type=locals().get("file_type"),
            file_size_bytes=len(locals().get("file_bytes", b"")),
            error_type=type(error).__name__,
        )
        raise

