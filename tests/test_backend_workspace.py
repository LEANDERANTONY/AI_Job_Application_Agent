import base64

import pytest
from fastapi.testclient import TestClient

from backend.app import app
from backend.request_auth import get_optional_auth_tokens, get_required_auth_tokens


client = TestClient(app)


@pytest.fixture(autouse=True)
def _satisfy_llm_route_auth():
    """T5 of the token-meter migration requires login on every LLM
    route — resume / JD parse, analyze, assistant, and résumé-builder
    start / message / generate (anonymous = no user_id = un-metered).

    These tests exercise that route FUNCTIONALITY, not the gate, and
    predate it: most run anonymously, a few pass dummy auth headers and
    monkeypatch the auth resolution. Overriding ``get_required_auth_
    tokens`` with ``get_optional_auth_tokens`` restores the EXACT
    pre-T5 dependency behaviour — real tokens when headers are present,
    ``(None, None)`` when absent, never a 401 — so every existing
    assertion in this file holds unchanged.

    Dedicated 401 coverage for the gate itself lives in
    ``tests/backend/test_llm_route_login_required.py`` (no override)."""
    app.dependency_overrides[get_required_auth_tokens] = get_optional_auth_tokens
    yield
    app.dependency_overrides.pop(get_required_auth_tokens, None)


@pytest.fixture(autouse=True)
def _export_gate_as_paid_tier(monkeypatch):
    """Resolve the export-entitlement tier as a paid user for every
    test in this module.

    The Free entitlement (PDF + professional_neutral only; DOCX and
    classic_ats are Pro/Business) is exhaustively covered in
    ``tests/backend/test_export_entitlement.py``. The export tests
    HERE pre-date the gate and exercise export *mechanics* — DOCX
    round-trip, snapshot/theme forwarding, multi-role rendering,
    unknown-session 400s. Without this they'd all 429 on DOCX /
    classic_ats and stop testing what they're for. ``_resolve_export_
    tier`` is only reached from the two export routes, so this is a
    no-op for every non-export test in the file.
    """
    monkeypatch.setattr(
        "backend.routers.workspace._resolve_export_tier",
        lambda *args, **kwargs: "business",
    )


def _encode_text_file_payload(filename: str, text: str):
    return {
        "filename": filename,
        "mime_type": "text/plain",
        "content_base64": base64.b64encode(text.encode("utf-8")).decode("utf-8"),
    }


def test_workspace_resume_upload_parses_candidate_profile():
    response = client.post(
        "/api/workspace/resume/upload",
        json=_encode_text_file_payload(
            "resume.txt",
            (
                "Leander Antony\n"
                "Chennai, India\n"
                "leander@example.com\n"
                "Python SQL Docker Communication\n"
                "Experience\n"
                "AI Engineer, Example Labs\n"
                "Jan 2023 - Jan 2025\n"
                "Built production ML APIs and evaluation workflows.\n"
            ),
        ),
    )

    assert response.status_code == 200
    payload = response.json()
    assert payload["resume_document"]["filetype"] == "TXT"
    assert payload["candidate_profile"]["full_name"] == "Leander Antony"
    assert "Python" in payload["candidate_profile"]["skills"]


def test_workspace_job_description_upload_parses_role_and_summary():
    response = client.post(
        "/api/workspace/job-description/upload",
        json=_encode_text_file_payload(
            "job.txt",
            (
                "Machine Learning Engineer\n"
                "Location: Bengaluru, India\n"
                "Required: Python, SQL, Docker, communication.\n"
                "Need 3+ years of experience.\n"
                "Preferred: AWS and production LLM systems.\n"
            ),
        ),
    )

    assert response.status_code == 200
    payload = response.json()
    assert payload["job_description"]["title"] == "Machine Learning Engineer"
    assert payload["job_description"]["requirements"]["hard_skills"]
    assert payload["jd_summary_view"]["sections"]


def test_resume_builder_session_can_progress_to_review():
    start_response = client.post("/api/workspace/resume-builder/start")

    assert start_response.status_code == 200
    session_id = start_response.json()["session_id"]

    message_payloads = [
        "Leander Antony\nChennai, India\nleander@example.com\n+91 9999999999\nlinkedin.com/in/leander",
        "Machine Learning Engineer\nAI engineer with product-focused ML experience across APIs, evaluation, and developer tooling.",
        "AI Engineer at Example Labs\nJan 2023 - Present\nBuilt ML APIs used by production teams.\nImproved evaluation workflows for model quality.",
        "Anna University | B.E. Computer Science\nAWS Certified Machine Learning Specialty",
        "Python, FastAPI, Docker, LLMs, SQL",
    ]

    latest_payload = None
    for message in message_payloads:
        response = client.post(
            "/api/workspace/resume-builder/message",
            json={
                "session_id": session_id,
                "message": message,
                "input_mode": "text",
            },
        )
        assert response.status_code == 200
        latest_payload = response.json()

    assert latest_payload is not None
    assert latest_payload["status"] == "reviewing"
    assert latest_payload["current_step"] == "review"
    assert latest_payload["ready_to_generate"] is True
    assert latest_payload["draft_profile"]["target_role"] == "Machine Learning Engineer"
    assert "Python" in latest_payload["draft_profile"]["skills"]


def test_resume_builder_generate_and_commit_returns_workspace_profile():
    start_response = client.post("/api/workspace/resume-builder/start")

    assert start_response.status_code == 200
    session_id = start_response.json()["session_id"]

    for message in [
        "Leander Antony\nChennai, India\nleander@example.com\n+91 9999999999",
        "Machine Learning Engineer\nAI engineer with ML platform and applied AI experience.",
        "AI Engineer at Example Labs\nJan 2023 - Present\nBuilt ML APIs and developer workflows.",
        "Anna University | B.E. Computer Science",
        "Python, FastAPI, Docker, SQL",
    ]:
        response = client.post(
            "/api/workspace/resume-builder/message",
            json={
                "session_id": session_id,
                "message": message,
                "input_mode": "text",
            },
        )
        assert response.status_code == 200

    generate_response = client.post(
        "/api/workspace/resume-builder/generate",
        json={"session_id": session_id},
    )

    assert generate_response.status_code == 200
    generated_payload = generate_response.json()
    assert generated_payload["status"] == "ready"
    assert generated_payload["generated_resume_markdown"]
    assert generated_payload["candidate_profile"]["full_name"] == "Leander Antony"

    commit_response = client.post(
        "/api/workspace/resume-builder/commit",
        json={"session_id": session_id},
    )

    assert commit_response.status_code == 200
    commit_payload = commit_response.json()
    assert commit_payload["resume_document"]["source"] == "assistant_builder"
    assert commit_payload["candidate_profile"]["source"] == "assistant_builder"
    assert commit_payload["generated_resume_markdown"]


def test_resume_builder_update_route_applies_draft_changes():
    start_response = client.post("/api/workspace/resume-builder/start")

    assert start_response.status_code == 200
    session_id = start_response.json()["session_id"]

    update_response = client.post(
        "/api/workspace/resume-builder/update",
        json={
            "session_id": session_id,
            "draft_profile": {
                "full_name": "Leander Antony",
                "location": "Chennai, India",
                "target_role": "Machine Learning Engineer",
                "professional_summary": "AI engineer with product and platform experience.",
                "skills": ["Python", "FastAPI", "Docker"],
                "contact_lines": ["leander@example.com", "+91 9999999999"],
            },
        },
    )

    assert update_response.status_code == 200
    payload = update_response.json()
    assert payload["draft_profile"]["full_name"] == "Leander Antony"
    assert payload["draft_profile"]["target_role"] == "Machine Learning Engineer"
    assert "Docker" in payload["draft_profile"]["skills"]


def test_resume_builder_update_clears_stale_generated_resume():
    """A draft edit after generation must clear the generated resume so
    the on-screen preview can't silently drift from a fresh export."""
    start_response = client.post("/api/workspace/resume-builder/start")
    assert start_response.status_code == 200
    session_id = start_response.json()["session_id"]

    for message in [
        "Leander Antony\nChennai, India\nleander@example.com",
        "Machine Learning Engineer\nAI engineer with applied ML experience.",
        "AI Engineer at Example Labs\nJan 2023 - Present\nBuilt ML APIs.",
        "Anna University | B.E. Computer Science",
        "Python, FastAPI, Docker, SQL",
    ]:
        resp = client.post(
            "/api/workspace/resume-builder/message",
            json={
                "session_id": session_id,
                "message": message,
                "input_mode": "text",
            },
        )
        assert resp.status_code == 200

    generate_response = client.post(
        "/api/workspace/resume-builder/generate",
        json={"session_id": session_id},
    )
    assert generate_response.status_code == 200
    assert generate_response.json()["generated_resume_markdown"]

    # Editing the draft must wipe the now-stale generated resume.
    update_response = client.post(
        "/api/workspace/resume-builder/update",
        json={
            "session_id": session_id,
            "draft_profile": {"target_role": "Senior ML Engineer"},
        },
    )
    assert update_response.status_code == 200
    update_payload = update_response.json()
    assert update_payload["generated_resume_markdown"] == ""
    assert update_payload["ready_to_commit"] is False
    # Still ready_to_generate so the user can immediately regenerate.
    assert update_payload["ready_to_generate"] is True


def test_resume_builder_reset_route_clears_session_to_first_state():
    """Start over: the reset route clears the draft + generated resume
    and reuses the SAME session_id (so no new resume_builder_sessions
    quota credit is charged)."""
    start_response = client.post("/api/workspace/resume-builder/start")
    assert start_response.status_code == 200
    session_id = start_response.json()["session_id"]

    client.post(
        "/api/workspace/resume-builder/update",
        json={
            "session_id": session_id,
            "draft_profile": {
                "full_name": "Leander Antony",
                "target_role": "Machine Learning Engineer",
                "skills": ["Python", "FastAPI"],
            },
        },
    )

    reset_response = client.post(
        "/api/workspace/resume-builder/reset",
        json={"session_id": session_id},
    )
    assert reset_response.status_code == 200
    payload = reset_response.json()
    # Same session_id — a clear/restart, not a new quota-charged session.
    assert payload["session_id"] == session_id
    # Draft wiped back to empty; back at the first step.
    assert payload["draft_profile"]["full_name"] == ""
    assert payload["draft_profile"]["target_role"] == ""
    assert payload["draft_profile"]["skills"] == []
    assert payload["generated_resume_markdown"] == ""
    assert payload["current_step"] == "basics"


def test_resume_builder_reset_route_rejects_unknown_session():
    response = client.post(
        "/api/workspace/resume-builder/reset",
        json={"session_id": "does-not-exist"},
    )
    assert response.status_code == 400


def test_resume_builder_latest_endpoint_returns_saved_session(monkeypatch):
    monkeypatch.setattr(
        "backend.routers.workspace.load_latest_resume_builder_session",
        lambda access_token, refresh_token: {
            "status": "available",
            "session": {
                "session_id": "builder-123",
                "status": "collecting",
                "current_step": "experience",
                "completed_steps": 2,
                "total_steps": 5,
                "progress_percent": 40,
                "assistant_message": "Tell me about your most relevant experience.",
                "draft_profile": {
                    "full_name": "Leander Antony",
                    "location": "Chennai, India",
                    "contact_lines": ["leander@example.com"],
                    "target_role": "Machine Learning Engineer",
                    "professional_summary": "AI engineer with ML platform experience.",
                    "experience_notes": "",
                    "education_notes": "",
                    "skills": ["Python"],
                    "certifications": [],
                },
                "generated_resume_markdown": "",
                "generated_resume_plain_text": "",
                "ready_to_generate": False,
                "ready_to_commit": False,
            },
        },
    )

    response = client.get(
        "/api/workspace/resume-builder/latest",
        headers={
            "X-Auth-Access-Token": "access-token",
            "X-Auth-Refresh-Token": "refresh-token",
        },
    )

    assert response.status_code == 200
    payload = response.json()
    assert payload["status"] == "available"
    assert payload["session"]["current_step"] == "experience"
    assert payload["session"]["draft_profile"]["target_role"] == "Machine Learning Engineer"


def test_resume_builder_export_round_trip_produces_docx_bytes():
    """Phase 5: drive a session through start -> message x5 -> generate
    -> /resume-builder/export and verify the response is base64-encoded
    DOCX bytes that python-docx can re-parse."""
    import base64
    from io import BytesIO

    from docx import Document

    start_response = client.post("/api/workspace/resume-builder/start")
    assert start_response.status_code == 200
    session_id = start_response.json()["session_id"]

    answers = [
        "Leander Antony, Chennai India. leander@example.com, +91 9999999999",
        "Senior ML Engineer. Built distributed Python systems for 5 years.",
        (
            "AI Engineer at Example Labs (Jan 2023 - Present). "
            "Built ML APIs. Reduced latency 30%."
        ),
        "Anna University, B.E. Computer Science (2016-2020)",
        "Python, FastAPI, AWS, Docker, SQL",
    ]
    for answer in answers:
        message_response = client.post(
            "/api/workspace/resume-builder/message",
            json={
                "session_id": session_id,
                "message": answer,
                "input_mode": "text",
            },
        )
        assert message_response.status_code == 200

    generate_response = client.post(
        "/api/workspace/resume-builder/generate",
        json={"session_id": session_id},
    )
    assert generate_response.status_code == 200

    export_response = client.post(
        "/api/workspace/resume-builder/export",
        json={
            "session_id": session_id,
            "export_format": "docx",
            "theme": "classic_ats",
        },
    )

    assert export_response.status_code == 200
    payload = export_response.json()
    assert payload["status"] == "ready"
    assert payload["export_format"] == "docx"
    assert payload["file_name"].endswith(".docx")
    assert payload["mime_type"] == (
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    assert payload["theme"] == "classic_ats"
    # The artifact title flows from the synthesized job_description
    # (which uses the user's target_role) through
    # `_synthesize_resume_builder_artifact`'s rename step.
    assert "Leander Antony" in payload["artifact_title"]
    assert "Resume" in payload["artifact_title"]

    # Bytes round-trip cleanly through python-docx — basic structural
    # check that the file isn't truncated or wrong-format.
    raw_bytes = base64.b64decode(payload["content_base64"])
    assert len(raw_bytes) > 5_000
    assert raw_bytes.startswith(b"PK")
    document = Document(BytesIO(raw_bytes))
    paragraph_texts = [p.text for p in document.paragraphs]
    # Header carries the user-typed name + email.
    assert any("Leander Antony" in text for text in paragraph_texts)
    assert any("leander@example.com" in text for text in paragraph_texts)
    # Required section headings render even on a sparse builder draft.
    assert "SUMMARY" in paragraph_texts
    assert "EDUCATION" in paragraph_texts


def test_resume_builder_export_supports_pdf_format():
    """The PDF branch should also work end-to-end. We don't try to
    parse the PDF (Phase 1 already covers that path); just verify the
    bytes are returned with the right mime + filename."""
    import base64

    start_response = client.post("/api/workspace/resume-builder/start")
    session_id = start_response.json()["session_id"]
    for answer in [
        "Mei Chen, Singapore. mei@example.sg",
        "Senior Data Engineer. ETL platform background.",
        "Data Engineer at Acme (2020-2023). Owned warehouse ingestion.",
        "NUS, B.Comp",
        "Python, Airflow, SQL",
    ]:
        client.post(
            "/api/workspace/resume-builder/message",
            json={"session_id": session_id, "message": answer, "input_mode": "text"},
        )
    client.post(
        "/api/workspace/resume-builder/generate",
        json={"session_id": session_id},
    )

    export_response = client.post(
        "/api/workspace/resume-builder/export",
        json={
            "session_id": session_id,
            "export_format": "pdf",
            "theme": "professional_neutral",
        },
    )

    assert export_response.status_code == 200
    payload = export_response.json()
    assert payload["export_format"] == "pdf"
    assert payload["file_name"].endswith(".pdf")
    assert payload["mime_type"] == "application/pdf"
    assert payload["theme"] == "professional_neutral"
    raw_bytes = base64.b64decode(payload["content_base64"])
    assert len(raw_bytes) > 1_000
    assert raw_bytes.startswith(b"%PDF")


def test_resume_builder_export_returns_400_when_session_unknown():
    """Old / wrong session_id → friendly 400, mirrors the other
    resume-builder routes' error contract."""
    response = client.post(
        "/api/workspace/resume-builder/export",
        json={
            "session_id": "session-that-never-existed",
            "export_format": "docx",
            "theme": "classic_ats",
        },
    )
    assert response.status_code == 400
    assert "Resume builder session" in response.json()["detail"]


def test_resume_builder_export_rejects_unsupported_format():
    """Pydantic Literal validates the format; markdown isn't a
    supported export format anywhere now."""
    start_response = client.post("/api/workspace/resume-builder/start")
    session_id = start_response.json()["session_id"]

    response = client.post(
        "/api/workspace/resume-builder/export",
        json={
            "session_id": session_id,
            "export_format": "markdown",
            "theme": "classic_ats",
        },
    )
    assert response.status_code == 422
    assert "export_format" in str(response.json())


def test_resume_builder_preview_route_returns_themed_html():
    """The themed-preview route renders the generated base resume as a
    standalone HTML document — no download, no LLM call. Requesting a
    Pro-gated theme (``modern_blue``) and getting a 200 also proves the
    preview is NOT entitlement-gated: every tier may preview every
    theme (that's the conversion surface); only /export is gated."""
    start_response = client.post("/api/workspace/resume-builder/start")
    session_id = start_response.json()["session_id"]
    for answer in [
        "Mei Chen, Singapore. mei@example.sg",
        "Senior Data Engineer. ETL platform background.",
        "Data Engineer at Acme (2020-2023). Owned warehouse ingestion.",
        "NUS, B.Comp",
        "Python, Airflow, SQL",
    ]:
        client.post(
            "/api/workspace/resume-builder/message",
            json={"session_id": session_id, "message": answer, "input_mode": "text"},
        )
    client.post(
        "/api/workspace/resume-builder/generate",
        json={"session_id": session_id},
    )

    preview_response = client.post(
        "/api/workspace/resume-builder/preview",
        json={"session_id": session_id, "theme": "modern_blue"},
    )

    assert preview_response.status_code == 200
    payload = preview_response.json()
    assert payload["status"] == "ready"
    # The picked theme is echoed back so the client can confirm what it
    # is looking at; "modern_blue" is a Pro-gated download theme yet
    # previews fine here — the gate lives only on /export.
    assert payload["resume_theme"] == "modern_blue"
    html = payload["html"]
    assert "<!doctype html" in html.lower()
    assert "Mei Chen" in html


def test_resume_builder_preview_route_returns_400_when_session_unknown():
    """Old / evicted session_id → friendly 400, mirroring the other
    resume-builder routes' error contract."""
    response = client.post(
        "/api/workspace/resume-builder/preview",
        json={
            "session_id": "session-that-never-existed",
            "theme": "professional_neutral",
        },
    )
    assert response.status_code == 400
    assert "Resume builder session" in response.json()["detail"]


def test_resume_builder_preview_route_rejects_unknown_theme():
    """The request model's theme Literal validates at the FastAPI
    boundary — a theme outside the supported set is a clean 422, never
    a fall-through to the renderer."""
    start_response = client.post("/api/workspace/resume-builder/start")
    session_id = start_response.json()["session_id"]

    response = client.post(
        "/api/workspace/resume-builder/preview",
        json={"session_id": session_id, "theme": "not_a_real_theme"},
    )
    assert response.status_code == 422
    assert "theme" in str(response.json())


def test_resume_builder_experience_parser_splits_multi_role_single_line():
    """Users often squash multiple roles onto one line ("X at A 2020-Present,
    prior at B 2017-2020"). The parser used to collapse this into ONE
    WorkExperience whose organization swallowed the entire suffix. Now
    we split on transition markers and emit one entry per role."""
    from backend.services.resume_builder_service import _build_experience_entries

    notes = (
        "Senior Backend Engineer at TechCorp from 2020-Present, "
        "prior at FinStart 2017-2020"
    )
    entries = _build_experience_entries(notes)

    assert len(entries) == 2
    first, second = entries
    assert first.title == "Senior Backend Engineer"
    assert first.organization == "TechCorp"
    assert first.start == "2020"
    assert first.end == "Present"
    # Second sub-headline starts with "at" — title falls back, org is set.
    assert second.organization == "FinStart"
    assert second.start == "2017"
    assert second.end == "2020"
    # Description must NOT regurgitate the headline — that was the source
    # of the duplicate-meta-line bug (organization had the whole blob and
    # description seeded an identical bullet downstream).
    assert first.description == ""
    assert second.description == ""


def test_resume_builder_experience_parser_separates_headline_from_bullets():
    """Single-line input where bullets follow the headline as sentences
    ("AI Engineer at Example Labs (Jan 2023 - Present). Built ML APIs.
    Reduced latency 30%.") should produce ONE entry with the dates
    extracted to start/end, the organization clean of date noise, and
    the bullet sentences in description (not the headline)."""
    from backend.services.resume_builder_service import _build_experience_entries

    notes = (
        "AI Engineer at Example Labs (Jan 2023 - Present). "
        "Built ML APIs. Reduced latency 30%."
    )
    entries = _build_experience_entries(notes)

    assert len(entries) == 1
    entry = entries[0]
    assert entry.title == "AI Engineer"
    assert entry.organization == "Example Labs"
    assert entry.start == "Jan 2023"
    assert entry.end == "Present"
    bullets = entry.description.splitlines()
    # Trailing punctuation may be preserved ("Built ML APIs." vs "Built ML APIs"),
    # so substring-match each expected bullet rather than checking equality.
    assert any("Built ML APIs" in line for line in bullets)
    assert any("Reduced latency" in line for line in bullets)
    # The headline itself is NOT in description.
    assert "AI Engineer at Example Labs" not in entry.description


def test_resume_builder_experience_parser_groups_multiline_blocks():
    """Multi-line input where each role starts a fresh headline line,
    optionally followed by indented bullets, should produce one entry
    per headline with the right bullets attached."""
    from backend.services.resume_builder_service import _build_experience_entries

    notes = (
        "Senior Backend Engineer at TechCorp (2020-Present)\n"
        "- Built distributed systems\n"
        "- Led team of 5\n"
        "Backend Engineer at FinStart (2017-2020)\n"
        "- Optimized payment pipeline"
    )
    entries = _build_experience_entries(notes)

    assert len(entries) == 2
    first, second = entries
    assert first.title == "Senior Backend Engineer"
    assert first.organization == "TechCorp"
    assert first.start == "2020"
    assert first.end == "Present"
    assert "Built distributed systems" in first.description
    assert "Led team of 5" in first.description
    assert second.title == "Backend Engineer"
    assert second.organization == "FinStart"
    assert second.start == "2017"
    assert second.end == "2020"
    assert "Optimized payment pipeline" in second.description


def test_resume_builder_education_parser_splits_multi_degree_single_line():
    """Users often pack two degrees onto one line ("MS Computer Science
    Stanford 2017, BTech CS IIT Madras 2015"). Previously the whole line
    became one EducationEntry; now we split on commas + degree patterns."""
    from backend.services.resume_builder_service import _build_education_entries

    notes = "MS Computer Science Stanford 2017, BTech CS IIT Madras 2015"
    entries = _build_education_entries(notes)

    assert len(entries) == 2
    titles_by_year = {entry.start: entry for entry in entries}
    assert "2017" in titles_by_year
    assert "2015" in titles_by_year
    ms_entry = titles_by_year["2017"]
    btech_entry = titles_by_year["2015"]
    assert "Stanford" in ms_entry.institution
    # Degree should retain the abbreviation; field-of-study can land in
    # either degree or institution depending on the parser, but at minimum
    # the abbreviation has to come through.
    assert ms_entry.degree.lower().startswith("ms")
    assert "IIT Madras" in btech_entry.institution
    assert btech_entry.degree.lower().startswith("btech") or btech_entry.degree.lower().startswith("b.tech")


def test_resume_builder_export_renders_multi_role_content_in_docx():
    """End-to-end check that the parser fix lands on the page.

    Regression coverage for the "very sparse / not a real resume"
    feedback: a single-line, multi-role experience and a single-line,
    multi-degree education had collapsed to one entry each, so the
    rendered DOCX showed only one company and one degree. We now drive
    the same shape through the full pipeline and assert both
    employers + both schools appear in the document text.
    """
    import base64
    from io import BytesIO

    from docx import Document

    start_response = client.post("/api/workspace/resume-builder/start")
    assert start_response.status_code == 200
    session_id = start_response.json()["session_id"]

    answers = [
        "Priya Sharma, Bangalore. priya@gmail.com, +91 8000000000",
        "Senior Backend Engineer. Distributed systems and payments background.",
        # Two roles squashed onto one line via the "prior at" transition.
        "Senior Backend Engineer at TechCorp from 2020-Present, prior at FinStart 2017-2020",
        # Two degrees on one line.
        "MS Computer Science Stanford 2017, BTech CS IIT Madras 2015",
        "Python, PostgreSQL, Kafka, Docker",
    ]
    for answer in answers:
        message_response = client.post(
            "/api/workspace/resume-builder/message",
            json={
                "session_id": session_id,
                "message": answer,
                "input_mode": "text",
            },
        )
        assert message_response.status_code == 200

    generate_response = client.post(
        "/api/workspace/resume-builder/generate",
        json={"session_id": session_id},
    )
    assert generate_response.status_code == 200

    export_response = client.post(
        "/api/workspace/resume-builder/export",
        json={
            "session_id": session_id,
            "export_format": "docx",
            "theme": "classic_ats",
        },
    )
    assert export_response.status_code == 200
    payload = export_response.json()
    assert payload["status"] == "ready"

    raw_bytes = base64.b64decode(payload["content_base64"])
    document = Document(BytesIO(raw_bytes))
    body = "\n".join(p.text for p in document.paragraphs)

    # Both employers must show up — previously only one role survived.
    assert "TechCorp" in body, "Missing first employer (TechCorp) — multi-role split regressed"
    assert "FinStart" in body, "Missing second employer (FinStart) — multi-role split regressed"
    # Dates land in their own start/end fields, not inside the org meta.
    assert "2020" in body
    assert "2017" in body
    # Both schools must show up — previously the two degrees merged.
    assert "Stanford" in body, "Missing first school (Stanford) — multi-degree split regressed"
    assert "IIT Madras" in body, "Missing second school (IIT Madras) — multi-degree split regressed"


def test_resume_builder_structuring_uses_llm_when_service_available(monkeypatch):
    """LLM-first structuring: when an OpenAIService is plumbed through
    the /generate route, the candidate_profile.experience and .education
    arrays come from the model's structured payload, not the regex
    parser. Bullets are LLM-rewritten into ATS voice — the regex path
    can't produce them at all (the user gave bare role headlines)."""
    from backend.services.resume_builder_service import _SESSIONS

    captured_calls: list[str] = []

    class _StubOpenAIService:
        def is_available(self):
            return True

        def run_json_prompt(self, system, user, **kwargs):
            task = kwargs.get("task_name", "")
            captured_calls.append(task)
            if task == "resume_builder_structuring":
                # The structured response the LLM would emit for the
                # user's "Senior Backend Engineer at TechCorp..." prose.
                return {
                    "experience": [
                        {
                            "title": "Senior Backend Engineer",
                            "organization": "TechCorp",
                            "location": "",
                            "start": "2020",
                            "end": "Present",
                            "bullets": [
                                "Owned the distributed systems platform serving 5M users.",
                                "Reduced p99 latency 30% through targeted query optimization.",
                            ],
                        },
                        {
                            "title": "Backend Engineer",
                            "organization": "FinStart",
                            "location": "",
                            "start": "2017",
                            "end": "2020",
                            "bullets": [
                                "Built the payments ingestion pipeline from scratch.",
                            ],
                        },
                    ],
                    "education": [
                        {
                            "institution": "Stanford",
                            "degree": "MS",
                            "field_of_study": "Computer Science",
                            "start": "2015",
                            "end": "2017",
                        },
                        {
                            "institution": "IIT Madras",
                            "degree": "BTech",
                            "field_of_study": "CS",
                            "start": "2011",
                            "end": "2015",
                        },
                    ],
                }
            # The conversational intake returns the empty default —
            # this test is about the structuring pass at /generate
            # time, not the per-turn intake calls.
            return {
                "draft_updates": {},
                "assistant_message": "Got it.",
                "status": "collecting",
                "focus_field": "",
            }

    monkeypatch.setattr(
        "backend.routers.workspace._resolve_openai_service",
        lambda access_token, refresh_token: _StubOpenAIService(),
    )

    start_response = client.post("/api/workspace/resume-builder/start")
    session_id = start_response.json()["session_id"]
    # Push the prose into the session via the update route — we don't
    # need to drive the conversational intake for this test.
    client.post(
        "/api/workspace/resume-builder/update",
        json={
            "session_id": session_id,
            "draft_profile": {
                "full_name": "Priya Sharma",
                "target_role": "Senior Backend Engineer",
                "experience_notes": (
                    "Senior Backend Engineer at TechCorp 2020-Present, "
                    "prior at FinStart 2017-2020"
                ),
                "education_notes": (
                    "MS Computer Science Stanford 2015-2017, "
                    "BTech CS IIT Madras 2011-2015"
                ),
                "skills": ["Python", "PostgreSQL"],
                "contact_lines": ["priya@example.com"],
            },
        },
    )

    generate_response = client.post(
        "/api/workspace/resume-builder/generate",
        json={"session_id": session_id},
    )
    assert generate_response.status_code == 200
    payload = generate_response.json()

    assert "resume_builder_structuring" in captured_calls, (
        "LLM structuring task was never invoked — service not plumbed through generate"
    )

    profile = payload["candidate_profile"]
    assert len(profile["experience"]) == 2
    titles = [entry["title"] for entry in profile["experience"]]
    assert titles == ["Senior Backend Engineer", "Backend Engineer"]
    # The LLM-rewritten bullet has to make it through the rendering
    # pipeline — this is the "rich feel" the user wants from LLM output.
    descriptions = "\n".join(entry["description"] for entry in profile["experience"])
    assert "Reduced p99 latency 30%" in descriptions
    assert len(profile["education"]) == 2
    assert profile["education"][0]["institution"] == "Stanford"
    assert profile["education"][1]["institution"] == "IIT Madras"

    # Cleanup so subsequent tests don't see the fixture's session.
    _SESSIONS.pop(session_id, None)


def test_resume_builder_structuring_falls_back_to_regex_when_llm_unavailable(monkeypatch):
    """No OpenAIService → regex parser fills the candidate profile.

    Same shape of output, just less polished. This is the safety net
    that keeps the resume builder working for unauthenticated users
    and during transient LLM outages."""
    from backend.services.resume_builder_service import _SESSIONS

    monkeypatch.setattr(
        "backend.routers.workspace._resolve_openai_service",
        lambda access_token, refresh_token: None,
    )

    start_response = client.post("/api/workspace/resume-builder/start")
    session_id = start_response.json()["session_id"]
    client.post(
        "/api/workspace/resume-builder/update",
        json={
            "session_id": session_id,
            "draft_profile": {
                "full_name": "Priya Sharma",
                "experience_notes": (
                    "Senior Backend Engineer at TechCorp 2020-Present, "
                    "prior at FinStart 2017-2020"
                ),
                "education_notes": (
                    "MS Computer Science Stanford 2017, BTech CS IIT Madras 2015"
                ),
                "skills": ["Python"],
                "contact_lines": ["priya@example.com"],
            },
        },
    )

    generate_response = client.post(
        "/api/workspace/resume-builder/generate",
        json={"session_id": session_id},
    )
    assert generate_response.status_code == 200
    profile = generate_response.json()["candidate_profile"]
    # Regex still produces the multi-role split (covered by the parser
    # tests above) — both companies and both schools are present.
    orgs = [entry["organization"] for entry in profile["experience"]]
    assert "TechCorp" in orgs
    assert "FinStart" in orgs
    schools = [entry["institution"] for entry in profile["education"]]
    assert "Stanford" in schools
    assert "IIT Madras" in schools

    _SESSIONS.pop(session_id, None)


def test_resume_builder_structuring_falls_back_when_llm_raises(monkeypatch):
    """Transient LLM error during structuring → regex fallback kicks in.

    Pins the contract that the export / generate endpoints never bubble
    up an LLM error to the user — there's always a deterministic
    fallback that produces a usable resume."""
    from backend.services.resume_builder_service import _SESSIONS

    class _RaisingOpenAIService:
        def is_available(self):
            return True

        def run_json_prompt(self, system, user, **kwargs):
            if kwargs.get("task_name") == "resume_builder_structuring":
                raise RuntimeError("OpenAI 503 — model overloaded")
            return {
                "draft_updates": {},
                "assistant_message": "ok",
                "status": "collecting",
                "focus_field": "",
            }

    monkeypatch.setattr(
        "backend.routers.workspace._resolve_openai_service",
        lambda access_token, refresh_token: _RaisingOpenAIService(),
    )

    start_response = client.post("/api/workspace/resume-builder/start")
    session_id = start_response.json()["session_id"]
    client.post(
        "/api/workspace/resume-builder/update",
        json={
            "session_id": session_id,
            "draft_profile": {
                "full_name": "Priya Sharma",
                "experience_notes": (
                    "Senior Backend Engineer at TechCorp 2020-Present, "
                    "prior at FinStart 2017-2020"
                ),
                "education_notes": "BTech CS IIT Madras 2015",
                "skills": ["Python"],
                "contact_lines": ["priya@example.com"],
            },
        },
    )

    generate_response = client.post(
        "/api/workspace/resume-builder/generate",
        json={"session_id": session_id},
    )
    assert generate_response.status_code == 200, (
        "LLM exception must NOT bubble up — fallback should produce a 200"
    )
    profile = generate_response.json()["candidate_profile"]
    orgs = [entry["organization"] for entry in profile["experience"]]
    assert "TechCorp" in orgs and "FinStart" in orgs

    _SESSIONS.pop(session_id, None)


def test_resume_builder_structuring_falls_back_per_section_on_partial_llm_output(monkeypatch):
    """LLM returned a usable experience array but an empty education
    array — we keep the LLM experience, fall back to regex for
    education. This is the layered defense: every section gets the
    best signal we can produce, independently."""
    from backend.services.resume_builder_service import _SESSIONS

    class _PartialOpenAIService:
        def is_available(self):
            return True

        def run_json_prompt(self, system, user, **kwargs):
            if kwargs.get("task_name") == "resume_builder_structuring":
                return {
                    "experience": [
                        {
                            "title": "Senior Backend Engineer",
                            "organization": "TechCorp",
                            "start": "2020",
                            "end": "Present",
                            "bullets": ["Built distributed systems."],
                        },
                    ],
                    # Education came back empty — must fall back to regex.
                    "education": [],
                }
            return {
                "draft_updates": {},
                "assistant_message": "ok",
                "status": "collecting",
                "focus_field": "",
            }

    monkeypatch.setattr(
        "backend.routers.workspace._resolve_openai_service",
        lambda access_token, refresh_token: _PartialOpenAIService(),
    )

    start_response = client.post("/api/workspace/resume-builder/start")
    session_id = start_response.json()["session_id"]
    client.post(
        "/api/workspace/resume-builder/update",
        json={
            "session_id": session_id,
            "draft_profile": {
                "full_name": "Priya Sharma",
                "experience_notes": "Senior Backend Engineer at TechCorp 2020-Present",
                "education_notes": "BTech CS IIT Madras 2015",
                "skills": ["Python"],
                "contact_lines": ["priya@example.com"],
            },
        },
    )

    generate_response = client.post(
        "/api/workspace/resume-builder/generate",
        json={"session_id": session_id},
    )
    assert generate_response.status_code == 200
    profile = generate_response.json()["candidate_profile"]
    # Experience came from the LLM — bullet text is the giveaway.
    descriptions = "\n".join(entry["description"] for entry in profile["experience"])
    assert "Built distributed systems" in descriptions
    # Education came from regex — no LLM rewrite, but the school is there.
    schools = [entry["institution"] for entry in profile["education"]]
    assert "IIT Madras" in schools

    _SESSIONS.pop(session_id, None)


def test_resume_builder_structuring_caches_payload_across_exports(monkeypatch):
    """Cache hit on re-export: a second /export against the SAME draft
    must reuse the cached structured payload and NOT re-call the LLM.

    Without this, switching theme or downloading the same artifact in
    a different format would re-burn a structuring call AND the LLM
    would subtly rephrase bullets (real models aren't deterministic
    across calls), so identical-content downloads showed inconsistent
    wording. The fix: hash the prompt's inputs and skip the call when
    the hash matches what we cached on the session."""
    from backend.services.resume_builder_service import _SESSIONS

    structuring_call_count = {"value": 0}

    class _CountingOpenAIService:
        def is_available(self):
            return True

        def run_json_prompt(self, system, user, **kwargs):
            if kwargs.get("task_name") == "resume_builder_structuring":
                structuring_call_count["value"] += 1
                return {
                    "experience": [
                        {
                            "title": "Senior Backend Engineer",
                            "organization": "TechCorp",
                            "start": "2020",
                            "end": "Present",
                            "bullets": [
                                "Owned distributed systems platform.",
                                "Reduced p99 latency 30%.",
                            ],
                        },
                    ],
                    "education": [
                        {
                            "institution": "Stanford",
                            "degree": "MS",
                            "field_of_study": "Computer Science",
                            "start": "2015",
                            "end": "2017",
                        },
                    ],
                }
            return {
                "draft_updates": {},
                "assistant_message": "ok",
                "status": "collecting",
                "focus_field": "",
            }

    monkeypatch.setattr(
        "backend.routers.workspace._resolve_openai_service",
        lambda access_token, refresh_token: _CountingOpenAIService(),
    )

    start_response = client.post("/api/workspace/resume-builder/start")
    session_id = start_response.json()["session_id"]
    client.post(
        "/api/workspace/resume-builder/update",
        json={
            "session_id": session_id,
            "draft_profile": {
                "full_name": "Priya Sharma",
                "experience_notes": "Senior Backend Engineer at TechCorp 2020-Present",
                "education_notes": "MS Computer Science Stanford 2015-2017",
                "skills": ["Python"],
                "contact_lines": ["priya@example.com"],
            },
        },
    )

    # First export — LLM fires once, cache populated.
    first = client.post(
        "/api/workspace/resume-builder/export",
        json={
            "session_id": session_id,
            "export_format": "docx",
            "theme": "classic_ats",
        },
    )
    assert first.status_code == 200
    assert structuring_call_count["value"] == 1

    # Second export at a different theme — cache hit, no new LLM call.
    second = client.post(
        "/api/workspace/resume-builder/export",
        json={
            "session_id": session_id,
            "export_format": "docx",
            "theme": "professional_neutral",
        },
    )
    assert second.status_code == 200
    assert structuring_call_count["value"] == 1, (
        "Re-export at a different theme must reuse cached structuring payload"
    )

    # Third export at a different format — still cache hit.
    third = client.post(
        "/api/workspace/resume-builder/export",
        json={
            "session_id": session_id,
            "export_format": "pdf",
            "theme": "classic_ats",
        },
    )
    assert third.status_code == 200
    assert structuring_call_count["value"] == 1, (
        "PDF re-export must reuse cached payload from the DOCX export"
    )

    _SESSIONS.pop(session_id, None)


def test_resume_builder_structuring_cache_invalidates_when_notes_change(monkeypatch):
    """Cache miss when the user edits experience_notes. The signature
    hash detects the change and forces a fresh LLM call so the
    rendered output reflects the new prose."""
    from backend.services.resume_builder_service import _SESSIONS

    structuring_calls: list[str] = []

    class _NotesAwareOpenAIService:
        def is_available(self):
            return True

        def run_json_prompt(self, system, user, **kwargs):
            if kwargs.get("task_name") == "resume_builder_structuring":
                structuring_calls.append(user)
                # Echo the title from the prose so the test can prove
                # the second call ran on the NEW notes, not the stale
                # cached payload.
                title = "Frontend Engineer" if "Frontend" in user else "Senior Backend Engineer"
                org = "Acme" if "Frontend" in user else "TechCorp"
                return {
                    "experience": [
                        {
                            "title": title,
                            "organization": org,
                            "start": "2020",
                            "end": "Present",
                            "bullets": ["Did things."],
                        },
                    ],
                    "education": [],
                }
            return {
                "draft_updates": {},
                "assistant_message": "ok",
                "status": "collecting",
                "focus_field": "",
            }

    monkeypatch.setattr(
        "backend.routers.workspace._resolve_openai_service",
        lambda access_token, refresh_token: _NotesAwareOpenAIService(),
    )

    start_response = client.post("/api/workspace/resume-builder/start")
    session_id = start_response.json()["session_id"]

    # First generate with backend prose.
    client.post(
        "/api/workspace/resume-builder/update",
        json={
            "session_id": session_id,
            "draft_profile": {
                "full_name": "Priya",
                "experience_notes": "Senior Backend Engineer at TechCorp 2020-Present",
                "skills": ["Python"],
                "contact_lines": ["priya@example.com"],
            },
        },
    )
    first = client.post(
        "/api/workspace/resume-builder/generate",
        json={"session_id": session_id},
    )
    assert first.status_code == 200
    first_orgs = [e["organization"] for e in first.json()["candidate_profile"]["experience"]]
    assert first_orgs == ["TechCorp"]
    assert len(structuring_calls) == 1

    # User edits to a different role — cache must invalidate.
    client.post(
        "/api/workspace/resume-builder/update",
        json={
            "session_id": session_id,
            "draft_profile": {
                "experience_notes": "Frontend Engineer at Acme 2020-Present",
            },
        },
    )
    second = client.post(
        "/api/workspace/resume-builder/generate",
        json={"session_id": session_id},
    )
    assert second.status_code == 200
    second_orgs = [e["organization"] for e in second.json()["candidate_profile"]["experience"]]
    assert second_orgs == ["Acme"], (
        "Edited prose must invalidate the cache and re-run the LLM"
    )
    assert len(structuring_calls) == 2

    _SESSIONS.pop(session_id, None)


def test_resume_builder_structuring_cache_survives_session_persistence_round_trip(monkeypatch):
    """The cache fields (signature + structured payloads) get persisted
    in the session payload and round-trip through Supabase. After a
    container restart hydrates the session from Supabase the cache
    still hits, no re-call to the LLM needed."""
    from backend.services.resume_builder_service import (
        _SESSIONS,
        export_resume_builder_session_payload,
        restore_resume_builder_session_payload,
    )

    call_counter = {"value": 0}

    class _CountingOpenAIService:
        def is_available(self):
            return True

        def run_json_prompt(self, system, user, **kwargs):
            if kwargs.get("task_name") == "resume_builder_structuring":
                call_counter["value"] += 1
                return {
                    "experience": [
                        {
                            "title": "Senior Backend Engineer",
                            "organization": "TechCorp",
                            "start": "2020",
                            "end": "Present",
                            "bullets": ["Built things."],
                        },
                    ],
                    "education": [],
                }
            return {
                "draft_updates": {},
                "assistant_message": "ok",
                "status": "collecting",
                "focus_field": "",
            }

    monkeypatch.setattr(
        "backend.routers.workspace._resolve_openai_service",
        lambda access_token, refresh_token: _CountingOpenAIService(),
    )

    start_response = client.post("/api/workspace/resume-builder/start")
    session_id = start_response.json()["session_id"]
    client.post(
        "/api/workspace/resume-builder/update",
        json={
            "session_id": session_id,
            "draft_profile": {
                "full_name": "Priya",
                "experience_notes": "Senior Backend Engineer at TechCorp 2020-Present",
                "skills": ["Python"],
                "contact_lines": ["priya@example.com"],
            },
        },
    )

    # Populate the cache.
    client.post(
        "/api/workspace/resume-builder/generate",
        json={"session_id": session_id},
    )
    assert call_counter["value"] == 1

    # Round-trip through the persistence boundary.
    payload = export_resume_builder_session_payload(session_id=session_id)
    _SESSIONS.pop(session_id, None)
    restore_resume_builder_session_payload(payload)

    # Generate again — cache must still hit, no new LLM call.
    client.post(
        "/api/workspace/resume-builder/generate",
        json={"session_id": session_id},
    )
    assert call_counter["value"] == 1, (
        "Cache must survive persistence round-trip — re-call after restore is wasted"
    )

    _SESSIONS.pop(session_id, None)


def test_resume_builder_project_parser_splits_blocks():
    """Regex fallback for projects splits blank-line-separated blocks
    into one ProjectEntry per project, preserving the headline + bullet
    structure and lifting any URL out of the headline into `link`."""
    from backend.services.resume_builder_service import _build_project_entries

    notes = (
        "Grounded RAG Q&A System github.com/LEANDERANTONY/HelpmateAI_RAG_QA_System\n"
        "- Designed production-grade RAG pipeline with hybrid retrieval.\n"
        "- Improved supported answer rate from 0.803 to 0.882 on RAGAS.\n"
        "\n"
        "GitHub Portfolio Reviewer Agent portfolio-reviewer-agent.streamlit.app\n"
        "- Combines deterministic repo analysis with LLM scoring.\n"
        "- Caches LLM calls to reduce cost and latency."
    )
    entries = _build_project_entries(notes)
    assert len(entries) == 2
    first, second = entries
    assert "Grounded RAG" in first.name
    assert "github.com" in first.link
    assert any("RAG pipeline" in b for b in first.bullets)
    assert "Portfolio Reviewer" in second.name
    assert "streamlit.app" in second.link


def test_resume_builder_structuring_emits_projects_via_llm(monkeypatch):
    """LLM structuring path produces ProjectEntry objects from
    projects_notes prose. Pins that the new contract surface — projects
    in the structuring response — flows through to the artifact."""
    from backend.services.resume_builder_service import _SESSIONS

    class _ProjectsAwareService:
        def is_available(self):
            return True

        def run_json_prompt(self, system, user, **kwargs):
            if kwargs.get("task_name") == "resume_builder_structuring":
                return {
                    "experience": [],
                    "education": [],
                    "projects": [
                        {
                            "name": "RAG Q&A System",
                            "description": "",
                            "bullets": [
                                "Hybrid retrieval with cross-encoder reranking.",
                                "Outperformed OpenAI File Search on faithfulness.",
                            ],
                            "technologies": ["FastAPI", "ChromaDB"],
                            "start": "2024",
                            "end": "Present",
                            "link": "github.com/LEANDERANTONY/HelpmateAI_RAG_QA_System",
                        },
                    ],
                }
            return {
                "draft_updates": {},
                "assistant_message": "ok",
                "status": "collecting",
                "focus_field": "",
            }

    monkeypatch.setattr(
        "backend.routers.workspace._resolve_openai_service",
        lambda access_token, refresh_token: _ProjectsAwareService(),
    )

    start_response = client.post("/api/workspace/resume-builder/start")
    session_id = start_response.json()["session_id"]
    client.post(
        "/api/workspace/resume-builder/update",
        json={
            "session_id": session_id,
            "draft_profile": {
                "full_name": "Priya Sharma",
                "skills": ["Python"],
                "contact_lines": ["priya@example.com"],
                "projects_notes": "RAG Q&A System github.com/LEANDERANTONY/HelpmateAI_RAG_QA_System - hybrid retrieval, beats OpenAI File Search on faithfulness",
            },
        },
    )

    generate_response = client.post(
        "/api/workspace/resume-builder/generate",
        json={"session_id": session_id},
    )
    assert generate_response.status_code == 200
    profile = generate_response.json()["candidate_profile"]
    assert len(profile["projects"]) == 1
    project = profile["projects"][0]
    assert project["name"] == "RAG Q&A System"
    assert "github.com" in project["link"]
    assert any("Hybrid retrieval" in b for b in project["bullets"])

    _SESSIONS.pop(session_id, None)


def test_resume_builder_publications_round_trip_into_artifact(monkeypatch):
    """Publications are a list-of-strings field on the draft (like
    certifications). They flow straight through to the artifact's
    publication_entries — no LLM structuring needed for citations."""
    from backend.services.resume_builder_service import _SESSIONS

    monkeypatch.setattr(
        "backend.routers.workspace._resolve_openai_service",
        lambda access_token, refresh_token: None,  # regex path is fine for this test
    )

    start_response = client.post("/api/workspace/resume-builder/start")
    session_id = start_response.json()["session_id"]
    client.post(
        "/api/workspace/resume-builder/update",
        json={
            "session_id": session_id,
            "draft_profile": {
                "full_name": "Priya Sharma",
                "experience_notes": "ML Engineer at Acme 2020-Present",
                "education_notes": "PhD AI Stanford 2020",
                "skills": ["Python"],
                "contact_lines": ["priya@example.com"],
                "publications": [
                    "Sharma, P. et al. — \"Multimodal Cancer Detection\" — Nature ML (2024).",
                    "Sharma, P. — \"RAG Evaluation Framework\" — NeurIPS (2023).",
                ],
            },
        },
    )

    generate_response = client.post(
        "/api/workspace/resume-builder/generate",
        json={"session_id": session_id},
    )
    assert generate_response.status_code == 200
    profile = generate_response.json()["candidate_profile"]
    assert len(profile["publications"]) == 2
    joined = " ".join(profile["publications"])
    assert "Multimodal Cancer Detection" in joined
    assert "NeurIPS" in joined

    _SESSIONS.pop(session_id, None)


def test_resume_builder_export_renders_projects_in_docx(monkeypatch):
    """End-to-end: a session with projects_notes flows through to the
    rendered DOCX bytes. PROJECTS section appears with the project
    names + bullets."""
    import base64
    from io import BytesIO

    from docx import Document
    from backend.services.resume_builder_service import _SESSIONS

    class _ProjectsService:
        def is_available(self):
            return True

        def run_json_prompt(self, system, user, **kwargs):
            if kwargs.get("task_name") == "resume_builder_structuring":
                return {
                    "experience": [
                        {
                            "title": "AI Engineer",
                            "organization": "Self-Directed",
                            "start": "2022",
                            "end": "Present",
                            "bullets": ["Built RAG and agentic systems."],
                        },
                    ],
                    "education": [
                        {"institution": "Stanford", "degree": "MS", "field_of_study": "CS", "start": "2017"},
                    ],
                    "projects": [
                        {
                            "name": "Grounded RAG Q&A System",
                            "bullets": ["Outperformed OpenAI File Search."],
                            "technologies": ["FastAPI"],
                            "link": "github.com/me/rag",
                        },
                    ],
                }
            return {"draft_updates": {}, "assistant_message": "ok", "status": "collecting", "focus_field": ""}

    monkeypatch.setattr(
        "backend.routers.workspace._resolve_openai_service",
        lambda access_token, refresh_token: _ProjectsService(),
    )

    start_response = client.post("/api/workspace/resume-builder/start")
    session_id = start_response.json()["session_id"]
    client.post(
        "/api/workspace/resume-builder/update",
        json={
            "session_id": session_id,
            "draft_profile": {
                "full_name": "Priya",
                "experience_notes": "AI Engineer Self-Directed 2022-Present, built RAG systems",
                "education_notes": "MS CS Stanford 2017",
                "skills": ["Python", "FastAPI"],
                "contact_lines": ["priya@example.com"],
                "projects_notes": "Grounded RAG Q&A System github.com/me/rag - outperforms OpenAI File Search",
                "publications": ["Sharma, P. — \"RAG Evaluation\" — Conf 2024."],
            },
        },
    )
    client.post("/api/workspace/resume-builder/generate", json={"session_id": session_id})

    export_response = client.post(
        "/api/workspace/resume-builder/export",
        json={"session_id": session_id, "export_format": "docx", "theme": "classic_ats"},
    )
    assert export_response.status_code == 200
    raw_bytes = base64.b64decode(export_response.json()["content_base64"])
    doc = Document(BytesIO(raw_bytes))
    body = "\n".join(p.text for p in doc.paragraphs)

    # Projects section landed.
    assert "Grounded RAG Q&A System" in body
    assert "OpenAI File Search" in body
    # Publications section landed.
    assert "RAG Evaluation" in body

    _SESSIONS.pop(session_id, None)


def test_resume_builder_structuring_expands_thin_summary(monkeypatch):
    """When the user types a one-liner summary, the structuring pass
    can replace it with a polished 2-3 sentence version (third-person
    ATS voice, grounded in facts already in the draft)."""
    import base64
    from io import BytesIO

    from docx import Document
    from backend.services.resume_builder_service import _SESSIONS

    class _ExpandsService:
        def is_available(self):
            return True

        def run_json_prompt(self, system, user, **kwargs):
            if kwargs.get("task_name") == "resume_builder_structuring":
                return {
                    "experience": [
                        {
                            "title": "Senior Backend Engineer",
                            "organization": "Stripe",
                            "start": "2022",
                            "end": "Present",
                            "bullets": ["Led the rate-limiter rewrite."],
                        },
                    ],
                    "education": [],
                    "projects": [],
                    "professional_summary": (
                        "Senior backend engineer with deep payments-platform "
                        "experience. Specializes in low-latency rate-limiting and "
                        "high-throughput data migrations on Stripe's billing "
                        "infrastructure."
                    ),
                }
            return {"draft_updates": {}, "assistant_message": "ok", "status": "collecting", "focus_field": ""}

    monkeypatch.setattr(
        "backend.routers.workspace._resolve_openai_service",
        lambda access_token, refresh_token: _ExpandsService(),
    )

    start_response = client.post("/api/workspace/resume-builder/start")
    session_id = start_response.json()["session_id"]
    client.post(
        "/api/workspace/resume-builder/update",
        json={
            "session_id": session_id,
            "draft_profile": {
                "full_name": "Priya",
                # User typed a one-liner.
                "professional_summary": "I have 6 years building payment systems.",
                "experience_notes": "Senior Backend Engineer at Stripe 2022-Present, led rate-limiter rewrite",
                "skills": ["Python"],
                "contact_lines": ["priya@example.com"],
            },
        },
    )

    export_response = client.post(
        "/api/workspace/resume-builder/export",
        json={"session_id": session_id, "export_format": "docx", "theme": "classic_ats"},
    )
    assert export_response.status_code == 200
    raw_bytes = base64.b64decode(export_response.json()["content_base64"])
    doc = Document(BytesIO(raw_bytes))
    body = "\n".join(p.text for p in doc.paragraphs)

    # The expanded summary replaces the one-liner.
    assert "Specializes in low-latency rate-limiting" in body
    assert "I have 6 years" not in body  # original verbatim is gone

    _SESSIONS.pop(session_id, None)


def test_resume_builder_structuring_keeps_user_summary_when_already_long(monkeypatch):
    """If the LLM returns an empty professional_summary or the user's
    own summary is already substantial, we keep the user's verbatim
    text in the rendered resume."""
    import base64
    from io import BytesIO

    from docx import Document
    from backend.services.resume_builder_service import _SESSIONS

    class _NoExpansionService:
        def is_available(self):
            return True

        def run_json_prompt(self, system, user, **kwargs):
            if kwargs.get("task_name") == "resume_builder_structuring":
                return {
                    "experience": [{"title": "Engineer", "organization": "Acme", "bullets": []}],
                    "education": [],
                    "projects": [],
                    "professional_summary": "",  # LLM declines to expand
                }
            return {"draft_updates": {}, "assistant_message": "ok", "status": "collecting", "focus_field": ""}

    monkeypatch.setattr(
        "backend.routers.workspace._resolve_openai_service",
        lambda access_token, refresh_token: _NoExpansionService(),
    )

    user_summary = (
        "Senior backend engineer with 6+ years building distributed payment "
        "systems. Led the rate-limiter rewrite at Stripe and shipped Vitess "
        "migrations across 3.2B legacy charges."
    )
    start_response = client.post("/api/workspace/resume-builder/start")
    session_id = start_response.json()["session_id"]
    client.post(
        "/api/workspace/resume-builder/update",
        json={
            "session_id": session_id,
            "draft_profile": {
                "full_name": "Priya",
                "professional_summary": user_summary,
                "experience_notes": "Engineer at Acme",
                "skills": ["Python"],
                "contact_lines": ["priya@example.com"],
            },
        },
    )

    export_response = client.post(
        "/api/workspace/resume-builder/export",
        json={"session_id": session_id, "export_format": "docx", "theme": "classic_ats"},
    )
    assert export_response.status_code == 200
    raw_bytes = base64.b64decode(export_response.json()["content_base64"])
    doc = Document(BytesIO(raw_bytes))
    body = "\n".join(p.text for p in doc.paragraphs)

    # User's own summary survived intact.
    assert "Vitess migrations across 3.2B" in body

    _SESSIONS.pop(session_id, None)


def test_resume_builder_skill_categories_round_trip_into_artifact(monkeypatch):
    """LLM emits skill_categories in the structuring payload; the
    sanitizer keeps only categories whose skills appear in the user's
    flat list; the artifact's skill_categories field gets populated;
    the rendered DOCX shows category-labelled rows instead of a single
    pipe-separated line."""
    import base64
    from io import BytesIO

    from docx import Document
    from backend.services.resume_builder_service import _SESSIONS

    class _CategoriesService:
        def is_available(self):
            return True

        def run_json_prompt(self, system, user, **kwargs):
            if kwargs.get("task_name") == "resume_builder_structuring":
                return {
                    "experience": [
                        {
                            "title": "AI Engineer",
                            "organization": "Self-Directed",
                            "start": "2022",
                            "end": "Present",
                            "bullets": ["Built RAG and agentic systems."],
                        },
                    ],
                    "education": [
                        {"institution": "Stanford", "degree": "MS", "field_of_study": "CS"},
                    ],
                    "projects": [],
                    "skill_categories": {
                        "Languages & Tools": ["Python", "SQL"],
                        "ML / DL Frameworks": ["PyTorch", "Scikit-learn"],
                        "GenAI & LLMs": ["LangChain", "OpenAI API"],
                        # Sanitizer must drop "Rust" — the user never typed it.
                        "Bogus": ["Rust"],
                    },
                }
            return {"draft_updates": {}, "assistant_message": "ok", "status": "collecting", "focus_field": ""}

    monkeypatch.setattr(
        "backend.routers.workspace._resolve_openai_service",
        lambda access_token, refresh_token: _CategoriesService(),
    )

    start_response = client.post("/api/workspace/resume-builder/start")
    session_id = start_response.json()["session_id"]
    client.post(
        "/api/workspace/resume-builder/update",
        json={
            "session_id": session_id,
            "draft_profile": {
                "full_name": "Priya",
                "experience_notes": "AI Engineer 2022-Present",
                "education_notes": "MS CS Stanford",
                "skills": [
                    "Python", "SQL", "PyTorch", "Scikit-learn",
                    "LangChain", "OpenAI API", "Docker", "FastAPI",
                ],
                "contact_lines": ["priya@example.com"],
            },
        },
    )

    export_response = client.post(
        "/api/workspace/resume-builder/export",
        json={"session_id": session_id, "export_format": "docx", "theme": "classic_ats"},
    )
    assert export_response.status_code == 200
    raw_bytes = base64.b64decode(export_response.json()["content_base64"])
    doc = Document(BytesIO(raw_bytes))
    body = "\n".join(p.text for p in doc.paragraphs)

    # Category labels rendered (not just a flat pipe list).
    assert "Languages & Tools:" in body
    assert "ML / DL Frameworks:" in body
    assert "GenAI & LLMs:" in body
    # The bogus 'Rust' category contained a skill the user didn't type
    # — sanitizer drops it before it can confuse the reader.
    assert "Rust" not in body
    assert "Bogus" not in body

    _SESSIONS.pop(session_id, None)


def test_resume_builder_skill_categories_falls_back_to_flat_when_llm_omits(monkeypatch):
    """When the LLM doesn't emit categories (or returns {}), the
    renderer falls back to the existing flat pipe-separated layout.
    Pins backwards compatibility for sparse skill sets."""
    import base64
    from io import BytesIO

    from docx import Document
    from backend.services.resume_builder_service import _SESSIONS

    class _NoCategoriesService:
        def is_available(self):
            return True

        def run_json_prompt(self, system, user, **kwargs):
            if kwargs.get("task_name") == "resume_builder_structuring":
                return {
                    "experience": [
                        {"title": "Engineer", "organization": "Acme", "bullets": []},
                    ],
                    "education": [],
                    "projects": [],
                    # No skill_categories key — sparse skill set, flat list is fine.
                }
            return {"draft_updates": {}, "assistant_message": "ok", "status": "collecting", "focus_field": ""}

    monkeypatch.setattr(
        "backend.routers.workspace._resolve_openai_service",
        lambda access_token, refresh_token: _NoCategoriesService(),
    )

    start_response = client.post("/api/workspace/resume-builder/start")
    session_id = start_response.json()["session_id"]
    client.post(
        "/api/workspace/resume-builder/update",
        json={
            "session_id": session_id,
            "draft_profile": {
                "full_name": "Priya",
                "experience_notes": "Engineer at Acme",
                "skills": ["Python", "SQL"],
                "contact_lines": ["priya@example.com"],
            },
        },
    )

    export_response = client.post(
        "/api/workspace/resume-builder/export",
        json={"session_id": session_id, "export_format": "docx", "theme": "classic_ats"},
    )
    assert export_response.status_code == 200
    raw_bytes = base64.b64decode(export_response.json()["content_base64"])
    doc = Document(BytesIO(raw_bytes))
    body = "\n".join(p.text for p in doc.paragraphs)

    # Skills present, but as a flat pipe list (no category labels).
    assert "Python" in body
    assert "SQL" in body
    assert "|" in body  # flat pipe separator visible
    assert ":" not in body or "Languages & Tools:" not in body  # no category labels

    _SESSIONS.pop(session_id, None)


def test_resume_builder_intake_recovers_full_name_when_llm_truncates(monkeypatch):
    """Safety net for the LLM dropping a surname.

    Real-world scenario from QA: the user typed
    'Priya Sharma, Bangalore. priya@gmail.com, +91 8000000000' but the
    LLM only captured full_name='Priya'. After applying the LLM
    updates, _augment_full_name_from_message looks at the literal
    message and promotes the longer name when it (a) starts with what
    the LLM captured and (b) is a valid name shape."""
    from backend.services.resume_builder_service import _SESSIONS

    class _TruncatingOpenAIService:
        def is_available(self):
            return True

        def run_json_prompt(self, system, user, **kwargs):
            return {
                "draft_updates": {
                    # LLM dropped 'Sharma' even though the user typed it.
                    "full_name": "Priya",
                    "location": "Bangalore",
                    "contact_lines": ["priya@gmail.com", "+91 8000000000"],
                },
                "assistant_message": "Got it Priya. What role are you targeting?",
                "status": "collecting",
                "focus_field": "target_role",
            }

    monkeypatch.setattr(
        "backend.routers.workspace._resolve_openai_service",
        lambda access_token, refresh_token: _TruncatingOpenAIService(),
    )

    start_response = client.post("/api/workspace/resume-builder/start")
    session_id = start_response.json()["session_id"]
    response = client.post(
        "/api/workspace/resume-builder/message",
        json={
            "session_id": session_id,
            "message": "Priya Sharma, Bangalore. priya@gmail.com, +91 8000000000",
            "input_mode": "text",
        },
    )
    assert response.status_code == 200
    draft = response.json()["draft_profile"]
    # Safety net should have promoted 'Priya' → 'Priya Sharma' from the
    # literal message.
    assert draft["full_name"] == "Priya Sharma", (
        "Name safety net failed to recover surname from literal message"
    )

    _SESSIONS.pop(session_id, None)


def test_resume_builder_intake_keeps_llm_name_when_user_typed_only_first_name(monkeypatch):
    """Safety net only kicks in when the literal message has MORE
    name. If the user only typed 'Priya', the LLM correctly captures
    'Priya' and the safety net is a no-op (no surname to promote).
    Pins the negative case."""
    from backend.services.resume_builder_service import _SESSIONS

    class _StubService:
        def is_available(self):
            return True

        def run_json_prompt(self, system, user, **kwargs):
            return {
                "draft_updates": {"full_name": "Priya"},
                "assistant_message": "Got it Priya.",
                "status": "collecting",
                "focus_field": "location",
            }

    monkeypatch.setattr(
        "backend.routers.workspace._resolve_openai_service",
        lambda access_token, refresh_token: _StubService(),
    )

    start_response = client.post("/api/workspace/resume-builder/start")
    session_id = start_response.json()["session_id"]
    response = client.post(
        "/api/workspace/resume-builder/message",
        json={
            "session_id": session_id,
            "message": "Priya",
            "input_mode": "text",
        },
    )
    assert response.status_code == 200
    assert response.json()["draft_profile"]["full_name"] == "Priya"

    _SESSIONS.pop(session_id, None)


def test_resume_builder_intake_does_not_overwrite_llm_correction(monkeypatch):
    """Safety net only fires on PREFIX matches. If the LLM disagrees
    with the literal first chunk (e.g., extracted 'Maya' from 'Priya
    typed it wrong, my name is Maya Sharma'), don't overwrite the
    LLM's correction with 'Priya' from the literal first chunk."""
    from backend.services.resume_builder_service import _SESSIONS

    class _CorrectingService:
        def is_available(self):
            return True

        def run_json_prompt(self, system, user, **kwargs):
            return {
                # LLM extracted the corrected name, not the literal first chunk.
                "draft_updates": {"full_name": "Maya Sharma"},
                "assistant_message": "Got it Maya.",
                "status": "collecting",
                "focus_field": "location",
            }

    monkeypatch.setattr(
        "backend.routers.workspace._resolve_openai_service",
        lambda access_token, refresh_token: _CorrectingService(),
    )

    start_response = client.post("/api/workspace/resume-builder/start")
    session_id = start_response.json()["session_id"]
    response = client.post(
        "/api/workspace/resume-builder/message",
        json={
            "session_id": session_id,
            # Literal message starts with "Priya" — but the LLM
            # correctly read further and pulled "Maya Sharma". Safety
            # net must NOT overwrite Maya with Priya.
            "message": "Priya — actually no, my name is Maya Sharma",
            "input_mode": "text",
        },
    )
    assert response.status_code == 200
    assert response.json()["draft_profile"]["full_name"] == "Maya Sharma", (
        "Safety net incorrectly overrode the LLM's correction"
    )

    _SESSIONS.pop(session_id, None)


def test_resume_builder_message_uses_llm_when_openai_service_available(monkeypatch):
    """LLM-first intake: when an OpenAIService is plumbed in, the
    model picks the next question and merges partial draft updates,
    instead of marching through the hardcoded 5 steps."""
    from backend.services.resume_builder_service import _SESSIONS

    captured_prompts: list[dict] = []

    class _StubOpenAIService:
        def is_available(self):
            return True

        def run_json_prompt(self, system, user, **kwargs):
            captured_prompts.append({"system": system, "user": user})
            return {
                "draft_updates": {
                    "full_name": "Leander Antony",
                    "location": "Chennai, India",
                    "contact_lines": ["leander@example.com", "+91 9999999999"],
                },
                "assistant_message": (
                    "Got it — Leander Antony in Chennai. What role are you "
                    "targeting?"
                ),
                "status": "collecting",
                "focus_field": "target_role",
            }

    monkeypatch.setattr(
        "backend.routers.workspace._resolve_openai_service",
        lambda access_token, refresh_token: _StubOpenAIService(),
    )

    start = client.post("/api/workspace/resume-builder/start")
    session_id = start.json()["session_id"]

    response = client.post(
        "/api/workspace/resume-builder/message",
        json={
            "session_id": session_id,
            "message": "I'm Leander Antony, based in Chennai, India. Reach me at leander@example.com or +91 9999999999.",
            "input_mode": "text",
        },
        headers={
            "X-Auth-Access-Token": "access-token",
            "X-Auth-Refresh-Token": "refresh-token",
        },
    )

    assert response.status_code == 200
    payload = response.json()
    assert payload["draft_profile"]["full_name"] == "Leander Antony"
    assert payload["draft_profile"]["location"] == "Chennai, India"
    assert "leander@example.com" in payload["draft_profile"]["contact_lines"]
    # Assistant message comes from the LLM, not the hardcoded step
    # acknowledgements.
    assert "Leander Antony" in payload["assistant_message"]
    # Only one LLM call; no regex step machine fallback ran.
    assert len(captured_prompts) == 1
    # The prompt body contains the user's literal message + the
    # current draft state so the model can ground its updates.
    assert "Leander Antony" in captured_prompts[0]["user"]

    # The session's conversation_history should now have one
    # user/assistant pair so subsequent turns have continuity.
    session = _SESSIONS[session_id]
    assert len(session.conversation_history) == 2
    assert session.conversation_history[0]["role"] == "user"
    assert session.conversation_history[1]["role"] == "assistant"


def test_resume_builder_message_falls_back_to_regex_when_llm_errors(monkeypatch):
    """When the LLM raises (no key, malformed JSON, network error, etc.),
    the deterministic regex step machine still runs so the feature
    keeps working."""
    from backend.services.resume_builder_service import _SESSIONS
    from src.errors import AgentExecutionError

    class _BrokenOpenAIService:
        def is_available(self):
            return True

        def run_json_prompt(self, system, user, **kwargs):
            raise AgentExecutionError("Simulated LLM failure for test.")

    monkeypatch.setattr(
        "backend.routers.workspace._resolve_openai_service",
        lambda access_token, refresh_token: _BrokenOpenAIService(),
    )

    start = client.post("/api/workspace/resume-builder/start")
    session_id = start.json()["session_id"]

    response = client.post(
        "/api/workspace/resume-builder/message",
        json={
            "session_id": session_id,
            "message": "Leander Antony, Chennai, leander@example.com",
            "input_mode": "text",
        },
        headers={
            "X-Auth-Access-Token": "access-token",
            "X-Auth-Refresh-Token": "refresh-token",
        },
    )

    assert response.status_code == 200
    payload = response.json()
    # Regex fallback fired: full_name + contact captured by
    # _apply_basics, current_step advanced to "role".
    assert payload["draft_profile"]["full_name"] == "Leander Antony"
    assert payload["current_step"] == "role"
    # Conversation history stays empty because the LLM path errored
    # before the user/assistant pair could be appended.
    assert _SESSIONS[session_id].conversation_history == []


def test_resume_builder_message_handles_backtracking_via_llm(monkeypatch):
    """Mid-flow correction: user says 'actually my role is X' on turn 2,
    LLM returns target_role=X in draft_updates, and the new value
    overwrites the prior one."""
    from backend.services.resume_builder_service import _SESSIONS

    turns: list[dict] = [
        {
            "draft_updates": {"target_role": "ML Engineer"},
            "assistant_message": "Got it — ML Engineer. What's your most relevant experience?",
            "status": "collecting",
            "focus_field": "experience_notes",
        },
        {
            "draft_updates": {"target_role": "Senior ML Engineer"},
            "assistant_message": "Updated to Senior ML Engineer. Tell me about your most relevant role.",
            "status": "collecting",
            "focus_field": "experience_notes",
        },
    ]
    turn_index = {"value": 0}

    class _StubOpenAIService:
        def is_available(self):
            return True

        def run_json_prompt(self, system, user, **kwargs):
            payload = turns[turn_index["value"]]
            turn_index["value"] += 1
            return payload

    monkeypatch.setattr(
        "backend.routers.workspace._resolve_openai_service",
        lambda access_token, refresh_token: _StubOpenAIService(),
    )

    start = client.post("/api/workspace/resume-builder/start")
    session_id = start.json()["session_id"]
    auth_headers = {
        "X-Auth-Access-Token": "access-token",
        "X-Auth-Refresh-Token": "refresh-token",
    }

    first = client.post(
        "/api/workspace/resume-builder/message",
        json={"session_id": session_id, "message": "I'm targeting an ML Engineer role.", "input_mode": "text"},
        headers=auth_headers,
    )
    assert first.json()["draft_profile"]["target_role"] == "ML Engineer"

    second = client.post(
        "/api/workspace/resume-builder/message",
        json={
            "session_id": session_id,
            "message": "Actually, I should aim for Senior ML Engineer.",
            "input_mode": "text",
        },
        headers=auth_headers,
    )
    # The corrected role overwrote the earlier one.
    assert second.json()["draft_profile"]["target_role"] == "Senior ML Engineer"
    # Conversation history has both turn pairs.
    assert len(_SESSIONS[session_id].conversation_history) == 4


def test_resume_builder_message_recovers_session_after_cache_miss(monkeypatch):
    """Container restart wipes `_SESSIONS`, but the draft is in Supabase.
    Mutating routes lazy-load from the store before returning the
    existing 400 — verify the recovery path actually round-trips."""
    from backend.services.resume_builder_service import (
        _SESSIONS,
        export_resume_builder_session_payload,
    )
    from src.schemas import ResumeBuilderSessionRecord

    start_response = client.post("/api/workspace/resume-builder/start")
    assert start_response.status_code == 200
    session_id = start_response.json()["session_id"]

    seed_response = client.post(
        "/api/workspace/resume-builder/message",
        json={
            "session_id": session_id,
            "message": "Leander Antony\nChennai, India\nleander@example.com",
            "input_mode": "text",
        },
    )
    assert seed_response.status_code == 200
    assert seed_response.json()["current_step"] == "role"

    captured_payload = export_resume_builder_session_payload(session_id=session_id)

    _SESSIONS.clear()
    assert session_id not in _SESSIONS

    class _FakeAppUser:
        id = "user-1"

    class _FakeContext:
        app_user = _FakeAppUser()

    class _FakeStore:
        def load_latest_session(self, access, refresh, user_id):
            return ResumeBuilderSessionRecord(
                user_id=user_id,
                session_id=session_id,
                status="collecting",
                current_step="role",
                session_payload_json=captured_payload,
                updated_at="",
            )

        def save_session(self, access, refresh, payload):
            return ResumeBuilderSessionRecord(
                user_id=payload["user_id"],
                session_id=payload["session_id"],
                status="",
                current_step="",
                session_payload_json=payload["session_payload_json"],
                updated_at="",
            )

    monkeypatch.setattr(
        "backend.services.resume_builder_persistence_service._resolve_store",
        lambda access_token, refresh_token: (_FakeContext(), _FakeStore()),
    )

    recovered_response = client.post(
        "/api/workspace/resume-builder/message",
        json={
            "session_id": session_id,
            "message": "Machine Learning Engineer. AI engineer with applied AI experience.",
            "input_mode": "text",
        },
        headers={
            "X-Auth-Access-Token": "access-token",
            "X-Auth-Refresh-Token": "refresh-token",
        },
    )

    assert recovered_response.status_code == 200
    payload = recovered_response.json()
    assert payload["current_step"] == "experience"
    assert payload["draft_profile"]["target_role"] == "Machine Learning Engineer"


def test_workspace_analyze_returns_fit_and_artifacts_without_assisted_run():
    response = client.post(
        "/api/workspace/analyze",
        json={
            "resume_text": (
                "Leander Antony\n"
                "Chennai, India\n"
                "Python SQL Docker Communication AWS\n"
                "Experience\n"
                "AI Engineer, Example Labs\n"
                "Jan 2023 - Jan 2025\n"
                "Built production ML APIs and evaluation workflows.\n"
            ),
            "resume_filetype": "TXT",
            "resume_source": "workspace",
            "job_description_text": (
                "Machine Learning Engineer\n"
                "Location: Chennai, India\n"
                "Required: Python, SQL, Docker, AWS, communication.\n"
                "Need 3+ years of experience.\n"
            ),
            "run_assisted": False,
        },
    )

    assert response.status_code == 200
    payload = response.json()
    assert payload["workflow"]["mode"] == "deterministic_preview"
    assert payload["fit_analysis"]["overall_score"] >= 0
    assert payload["artifacts"]["tailored_resume"]["markdown"]
    assert payload["artifacts"]["cover_letter"]["markdown"]


def test_workspace_analyze_job_start_returns_job_handle(monkeypatch):
    monkeypatch.setattr(
        "backend.routers.workspace.start_workspace_analysis_job",
        lambda **kwargs: {
            "job_id": "job-123",
            "status": "queued",
            "stage_title": "Workflow crew",
            "stage_detail": "Preparing the first agent.",
            "progress_percent": 3,
            "result": None,
            "error_message": None,
        },
    )

    response = client.post(
        "/api/workspace/analyze-jobs",
        json={
            "resume_text": "Leander Antony\nPython\nExperience\nAI Engineer",
            "resume_filetype": "TXT",
            "resume_source": "workspace",
            "job_description_text": "Machine Learning Engineer\nRequired: Python",
            "run_assisted": True,
        },
    )

    assert response.status_code == 200
    payload = response.json()
    assert payload["job_id"] == "job-123"
    assert payload["status"] == "queued"
    assert payload["stage_title"] == "Workflow crew"
    assert payload["result"] is None
    assert payload["error_message"] is None


def test_workspace_analyze_job_status_returns_completed_result(monkeypatch):
    monkeypatch.setattr(
        "backend.routers.workspace._require_user_id",
        lambda *args, **kwargs: "owner-1",
    )
    monkeypatch.setattr(
        "backend.routers.workspace.get_workspace_analysis_job",
        lambda job_id, owner_user_id=None: {
            "job_id": job_id,
            "status": "completed",
            "stage_title": "Workflow crew",
            "stage_detail": "All agents are done.",
            "progress_percent": 100,
            "error_message": None,
            "result": {
                "resume_document": {"text": "resume", "filetype": "TXT", "source": "workspace"},
                "candidate_profile": {
                    "full_name": "Leander Antony",
                    "location": "Chennai, India",
                    "contact_lines": [],
                    "source": "workspace",
                    "resume_text": "resume",
                    "skills": [],
                    "experience": [],
                    "education": [],
                    "certifications": [],
                    "source_signals": [],
                },
                "job_description": {
                    "title": "Machine Learning Engineer",
                    "raw_text": "jd",
                    "cleaned_text": "jd",
                    "location": None,
                    "salary": None,
                    "requirements": {
                        "hard_skills": [],
                        "soft_skills": [],
                        "experience_requirement": None,
                        "must_haves": [],
                        "nice_to_haves": [],
                    },
                },
                "jd_summary_view": {
                    "headline": "headline",
                    "summary": "summary",
                    "sections": [],
                    "fit_signals": [],
                    "interview_focus": [],
                },
                "fit_analysis": {
                    "target_role": "Machine Learning Engineer",
                    "overall_score": 50,
                    "readiness_label": "Promising",
                    "matched_hard_skills": [],
                    "missing_hard_skills": [],
                    "matched_soft_skills": [],
                    "missing_soft_skills": [],
                    "experience_signal": "",
                    "strengths": [],
                    "gaps": [],
                    "recommendations": [],
                },
                "tailored_draft": {
                    "target_role": "Machine Learning Engineer",
                    "professional_summary": "",
                    "highlighted_skills": [],
                    "priority_bullets": [],
                    "gap_mitigation_steps": [],
                },
                "agent_result": None,
                "artifacts": {
                    "tailored_resume": {
                        "title": "resume",
                        "filename_stem": "resume",
                        "summary": "",
                        "markdown": "",
                        "plain_text": "",
                        "theme": "classic_ats",
                        "header": {
                            "full_name": "Leander Antony",
                            "location": "Chennai, India",
                            "contact_lines": [],
                        },
                        "target_role": "Machine Learning Engineer",
                        "professional_summary": "",
                        "highlighted_skills": [],
                        "experience_entries": [],
                        "education_entries": [],
                        "certifications": [],
                        "change_log": [],
                        "validation_notes": [],
                    },
                    "cover_letter": {
                        "title": "cover",
                        "filename_stem": "cover",
                        "summary": "",
                        "markdown": "",
                        "plain_text": "",
                    },
                    "report": {
                        "title": "report",
                        "filename_stem": "report",
                        "summary": "",
                        "markdown": "",
                        "plain_text": "",
                    },
                },
                "workflow": {
                    "mode": "openai",
                    "assisted_requested": True,
                    "assisted_available": True,
                    "review_approved": True,
                    "fallback_reason": "",
                },
                "imported_job_posting": None,
            },
        },
    )

    response = client.get("/api/workspace/analyze-jobs/job-123")

    assert response.status_code == 200
    payload = response.json()
    assert payload["status"] == "completed"
    assert payload["progress_percent"] == 100
    assert payload["result"]["job_description"]["title"] == "Machine Learning Engineer"


def test_workspace_analyze_job_start_returns_503_when_capacity_exhausted(monkeypatch):
    """When all concurrent run slots are taken, the start endpoint must
    fast-fail with 503 + Retry-After instead of queueing the request
    behind a thread spawn or returning a generic 500."""
    from backend.services.workspace_run_jobs import WorkspaceRunJobCapacityError

    def _raise_capacity(**_kwargs):
        raise WorkspaceRunJobCapacityError(
            "Too many agentic workflow runs are in flight right now."
        )

    monkeypatch.setattr(
        "backend.routers.workspace.start_workspace_analysis_job",
        _raise_capacity,
    )

    response = client.post(
        "/api/workspace/analyze-jobs",
        json={
            "resume_text": "Leander Antony\nPython\nExperience\nAI Engineer",
            "resume_filetype": "TXT",
            "resume_source": "workspace",
            "job_description_text": "Machine Learning Engineer\nRequired: Python",
            "run_assisted": True,
        },
    )

    assert response.status_code == 503
    assert response.headers.get("Retry-After")
    assert int(response.headers["Retry-After"]) > 0
    assert "busy" in response.json()["detail"].lower()


def test_workspace_analyze_job_status_returns_actionable_404_when_job_missing(monkeypatch):
    """Container restart drops `_JOBS`. The poll-side response should
    explain the cause and prompt a re-run, not return a bare
    'not found' that the hook surfaces unchanged."""
    monkeypatch.setattr(
        "backend.routers.workspace._require_user_id",
        lambda *args, **kwargs: "owner-1",
    )
    monkeypatch.setattr(
        "backend.routers.workspace.get_workspace_analysis_job",
        lambda job_id, owner_user_id=None: None,
    )

    response = client.get("/api/workspace/analyze-jobs/missing-job-id")

    assert response.status_code == 404
    detail = response.json()["detail"]
    assert "no longer available" in detail.lower()
    assert "run the workflow again" in detail.lower()


def test_workspace_analyze_job_cancel_route_returns_job_state(monkeypatch):
    """POST .../cancel returns the (serialized) job. It usually comes
    back still 'running' — the worker observes the flag at its next
    stage boundary — so the frontend keeps polling until 'cancelled'."""
    captured = {}

    def _cancel(job_id, owner_user_id=None):
        captured["job_id"] = job_id
        captured["owner_user_id"] = owner_user_id
        return {
            "job_id": job_id,
            "status": "running",
            "stage_title": "Forge agent",
            "stage_detail": "Stopping after the current step.",
            "progress_percent": 41,
            "result": None,
            "error_message": None,
        }

    monkeypatch.setattr(
        "backend.routers.workspace._require_user_id",
        lambda *args, **kwargs: "owner-1",
    )
    monkeypatch.setattr(
        "backend.routers.workspace.cancel_workspace_analysis_job",
        _cancel,
    )

    response = client.post("/api/workspace/analyze-jobs/job-xyz/cancel")

    assert response.status_code == 200
    assert captured["job_id"] == "job-xyz"
    # The route resolves the caller and scopes the cancel to them.
    assert captured["owner_user_id"] == "owner-1"
    payload = response.json()
    assert payload["job_id"] == "job-xyz"
    assert payload["status"] == "running"


def test_workspace_analyze_job_cancel_route_404_when_missing(monkeypatch):
    """An already-finished / pruned / wrong job id returns an
    actionable 404 the polling hook can surface verbatim."""
    monkeypatch.setattr(
        "backend.routers.workspace._require_user_id",
        lambda *args, **kwargs: "owner-1",
    )
    monkeypatch.setattr(
        "backend.routers.workspace.cancel_workspace_analysis_job",
        lambda job_id, owner_user_id=None: None,
    )

    response = client.post("/api/workspace/analyze-jobs/gone/cancel")

    assert response.status_code == 404
    detail = response.json()["detail"]
    assert "no longer available" in detail.lower()
    assert "nothing to stop" in detail.lower()


def test_workspace_analyze_prefers_imported_job_title_when_parser_cannot_extract_it():
    response = client.post(
        "/api/workspace/analyze",
        json={
            "resume_text": (
                "Leander Antony\n"
                "Chennai, India\n"
                "Python SQL Docker Communication AWS\n"
                "Experience\n"
                "AI Engineer, Example Labs\n"
                "Jan 2023 - Jan 2025\n"
                "Built production ML APIs and evaluation workflows.\n"
            ),
            "resume_filetype": "TXT",
            "resume_source": "workspace",
            "job_description_text": (
                "We are hiring a senior engineer to build data products and machine learning systems.\n"
                "Required: Python, SQL, Docker, AWS, communication.\n"
                "Need 5+ years of experience.\n"
            ),
            "imported_job_posting": {
                "id": "greenhouse:narvar:6930410",
                "source": "greenhouse",
                "title": "Staff Software Engineer, Machine Learning",
                "company": "Narvar",
                "location": "Bengaluru, Karnataka, India",
                "employment_type": "",
                "url": "https://job-boards.greenhouse.io/narvar/jobs/6930410",
                "summary": "Machine learning platform role.",
                "description_text": "Imported posting description",
                "posted_at": "",
                "scraped_at": "",
                "metadata": {},
            },
            "run_assisted": False,
        },
    )

    assert response.status_code == 200
    payload = response.json()
    assert payload["job_description"]["title"] == "Staff Software Engineer, Machine Learning"
    assert payload["job_description"]["location"] == "Bengaluru, Karnataka, India"
    assert "Staff Software Engineer, Machine Learning" in payload["artifacts"]["cover_letter"]["title"]
    assert "Staff Software Engineer, Machine Learning" in payload["artifacts"]["tailored_resume"]["title"]


def test_workspace_assistant_answer_uses_workspace_snapshot_context():
    analysis_response = client.post(
        "/api/workspace/analyze",
        json={
            "resume_text": (
                "Leander Antony\n"
                "Chennai, India\n"
                "Python SQL Docker Communication AWS\n"
                "Experience\n"
                "AI Engineer, Example Labs\n"
                "Jan 2023 - Jan 2025\n"
                "Built production ML APIs and evaluation workflows.\n"
            ),
            "resume_filetype": "TXT",
            "resume_source": "workspace",
            "job_description_text": (
                "Machine Learning Engineer\n"
                "Location: Chennai, India\n"
                "Required: Python, SQL, Docker, AWS, communication.\n"
                "Need 3+ years of experience.\n"
            ),
            "run_assisted": False,
        },
    )

    assert analysis_response.status_code == 200
    snapshot = analysis_response.json()

    response = client.post(
        "/api/workspace/assistant/answer",
        json={
            "question": "What are my biggest gaps for this role?",
            "current_page": "Workspace",
            "workspace_snapshot": snapshot,
            "history": [],
        },
    )

    assert response.status_code == 200
    payload = response.json()
    assert payload["answer"]
    assert payload["sources"]


def test_workspace_signature_is_stable_64_char_hash():
    """`workflow_signature` is the change-detection fingerprint stored
    on every saved-workspace row. Pin: it's 64 chars (sha256 hex) and
    deterministic across runs for the same canonical payload."""
    from backend.services.workspace_persistence_service import _workspace_signature

    snapshot = {
        "candidate_profile": {"full_name": "Leander Antony", "skills": ["Python", "AWS"]},
        "job_description": {"title": "Senior ML Engineer"},
        "fit_analysis": {"overall_score": 55, "matched_hard_skills": ["Python"]},
        "tailored_draft": {"target_role": "Senior ML Engineer"},
        # extra keys the signature must ignore
        "artifacts": {"tailored_resume": {"markdown": "# Resume"}},
        "agent_result": {"mode": "openai"},
    }

    signature = _workspace_signature(snapshot)

    assert isinstance(signature, str)
    assert len(signature) == 64
    assert all(c in "0123456789abcdef" for c in signature)

    # Same canonical payload → same hash. Re-ordering keys shouldn't
    # change the result because `json.dumps(sort_keys=True)` is the
    # canonicalization step.
    snapshot_reordered = dict(reversed(list(snapshot.items())))
    assert _workspace_signature(snapshot_reordered) == signature

    # Adding/changing a tracked field DOES change the hash.
    mutated = dict(snapshot)
    mutated["candidate_profile"] = {**snapshot["candidate_profile"], "skills": ["Python"]}
    assert _workspace_signature(mutated) != signature

    # Adding/changing an UNTRACKED field doesn't (artifacts, agent_result
    # aren't in the canonical payload).
    untracked = dict(snapshot)
    untracked["artifacts"] = {"tailored_resume": {"markdown": "# Different"}}
    assert _workspace_signature(untracked) == signature


def test_workspace_save_then_load_round_trips_artifact_markdown(monkeypatch):
    """The save/load round-trip should preserve the rendered tailored
    resume + cover letter markdown exactly. Container restarts (saved
    sessions outliving in-memory state) rely on this — if any field
    in the snapshot drops on serialize or fails to rehydrate, the
    user sees a different document than they saved."""
    from types import SimpleNamespace

    from backend.services import workspace_persistence_service
    from src.schemas import SavedWorkspaceRecord
    from src.services.fit_service import build_fit_analysis
    from src.services.job_service import build_job_description_from_text
    from src.services.profile_service import build_candidate_profile_from_resume
    from src.services.tailoring_service import build_tailored_resume_draft
    from src.resume_builder import build_tailored_resume_artifact
    from src.cover_letter_builder import build_cover_letter_artifact
    from src.schemas import ResumeDocument

    # Build a realistic snapshot deterministically (no LLM cost) so the
    # round-trip has substantive payload to preserve.
    resume_text = (
        "Leander Antony\n"
        "Chennai, India\n"
        "leander@example.com\n"
        "+91 9999999999\n"
        "Skills: Python, FastAPI, Docker, SQL, AWS\n"
        "Experience\n"
        "AI Engineer at Example Labs (Jan 2023 - Present)\n"
        "Built ML evaluation pipelines.\n"
        "Education\n"
        "Anna University, B.E. Computer Science\n"
    )
    jd_text = (
        "Senior ML Engineer\n"
        "Required: Python, AWS, SQL.\n"
        "Need 5+ years.\n"
    )
    candidate_profile = build_candidate_profile_from_resume(
        ResumeDocument(text=resume_text, filetype="TXT", source="test")
    )
    job_description = build_job_description_from_text(jd_text)
    fit_analysis = build_fit_analysis(candidate_profile, job_description)
    tailored_draft = build_tailored_resume_draft(
        candidate_profile, job_description, fit_analysis
    )
    tailored_resume = build_tailored_resume_artifact(
        candidate_profile, job_description, fit_analysis, tailored_draft
    )
    cover_letter = build_cover_letter_artifact(
        candidate_profile, job_description, fit_analysis, tailored_draft
    )

    # Snapshot shape mirrors what `run_workspace_analysis` produces.
    workspace_snapshot = {
        "candidate_profile": _to_serializable(candidate_profile),
        "job_description": _to_serializable(job_description),
        "fit_analysis": _to_serializable(fit_analysis),
        "tailored_draft": _to_serializable(tailored_draft),
        "agent_result": None,
        "imported_job_posting": None,
        "artifacts": {
            "tailored_resume": _to_serializable(tailored_resume),
            "cover_letter": _to_serializable(cover_letter),
        },
    }

    # In-memory mock store. Records are keyed by user_id and survive
    # the "session restart" — this is the fixture for verifying the
    # round-trip.
    storage: dict[str, dict] = {}

    class _MockStore:
        def __init__(self, *_args, **_kwargs):
            pass

        def is_configured(self):
            return True

        def save_workspace(self, access_token, refresh_token, payload):
            user_id = payload["user_id"]
            record = SavedWorkspaceRecord(
                user_id=user_id,
                job_title=payload["job_title"],
                workflow_signature=payload["workflow_signature"],
                workflow_snapshot_json=payload["workflow_snapshot_json"],
                cover_letter_payload_json=payload["cover_letter_payload_json"],
                tailored_resume_payload_json=payload["tailored_resume_payload_json"],
                expires_at="2099-01-01T00:00:00+00:00",
                updated_at="2026-05-07T12:00:00+00:00",
            )
            # Capture exactly what was persisted — this is the source
            # of truth the load path will rehydrate from.
            storage[user_id] = {
                "record": record,
                "raw_payload": dict(payload),
            }
            return record

        def load_workspace(self, access_token, refresh_token, user_id, now=None):
            entry = storage.get(user_id)
            if entry is None:
                return None, "missing"
            return entry["record"], "available"

    fake_context = SimpleNamespace(
        app_user=SimpleNamespace(id="test-user"),
        auth_service=SimpleNamespace(is_configured=lambda: True),
    )

    monkeypatch.setattr(
        workspace_persistence_service,
        "resolve_authenticated_context",
        lambda *, access_token, refresh_token: fake_context,
    )
    monkeypatch.setattr(
        workspace_persistence_service,
        "SavedWorkspaceStore",
        _MockStore,
    )

    save_response = workspace_persistence_service.save_workspace_snapshot(
        access_token="access-token",
        refresh_token="refresh-token",
        workspace_snapshot=workspace_snapshot,
    )
    assert save_response["status"] == "saved"
    assert save_response["saved_workspace"]["job_title"] == "Senior ML Engineer"

    # Simulate "session expiry": local in-memory state is gone, but
    # the DB record persists. The load endpoint pulls from storage
    # and re-renders the artifacts from the saved snapshot.
    load_response = workspace_persistence_service.load_saved_workspace_snapshot(
        access_token="access-token",
        refresh_token="refresh-token",
    )
    assert load_response["status"] == "available"

    rehydrated_snapshot = load_response["workspace_snapshot"]
    rehydrated_resume = rehydrated_snapshot["artifacts"]["tailored_resume"]
    rehydrated_cover_letter = rehydrated_snapshot["artifacts"]["cover_letter"]

    # Markdown is the user-visible payload — these must match exactly.
    assert rehydrated_resume["markdown"] == tailored_resume.markdown
    assert rehydrated_cover_letter["markdown"] == cover_letter.markdown
    # Title and filename_stem flow through the saved payload to the
    # re-rendered artifact. Mismatches here mean the export filename
    # would change between save and reload.
    assert rehydrated_resume["title"] == tailored_resume.title
    assert rehydrated_resume["filename_stem"] == tailored_resume.filename_stem
    assert rehydrated_cover_letter["title"] == cover_letter.title

    # Candidate identity round-trips byte-equal — name + email + skills
    # are the irreducible parts a recruiter sees.
    rehydrated_profile = rehydrated_snapshot["candidate_profile"]
    assert rehydrated_profile["full_name"] == candidate_profile.full_name
    assert rehydrated_profile["contact_lines"] == candidate_profile.contact_lines
    assert rehydrated_profile["skills"] == candidate_profile.skills
    # JD title round-trips so the saved-workspaces sidebar label stays
    # accurate after reload.
    assert rehydrated_snapshot["job_description"]["title"] == job_description.title


def _to_serializable(value):
    """Inline dataclass -> dict helper. Mirrors the persistence
    service's `_serialize` so the test doesn't import a private."""
    from dataclasses import asdict, is_dataclass

    if is_dataclass(value):
        return {k: _to_serializable(v) for k, v in asdict(value).items()}
    if isinstance(value, dict):
        return {k: _to_serializable(v) for k, v in value.items()}
    if isinstance(value, list):
        return [_to_serializable(item) for item in value]
    return value


def test_workspace_save_endpoint_forwards_auth_headers(monkeypatch):
    captured = {}

    def fake_save_workspace_snapshot(*, access_token, refresh_token, workspace_snapshot):
        captured["access_token"] = access_token
        captured["refresh_token"] = refresh_token
        captured["workspace_snapshot"] = workspace_snapshot
        return {
            "status": "saved",
            "saved_workspace": {
                "job_title": "Machine Learning Engineer",
                "expires_at": "2026-04-22T00:00:00+00:00",
                "updated_at": "2026-04-21T00:00:00+00:00",
            },
        }

    monkeypatch.setattr(
        "backend.routers.workspace.save_workspace_snapshot",
        fake_save_workspace_snapshot,
    )

    response = client.post(
        "/api/workspace/save",
        headers={
            "X-Auth-Access-Token": "access-token",
            "X-Auth-Refresh-Token": "refresh-token",
        },
        json={
            "workspace_snapshot": {
                "candidate_profile": {"full_name": "Leander Antony"},
                "job_description": {"title": "Machine Learning Engineer"},
                "fit_analysis": {"overall_score": 82},
                "tailored_draft": {"target_role": "Machine Learning Engineer"},
                "artifacts": {},
            }
        },
    )

    assert response.status_code == 200
    assert captured["access_token"] == "access-token"
    assert captured["refresh_token"] == "refresh-token"
    assert captured["workspace_snapshot"]["job_description"]["title"] == "Machine Learning Engineer"


def test_workspace_saved_endpoint_returns_saved_snapshot(monkeypatch):
    monkeypatch.setattr(
        "backend.routers.workspace.load_saved_workspace_snapshot",
        lambda access_token, refresh_token: {
            "status": "available",
            "saved_workspace": {
                "job_title": "Machine Learning Engineer",
                "expires_at": "2026-04-22T00:00:00+00:00",
                "updated_at": "2026-04-21T00:00:00+00:00",
            },
            "workspace_snapshot": {
                "resume_document": {
                    "text": "resume text",
                    "filetype": "Saved Workspace",
                    "source": "saved_workspace",
                },
                "candidate_profile": {"full_name": "Leander Antony"},
                "job_description": {
                    "title": "Machine Learning Engineer",
                    "raw_text": "raw jd",
                    "cleaned_text": "cleaned jd",
                    "location": "Chennai",
                    "requirements": {
                        "hard_skills": ["Python"],
                        "soft_skills": ["Communication"],
                        "experience_requirement": "3+ years",
                        "must_haves": [],
                        "nice_to_haves": [],
                    },
                },
                "jd_summary_view": {"sections": [{"title": "Overview", "items": ["ML role"]}]},
                "fit_analysis": {
                    "target_role": "Machine Learning Engineer",
                    "overall_score": 82,
                    "readiness_label": "Strong",
                    "matched_hard_skills": ["Python"],
                    "missing_hard_skills": [],
                    "matched_soft_skills": ["Communication"],
                    "missing_soft_skills": [],
                    "experience_signal": "3+ years",
                    "strengths": ["Production Python"],
                    "gaps": [],
                    "recommendations": [],
                },
                "tailored_draft": {
                    "target_role": "Machine Learning Engineer",
                    "professional_summary": "summary",
                    "highlighted_skills": ["Python"],
                    "priority_bullets": ["Built ML APIs"],
                    "gap_mitigation_steps": [],
                },
                "agent_result": None,
                "artifacts": {
                    "tailored_resume": {
                        "title": "Tailored Resume",
                        "filename_stem": "tailored-resume",
                        "summary": "summary",
                        "markdown": "markdown",
                        "plain_text": "plain text",
                        "theme": "classic_ats",
                        "header": {
                            "full_name": "Leander Antony",
                            "location": "Chennai",
                            "contact_lines": [],
                        },
                        "target_role": "Machine Learning Engineer",
                        "professional_summary": "summary",
                        "highlighted_skills": ["Python"],
                        "experience_entries": [],
                        "education_entries": [],
                        "certifications": [],
                        "change_log": [],
                        "validation_notes": [],
                    },
                    "cover_letter": {
                        "title": "Cover Letter",
                        "filename_stem": "cover-letter",
                        "summary": "summary",
                        "markdown": "markdown",
                        "plain_text": "plain text",
                    },
                    "report": {
                        "title": "Report",
                        "filename_stem": "report",
                        "summary": "summary",
                        "markdown": "markdown",
                        "plain_text": "plain text",
                    },
                },
                "workflow": {
                    "mode": "saved_workspace",
                    "assisted_requested": False,
                    "assisted_available": True,
                    "review_approved": False,
                    "fallback_reason": "",
                },
                "imported_job_posting": None,
            },
        },
    )

    response = client.get(
        "/api/workspace/saved",
        headers={
            "X-Auth-Access-Token": "access-token",
            "X-Auth-Refresh-Token": "refresh-token",
        },
    )

    assert response.status_code == 200
    payload = response.json()
    assert payload["status"] == "available"
    assert payload["workspace_snapshot"]["job_description"]["title"] == "Machine Learning Engineer"


def test_workspace_saved_jobs_endpoint_returns_shortlist(monkeypatch):
    monkeypatch.setattr(
        "backend.routers.workspace.list_saved_jobs",
        lambda access_token, refresh_token: {
            "status": "available",
            "saved_jobs": [
                {
                    "id": "job-123",
                    "title": "Machine Learning Engineer",
                    "company": "Example Labs",
                    "source": "greenhouse",
                    "location": "Chennai, India",
                    "employment_type": "Full-time",
                    "url": "https://example.com/jobs/job-123",
                    "summary": "Own production ML systems.",
                    "description_text": "Python SQL Docker",
                    "posted_at": "2026-04-20T10:00:00+00:00",
                    "scraped_at": "2026-04-21T10:00:00+00:00",
                    "metadata": {"departments": ["ML Platform"]},
                    "saved_at": "2026-04-21T10:05:00+00:00",
                    "updated_at": "2026-04-21T10:05:00+00:00",
                }
            ],
            "total_saved_jobs": 1,
            "latest_saved_at": "2026-04-21T10:05:00+00:00",
        },
    )

    response = client.get(
        "/api/workspace/saved-jobs",
        headers={
            "X-Auth-Access-Token": "access-token",
            "X-Auth-Refresh-Token": "refresh-token",
        },
    )

    assert response.status_code == 200
    payload = response.json()
    assert payload["total_saved_jobs"] == 1
    assert payload["saved_jobs"][0]["id"] == "job-123"


def test_workspace_saved_jobs_save_endpoint_forwards_auth_headers(monkeypatch):
    captured = {}

    def fake_save_saved_job(*, access_token, refresh_token, job_posting):
        captured["access_token"] = access_token
        captured["refresh_token"] = refresh_token
        captured["job_posting"] = job_posting
        return {
            "status": "saved",
            "saved_job": {
                **job_posting,
                "saved_at": "2026-04-21T10:05:00+00:00",
                "updated_at": "2026-04-21T10:05:00+00:00",
            },
            "message": "Saved Machine Learning Engineer to your shortlist.",
        }

    monkeypatch.setattr(
        "backend.routers.workspace.save_saved_job",
        fake_save_saved_job,
    )

    response = client.post(
        "/api/workspace/saved-jobs",
        headers={
            "X-Auth-Access-Token": "access-token",
            "X-Auth-Refresh-Token": "refresh-token",
        },
        json={
            "job_posting": {
                "id": "job-123",
                "title": "Machine Learning Engineer",
                "company": "Example Labs",
                "source": "greenhouse",
                "description_text": "Python SQL Docker",
                "metadata": {},
            }
        },
    )

    assert response.status_code == 200
    assert captured["access_token"] == "access-token"
    assert captured["refresh_token"] == "refresh-token"
    assert captured["job_posting"]["id"] == "job-123"


def test_workspace_saved_jobs_delete_endpoint_forwards_auth_headers(monkeypatch):
    captured = {}

    def fake_remove_saved_job(*, access_token, refresh_token, job_id):
        captured["access_token"] = access_token
        captured["refresh_token"] = refresh_token
        captured["job_id"] = job_id
        return {
            "status": "removed",
            "job_id": job_id,
        }

    monkeypatch.setattr(
        "backend.routers.workspace.remove_saved_job",
        fake_remove_saved_job,
    )

    response = client.delete(
        "/api/workspace/saved-jobs/job-123",
        headers={
            "X-Auth-Access-Token": "access-token",
            "X-Auth-Refresh-Token": "refresh-token",
        },
    )

    assert response.status_code == 200
    assert captured["access_token"] == "access-token"
    assert captured["refresh_token"] == "refresh-token"
    assert captured["job_id"] == "job-123"


def test_workspace_artifact_export_endpoint_forwards_snapshot(monkeypatch):
    captured = {}

    def fake_export_workspace_artifact(
        *, workspace_snapshot, artifact_kind, export_format, resume_theme, cover_letter_theme
    ):
        captured["workspace_snapshot"] = workspace_snapshot
        captured["artifact_kind"] = artifact_kind
        captured["export_format"] = export_format
        captured["resume_theme"] = resume_theme
        captured["cover_letter_theme"] = cover_letter_theme
        return {
            "status": "ready",
            "artifact_kind": artifact_kind,
            "export_format": export_format,
            "file_name": "leander-antony-machine-learning-engineer-tailored-resume.pdf",
            "mime_type": "application/pdf",
            "content_base64": "cGRm",
            "resume_theme": resume_theme,
            "cover_letter_theme": cover_letter_theme,
            "artifact_title": "Leander Antony - Machine Learning Engineer Tailored Resume",
        }

    monkeypatch.setattr(
        "backend.routers.workspace.export_workspace_artifact",
        fake_export_workspace_artifact,
    )

    response = client.post(
        "/api/workspace/artifacts/export",
        json={
            "workspace_snapshot": {
                "candidate_profile": {"full_name": "Leander Antony"},
                "job_description": {"title": "Machine Learning Engineer"},
                "fit_analysis": {"overall_score": 82},
                "tailored_draft": {"target_role": "Machine Learning Engineer"},
            },
            "artifact_kind": "tailored_resume",
            "export_format": "pdf",
            "resume_theme": "classic_ats",
        },
    )

    assert response.status_code == 200
    payload = response.json()
    assert payload["status"] == "ready"
    assert captured["artifact_kind"] == "tailored_resume"
    assert captured["export_format"] == "pdf"
    assert captured["resume_theme"] == "classic_ats"


def test_workspace_artifact_export_endpoint_accepts_docx(monkeypatch):
    """DOCX is the new download surface (Phase 2). Verify the route
    accepts export_format='docx' end to end and surfaces the file_name
    + mime_type the frontend expects."""
    captured = {}

    def fake_export_workspace_artifact(
        *, workspace_snapshot, artifact_kind, export_format, resume_theme, cover_letter_theme
    ):
        captured["export_format"] = export_format
        return {
            "status": "ready",
            "artifact_kind": artifact_kind,
            "export_format": export_format,
            "file_name": "leander-antony-machine-learning-engineer-tailored-resume.docx",
            "mime_type": (
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            ),
            "content_base64": "UEsDBA==",
            "resume_theme": resume_theme,
            "cover_letter_theme": cover_letter_theme,
            "artifact_title": "Leander Antony - Machine Learning Engineer Tailored Resume",
        }

    monkeypatch.setattr(
        "backend.routers.workspace.export_workspace_artifact",
        fake_export_workspace_artifact,
    )

    response = client.post(
        "/api/workspace/artifacts/export",
        json={
            "workspace_snapshot": {
                "candidate_profile": {"full_name": "Leander Antony"},
                "job_description": {"title": "Machine Learning Engineer"},
                "fit_analysis": {"overall_score": 82},
                "tailored_draft": {"target_role": "Machine Learning Engineer"},
            },
            "artifact_kind": "tailored_resume",
            "export_format": "docx",
            "resume_theme": "classic_ats",
        },
    )

    assert response.status_code == 200
    payload = response.json()
    assert captured["export_format"] == "docx"
    assert payload["file_name"].endswith(".docx")
    assert payload["mime_type"] == (
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )


def test_workspace_artifact_export_endpoint_rejects_markdown_format():
    """Markdown was removed as a download format in Phase 2. The
    request validator must reject it with 422 so old clients see a
    clear failure rather than silent fallthrough."""
    response = client.post(
        "/api/workspace/artifacts/export",
        json={
            "workspace_snapshot": {
                "candidate_profile": {"full_name": "Leander Antony"},
                "job_description": {"title": "Machine Learning Engineer"},
                "fit_analysis": {"overall_score": 82},
                "tailored_draft": {"target_role": "Machine Learning Engineer"},
            },
            "artifact_kind": "tailored_resume",
            "export_format": "markdown",
            "resume_theme": "classic_ats",
        },
    )

    assert response.status_code == 422
    body = response.json()
    # The Pydantic validator points at the export_format field with a
    # Literal mismatch — verify the error mentions either the field
    # name or the literal options so a debugger can find it fast.
    assert "export_format" in str(body)


def test_workspace_artifact_preview_endpoint_forwards_snapshot(monkeypatch):
    captured = {}

    def fake_preview_workspace_artifact(
        *, workspace_snapshot, artifact_kind, resume_theme, cover_letter_theme
    ):
        captured["workspace_snapshot"] = workspace_snapshot
        captured["artifact_kind"] = artifact_kind
        captured["resume_theme"] = resume_theme
        captured["cover_letter_theme"] = cover_letter_theme
        return {
            "status": "ready",
            "artifact_kind": artifact_kind,
            "resume_theme": resume_theme,
            "cover_letter_theme": cover_letter_theme,
            "artifact_title": "Leander Antony - Machine Learning Engineer Tailored Resume",
            "html": "<html><body><h1>Preview</h1></body></html>",
        }

    monkeypatch.setattr(
        "backend.routers.workspace.preview_workspace_artifact",
        fake_preview_workspace_artifact,
    )

    response = client.post(
        "/api/workspace/artifacts/preview",
        json={
            "workspace_snapshot": {
                "candidate_profile": {"full_name": "Leander Antony"},
                "job_description": {"title": "Machine Learning Engineer"},
                "fit_analysis": {"overall_score": 82},
                "tailored_draft": {"target_role": "Machine Learning Engineer"},
            },
            "artifact_kind": "tailored_resume",
            "resume_theme": "classic_ats",
        },
    )

    assert response.status_code == 200
    payload = response.json()
    assert payload["status"] == "ready"
    assert payload["artifact_kind"] == "tailored_resume"
    assert captured["resume_theme"] == "classic_ats"


# ─── PostHog funnel events ──────────────────────────────────────────────


def test_resume_upload_emits_resume_uploaded_event(monkeypatch):
    """A successful /resume/upload emits a `resume_uploaded` PostHog
    funnel event carrying the file type."""
    events = []
    monkeypatch.setattr(
        "backend.routers.workspace.capture_event",
        lambda **kwargs: events.append(kwargs),
    )
    response = client.post(
        "/api/workspace/resume/upload",
        json=_encode_text_file_payload(
            "resume.txt",
            "Leander Antony\nChennai\nPython SQL Docker\n"
            "Experience\nAI Engineer, Example Labs\n",
        ),
    )

    assert response.status_code == 200
    assert [e["event"] for e in events] == ["resume_uploaded"]
    assert events[0]["properties"]["file_type"] == "txt"


def test_analyze_emits_analysis_started_event(monkeypatch):
    """A successful /analyze emits an `analysis_started` event tagged
    with the sync mode."""
    events = []
    monkeypatch.setattr(
        "backend.routers.workspace.capture_event",
        lambda **kwargs: events.append(kwargs),
    )
    response = client.post(
        "/api/workspace/analyze",
        json={
            "resume_text": (
                "Leander Antony\nPython SQL Docker AWS\n"
                "Experience\nAI Engineer, Example Labs\n"
                "Jan 2023 - Jan 2025\nBuilt production ML APIs.\n"
            ),
            "resume_filetype": "TXT",
            "resume_source": "workspace",
            "job_description_text": (
                "Machine Learning Engineer\n"
                "Required: Python, SQL, Docker, AWS.\n"
            ),
            "run_assisted": False,
        },
    )

    assert response.status_code == 200
    assert [e["event"] for e in events] == ["analysis_started"]
    assert events[0]["properties"]["mode"] == "sync"
    assert events[0]["properties"]["run_assisted"] is False


def test_artifact_export_emits_artifact_exported_event(monkeypatch):
    """A successful /artifacts/export emits an `artifact_exported`
    event with the artifact kind + format."""
    events = []
    monkeypatch.setattr(
        "backend.routers.workspace.capture_event",
        lambda **kwargs: events.append(kwargs),
    )

    def fake_export_workspace_artifact(
        *, workspace_snapshot, artifact_kind, export_format,
        resume_theme, cover_letter_theme
    ):
        return {
            "status": "ready",
            "artifact_kind": artifact_kind,
            "export_format": export_format,
            "file_name": "resume.pdf",
            "mime_type": "application/pdf",
            "content_base64": "cGRm",
            "resume_theme": resume_theme,
            "cover_letter_theme": cover_letter_theme,
            "artifact_title": "Resume",
        }

    monkeypatch.setattr(
        "backend.routers.workspace.export_workspace_artifact",
        fake_export_workspace_artifact,
    )
    response = client.post(
        "/api/workspace/artifacts/export",
        json={
            "workspace_snapshot": {
                "candidate_profile": {"full_name": "Leander Antony"},
                "job_description": {"title": "Machine Learning Engineer"},
                "fit_analysis": {"overall_score": 82},
                "tailored_draft": {"target_role": "Machine Learning Engineer"},
            },
            "artifact_kind": "tailored_resume",
            "export_format": "pdf",
            "resume_theme": "classic_ats",
        },
    )

    assert response.status_code == 200
    assert [e["event"] for e in events] == ["artifact_exported"]
    assert events[0]["properties"]["artifact_kind"] == "tailored_resume"
    assert events[0]["properties"]["export_format"] == "pdf"
