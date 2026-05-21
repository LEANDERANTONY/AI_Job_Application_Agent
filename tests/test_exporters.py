from io import BytesIO
from unittest.mock import patch

from src.errors import ExportError

from src.exporters import (
    _build_resume_contact_inline_html,
    _build_resume_html,
    _looks_like_contact_link,
    build_cover_letter_preview_html,
    export_docx_bytes,
    export_pdf_bytes,
    generate_pdf,
)
from src.schemas import (
    CoverLetterArtifact,
    EducationEntry,
    ProjectEntry,
    ResumeDocument,
    ResumeExperienceEntry,
    ResumeHeader,
    TailoredResumeArtifact,
    WorkExperience,
)


def test_build_cover_letter_preview_html_contains_structure():
    artifact = CoverLetterArtifact(
        title="Leander Antony - Data Analyst Cover Letter",
        filename_stem="candidate-cover-letter",
        summary="Grounded cover letter draft.",
        markdown="# Leander Antony - Data Analyst Cover Letter\n\nDear Hiring Team,\n\nI am excited to apply for the role.",
        plain_text="Leander Antony - Data Analyst Cover Letter\n\nDear Hiring Team,\n\nI am excited to apply for the role.",
    )

    html_output = build_cover_letter_preview_html(artifact)

    assert 'class="cover-letter-title"' in html_output
    assert "Leander Antony" in html_output
    assert 'class="cover-letter-role"' in html_output
    assert "Data Analyst" in html_output
    assert "Dear Hiring Team" in html_output
    assert 'class="cover-letter-greeting-break"' in html_output
    assert "font-size: 11.4pt;" in html_output


def test_build_resume_html_uses_classic_template_with_warm_neutral_palette():
    html_output = _build_resume_html("# Candidate\n\n## Experience", theme="classic_ats")

    assert "resume-shell--classic" in html_output
    assert "@page {{ size: A4; margin: 0; }}" not in html_output
    assert "@page { size: A4; margin: 0; }" in html_output
    # Warm-brown palette values must appear in the :root block; the
    # rest of the CSS now references them via var(--*) so we no
    # longer assert on the literal property declarations.
    assert "#221912" in html_output  # --ink
    assert "#fffdf9" in html_output  # --paper
    assert "#fffdfa" in html_output  # --surface
    assert "#8f6845" in html_output  # --accent (warm brown)
    assert "#d7c2af" in html_output  # --line
    assert "margin: 0 -15mm 6px;" in html_output
    # Structural classes are still here.
    assert ".resume-experience-card + .resume-experience-card" in html_output
    assert ".resume-classic-section--plain-head h2" in html_output
    # Pagination guards added when we lifted the one-page clip.
    assert "page-break-inside: avoid" in html_output


def test_build_resume_html_omits_empty_contact_card():
    artifact = TailoredResumeArtifact(
        title="Tailored Resume",
        filename_stem="tailored-resume",
        summary="Structured resume",
        markdown="# Resume",
        plain_text="Resume",
        theme="classic_ats",
        header=ResumeHeader(
            full_name="Leander Antony",
            location="Chennai, India",
            contact_lines=[],
        ),
        target_role="AI Engineer",
        professional_summary="Builds grounded AI workflow products.",
        highlighted_skills=["Python", "SQL"],
        experience_entries=[],
        education_entries=[],
        certifications=[],
        validation_notes=["Internal only"],
    )

    html_output = _build_resume_html(artifact.markdown, theme="classic_ats", artifact=artifact)

    assert "<h2>Contact</h2>" not in html_output


def test_build_resume_html_renders_classic_ats_as_single_column_blue_layout():
    artifact = TailoredResumeArtifact(
        title="Tailored Resume",
        filename_stem="tailored-resume",
        summary="Structured resume",
        markdown="# Resume",
        plain_text="Resume",
        theme="classic_ats",
        header=ResumeHeader(
            full_name="Leander Antony",
            location="Chennai, India",
            contact_lines=["leander@example.com", "+91 99999 99999"],
        ),
        target_role="AI Engineer",
        professional_summary="Builds grounded AI workflow products.",
        highlighted_skills=["Python", "SQL"],
        experience_entries=[
            ResumeExperienceEntry(
                title="AI Engineer",
                organization="Example Labs",
                location="Remote",
                start="2023",
                end="Present",
                bullets=["Built ML workflow products."],
            )
        ],
        education_entries=[],
        certifications=[],
    )

    html_output = _build_resume_html(artifact.markdown, theme="classic_ats", artifact=artifact)

    assert "resume-classic-header" in html_output
    assert "resume-contact-inline" in html_output
    assert "resume-classic-section" in html_output
    assert "resume-classic-sidebar" not in html_output
    assert "resume-modern-contact" not in html_output
    # Mode-aware headline: this artifact sets target_role="AI Engineer",
    # so the (formerly dormant) role line now renders — between the
    # name and the contact block. (When target_role is "" it is omitted
    # entirely; see test_resume_headline_*.) NOTE: assert on the
    # ELEMENT tag, not the bare class token — `.resume-classic-role`
    # also appears in the <style> block.
    assert '<p class="resume-classic-role">AI Engineer</p>' in html_output
    assert html_output.index("<h1>") < html_output.index(
        '<p class="resume-classic-role">'
    ) < html_output.index('<p class="resume-contact-inline')
    assert "Chennai, India" in html_output
    assert "leander@example.com" in html_output


def test_export_pdf_bytes_passes_tailored_resume_artifact_to_generator():
    artifact = TailoredResumeArtifact(
        title="Tailored Resume",
        filename_stem="tailored-resume",
        summary="Structured resume",
        markdown="# Resume",
        plain_text="Resume",
        theme="classic_ats",
    )

    with patch("src.exporters.generate_pdf", return_value=BytesIO(b"%PDF-structured")) as mock_generate_pdf:
        pdf_bytes = export_pdf_bytes(artifact)

    assert pdf_bytes == b"%PDF-structured"
    assert mock_generate_pdf.call_args.kwargs["artifact"] is artifact
    assert mock_generate_pdf.call_args.kwargs["document_kind"] == "tailored_resume"


def test_export_pdf_bytes_treats_cover_letter_as_cover_letter_document():
    artifact = CoverLetterArtifact(
        title="Candidate Cover Letter",
        filename_stem="candidate-cover-letter",
        summary="Grounded cover letter draft.",
        markdown="# Candidate Cover Letter\n\nDear Hiring Team,\n\nI am excited to apply for the role.",
        plain_text="Candidate Cover Letter\n\nDear Hiring Team,\n\nI am excited to apply for the role.",
    )

    with patch("src.exporters.generate_pdf", return_value=BytesIO(b"%PDF-cover-letter")) as mock_generate_pdf:
        pdf_bytes = export_pdf_bytes(artifact)

    assert pdf_bytes == b"%PDF-cover-letter"
    assert mock_generate_pdf.call_args.kwargs["artifact"] is None
    assert mock_generate_pdf.call_args.kwargs["document_kind"] == "cover_letter"


@patch("src.exporters._generate_pdf_with_weasyprint", side_effect=RuntimeError("renderer unavailable"))
def test_export_pdf_bytes_falls_back_when_weasyprint_backend_fails(_mock_generate_pdf):
    artifact = CoverLetterArtifact(
        title="Candidate Cover Letter",
        filename_stem="candidate-cover-letter",
        summary="Grounded cover letter draft.",
        markdown="# Candidate Cover Letter\n\nDear Hiring Team,\n\nI am excited to apply for the role.",
        plain_text="Candidate Cover Letter\n\nDear Hiring Team,\n\nI am excited to apply for the role.",
    )

    pdf_bytes = export_pdf_bytes(artifact)

    assert pdf_bytes.startswith(b"%PDF")
    assert len(pdf_bytes) > 500


@patch("src.exporters._generate_pdf_with_reportlab", return_value=BytesIO(b"%PDF-skip"))
@patch("src.exporters._generate_pdf_with_weasyprint", side_effect=OSError("libgobject missing"))
def test_generate_pdf_falls_back_when_weasyprint_runtime_is_unavailable(
    _mock_weasyprint,
    mock_reportlab,
):
    pdf_buffer = generate_pdf("# Resume", title="Test Resume")

    assert pdf_buffer.getvalue() == b"%PDF-skip"
    mock_reportlab.assert_called_once_with("# Resume", "Test Resume")


@patch("src.exporters._generate_pdf_with_reportlab", side_effect=RuntimeError("reportlab failed"))
@patch("src.exporters._generate_pdf_with_weasyprint", side_effect=RuntimeError("weasyprint failed"))
def test_export_pdf_bytes_raises_export_error_when_both_backends_fail(
    _mock_weasyprint,
    _mock_reportlab,
):
    artifact = CoverLetterArtifact(
        title="Candidate Cover Letter",
        filename_stem="candidate-cover-letter",
        summary="Grounded cover letter draft.",
        markdown="# Candidate Cover Letter\n",
        plain_text="Candidate Cover Letter\n",
    )

    try:
        export_pdf_bytes(artifact)
        assert False, "Expected ExportError"
    except ExportError:
        assert True


# ---------------------------------------------------------------------------
# Edge cases — multi-page volume, Unicode names, RTL text, very long
# bullets, sparse profiles. Focus on the HTML layer (where structural
# bugs surface); WeasyPrint runtime is platform-fragile and already
# covered by fallback tests above.
# ---------------------------------------------------------------------------


def _make_resume_artifact(**overrides) -> TailoredResumeArtifact:
    """Builder for edge-case test artifacts. Defaults are minimal but
    the resulting HTML always contains a full structured shell."""
    base = dict(
        title="Tailored Resume",
        filename_stem="tailored-resume",
        summary="Structured resume",
        markdown="# Resume",
        plain_text="Resume",
        theme="classic_ats",
        header=ResumeHeader(
            full_name="Test Candidate",
            location="Earth",
            contact_lines=["test@example.com"],
        ),
        target_role="Test Role",
        professional_summary="Built things.",
        highlighted_skills=["Python"],
        experience_entries=[],
        education_entries=[],
        certifications=[],
    )
    base.update(overrides)
    return TailoredResumeArtifact(**base)


def test_resume_html_renders_long_resume_without_truncating_entries():
    """A 10-entry resume should not silently lose later entries — every
    job's bullets must surface in the rendered HTML."""
    entries = [
        ResumeExperienceEntry(
            title=f"Engineer {idx}",
            organization=f"Company {idx}",
            location="Remote",
            start=f"20{idx:02d}",
            end=f"20{idx + 1:02d}",
            bullets=[
                f"Built a system handling traffic at company {idx}.",
                f"Reduced latency by {idx * 5}% across the platform.",
                f"Owned the on-call rotation in role {idx}.",
            ],
        )
        for idx in range(10)
    ]
    artifact = _make_resume_artifact(experience_entries=entries)

    html_output = _build_resume_html(
        artifact.markdown, theme="classic_ats", artifact=artifact
    )

    for idx in range(10):
        assert f"Company {idx}" in html_output, (
            f"company {idx} dropped from rendered HTML"
        )
        assert f"latency by {idx * 5}%" in html_output


def test_resume_html_preserves_unicode_in_candidate_header():
    """Candidate names in non-Latin scripts must round-trip through the
    HTML without being escaped to entities or stripped."""
    artifact = _make_resume_artifact(
        header=ResumeHeader(
            full_name="李明 (Li Ming)",
            location="北京, China",
            contact_lines=["liming@example.cn"],
        ),
    )

    html_output = _build_resume_html(
        artifact.markdown, theme="classic_ats", artifact=artifact
    )

    assert "李明" in html_output
    assert "北京" in html_output


def test_resume_html_preserves_arabic_name_and_rtl_text():
    """Arabic candidates should at least see their name preserved.
    The current renderer does NOT add `dir=\"rtl\"` to the root, so
    Arabic text renders left-to-right by default — pin that behavior
    so any future RTL handling change is intentional."""
    artifact = _make_resume_artifact(
        header=ResumeHeader(
            full_name="محمد علي",
            location="القاهرة, Egypt",
            contact_lines=["mohamed@example.eg"],
        ),
        professional_summary="Senior engineer with experience leading Arabic-language NLP projects.",
    )

    html_output = _build_resume_html(
        artifact.markdown, theme="classic_ats", artifact=artifact
    )

    assert "محمد علي" in html_output
    assert "القاهرة" in html_output
    # Documented gap: no auto-RTL. If the renderer learns RTL, this
    # assertion flips and the test surfaces the change as a deliberate
    # decision rather than silent regression.
    assert 'dir="rtl"' not in html_output


def test_resume_html_preserves_accented_european_names():
    artifact = _make_resume_artifact(
        header=ResumeHeader(
            full_name="François Müller",
            location="Zürich, Schweiz",
            contact_lines=["francois@example.ch"],
        ),
    )

    html_output = _build_resume_html(
        artifact.markdown, theme="classic_ats", artifact=artifact
    )

    assert "François Müller" in html_output
    assert "Zürich" in html_output


def test_resume_html_preserves_very_long_bullet_without_truncation():
    """A pathologically long bullet (5000 chars) shouldn't get silently
    sliced. The visible result will overflow the page in PDF — but
    that's a layout problem, not a data-loss bug. We pin the
    no-truncation behavior so any future cap is added intentionally."""
    long_bullet = "Built " + ("complex distributed-systems infrastructure " * 120)
    assert len(long_bullet) > 5000
    entries = [
        ResumeExperienceEntry(
            title="Engineer",
            organization="Acme",
            location="Remote",
            start="2020",
            end="Present",
            bullets=[long_bullet],
        )
    ]
    artifact = _make_resume_artifact(experience_entries=entries)

    html_output = _build_resume_html(
        artifact.markdown, theme="classic_ats", artifact=artifact
    )

    # The final segment of the bullet survives.
    assert long_bullet[-200:] in html_output


def test_resume_html_renders_sparse_profile_without_error():
    """When the candidate has nothing but a name + email, the renderer
    must still produce a valid HTML shell (empty-state friendly) rather
    than crashing on missing sections."""
    artifact = _make_resume_artifact(
        header=ResumeHeader(
            full_name="Sparse Candidate",
            location="",
            contact_lines=["sparse@example.com"],
        ),
        target_role="",
        professional_summary="",
        highlighted_skills=[],
        experience_entries=[],
        education_entries=[],
        certifications=[],
    )

    html_output = _build_resume_html(
        artifact.markdown, theme="classic_ats", artifact=artifact
    )

    # Shell + header still render.
    assert "Sparse Candidate" in html_output
    assert "resume-shell--classic" in html_output
    # The renderer keeps section scaffolding consistent and surfaces
    # empty-state copy instead of dropping headers — verify the
    # scaffold is present so a future "drop empty sections" change is
    # intentional, and that the empty-state message appears.
    assert "resume-empty" in html_output
    assert "<h2>Education</h2>" in html_output


def test_cover_letter_html_preserves_unicode_signoff():
    """Signoff with a non-ASCII name (e.g., 'Sincerely, François')
    must round-trip the accented character in the rendered greeting."""
    artifact = CoverLetterArtifact(
        title="François Müller - Senior Engineer Cover Letter",
        filename_stem="francois-cover-letter",
        summary="Grounded cover letter draft.",
        markdown=(
            "# François Müller - Senior Engineer Cover Letter\n\n"
            "Dear Hiring Team,\n\n"
            "I am excited to apply for the role.\n\n"
            "Sincerely,\n\nFrançois"
        ),
        plain_text="placeholder",
    )

    html_output = build_cover_letter_preview_html(artifact)

    assert "François Müller" in html_output
    assert "Sincerely" in html_output
    assert "François" in html_output


# ---------------------------------------------------------------------------
# DOCX export — Phase 1
#
# Tests parse the rendered DOCX bytes back through `python-docx` and
# assert structural shape (heading text, bullet count, section
# ordering). We don't try to assert on visual styling — that's the
# manual QA loop in Phase 4.
# ---------------------------------------------------------------------------


def _parse_docx(data: bytes):
    from docx import Document

    return Document(BytesIO(data))


def _docx_paragraph_pairs(doc):
    """Return list of (style_name, text) for every paragraph."""
    return [(p.style.name, p.text) for p in doc.paragraphs]


def _make_full_resume_artifact() -> TailoredResumeArtifact:
    return TailoredResumeArtifact(
        title="Leander Antony - Senior ML Engineer Tailored Resume",
        filename_stem="leander-tailored-resume",
        summary="Tailored summary",
        markdown="# Resume",
        plain_text="Resume",
        theme="classic_ats",
        header=ResumeHeader(
            full_name="Leander Antony",
            location="Chennai, India",
            contact_lines=["leander@example.com", "+91 99999 99999", "linkedin.com/in/leander"],
        ),
        target_role="Senior ML Engineer",
        professional_summary=(
            "Senior ML engineer with 5 years building distributed Python "
            "systems on AWS and Postgres."
        ),
        highlighted_skills=["Python", "AWS", "Docker", "Postgres", "FastAPI"],
        experience_entries=[
            ResumeExperienceEntry(
                title="AI Engineer",
                organization="Example Labs",
                location="Remote",
                start="Jan 2023",
                end="Present",
                bullets=[
                    "Built FastAPI services that ship LLM evaluation reports.",
                    "Reduced inference latency 30% via batching and caching.",
                    "Owned the on-call rotation for the model API.",
                ],
            ),
            ResumeExperienceEntry(
                title="ML Intern",
                organization="Acme Co",
                location="Bangalore",
                start="2022",
                end="2023",
                bullets=["Wrote eval harnesses for 3 production models."],
            ),
        ],
        project_entries=[
            ProjectEntry(
                name="Open-source resume parser",
                description="LLM-backed resume parser used in ~5K downloads.",
                bullets=["Added Unicode-aware name detection.", "PR-reviewed by 4 contributors."],
                technologies=["Python", "FastAPI"],
                link="github.com/leander/resume-parser",
                start="2024",
                end="Present",
            ),
        ],
        education_entries=[
            EducationEntry(
                institution="Anna University",
                degree="B.E.",
                field_of_study="Computer Science",
                start="2016",
                end="2020",
            ),
        ],
        certifications=["AWS Certified ML Specialty", "GCP Professional ML Engineer"],
        publication_entries=["Distributed Eval at Scale (2024)"],
        section_order=[
            "summary",
            "skills",
            "experience",
            "projects",
            "education",
            "publications",
            "certifications",
        ],
    )


def test_export_docx_bytes_renders_full_resume_with_all_sections():
    artifact = _make_full_resume_artifact()

    data = export_docx_bytes(artifact)

    assert isinstance(data, bytes)
    # Real .docx files start with the PK ZIP magic. Catches the case
    # where we accidentally return a string or a buffer.
    assert data.startswith(b"PK")
    assert len(data) > 5_000

    doc = _parse_docx(data)
    pairs = _docx_paragraph_pairs(doc)

    # Header: name -> role headline -> contact line. This artifact
    # sets target_role="Senior ML Engineer", so the DOCX header
    # renders the uppercase role line between the name and contact
    # (mirrors the HTML/PDF `.resume-classic-role`; the DOCX builder
    # was fixed 2026-05-21 to match). When target_role is "" the role
    # line is omitted — covered by test_export_docx_bytes_drops_* /
    # the headline tests.
    paragraph_texts = [text for _, text in pairs]
    assert "Leander Antony" in paragraph_texts[0]
    assert paragraph_texts[1] == "SENIOR ML ENGINEER"
    assert "Chennai, India" in paragraph_texts[2]
    assert "leander@example.com" in paragraph_texts[2]

    # Section headings render as uppercase labels in the order asked
    # for. Names match the artifact.section_order list 1:1.
    heading_texts = [
        text
        for _, text in pairs
        if text in {"SUMMARY", "CORE SKILLS", "EXPERIENCE", "PROJECTS", "EDUCATION", "PUBLICATIONS", "CERTIFICATIONS"}
    ]
    assert heading_texts == [
        "SUMMARY",
        "CORE SKILLS",
        "EXPERIENCE",
        "PROJECTS",
        "EDUCATION",
        "PUBLICATIONS",
        "CERTIFICATIONS",
    ]

    # Bullets render with Word's built-in 'List Bullet' style so they
    # open as a proper bullet list in Word and Google Docs.
    bullet_texts = [text for style, text in pairs if style == "List Bullet"]
    # 3 experience bullets + 1 intern bullet + 2 project bullets +
    # 1 publication + 2 certifications = 9.
    assert len(bullet_texts) == 9
    assert "Built FastAPI services that ship LLM evaluation reports." in bullet_texts
    assert "Distributed Eval at Scale (2024)" in bullet_texts
    assert "AWS Certified ML Specialty" in bullet_texts


def test_export_docx_bytes_drops_empty_optional_sections():
    """Sparse profiles legitimately miss Experience / Projects /
    Publications / Certifications. Those sections should drop entirely
    rather than render an empty header — matches the HTML render's
    behavior."""
    artifact = TailoredResumeArtifact(
        title="Sparse Candidate Resume",
        filename_stem="sparse",
        summary="",
        markdown="",
        plain_text="",
        theme="classic_ats",
        header=ResumeHeader(
            full_name="Sparse Candidate",
            location="Remote",
            contact_lines=["sparse@example.com"],
        ),
        professional_summary="Recent graduate seeking entry-level ML roles.",
        highlighted_skills=["Python"],
        education_entries=[
            EducationEntry(institution="Anna University", degree="B.E. CS"),
        ],
        section_order=["summary", "skills", "experience", "projects", "education", "publications", "certifications"],
    )

    data = export_docx_bytes(artifact)
    doc = _parse_docx(data)

    headings = [text for _, text in _docx_paragraph_pairs(doc)]
    # Required sections render even when sparse.
    assert "SUMMARY" in headings
    assert "CORE SKILLS" in headings
    assert "EDUCATION" in headings
    # Optional / empty sections should NOT render their heading.
    assert "EXPERIENCE" not in headings
    assert "PROJECTS" not in headings
    assert "PUBLICATIONS" not in headings
    assert "CERTIFICATIONS" not in headings


def test_export_docx_bytes_honors_custom_section_order():
    """`section_order` drives the section sequence so students can lead
    with Education, academics with Publications, seniors with
    Experience after Skills. Verify a non-default order rounds through
    to the rendered DOCX."""
    artifact = _make_full_resume_artifact()
    artifact.section_order = [
        "summary",
        "education",
        "skills",
        "experience",
        "projects",
        "certifications",
        "publications",
    ]

    data = export_docx_bytes(artifact)
    doc = _parse_docx(data)

    headings_in_doc_order = [
        text
        for _, text in _docx_paragraph_pairs(doc)
        if text in {"SUMMARY", "CORE SKILLS", "EXPERIENCE", "PROJECTS", "EDUCATION", "PUBLICATIONS", "CERTIFICATIONS"}
    ]
    assert headings_in_doc_order == [
        "SUMMARY",
        "EDUCATION",
        "CORE SKILLS",
        "EXPERIENCE",
        "PROJECTS",
        "CERTIFICATIONS",
        "PUBLICATIONS",
    ]


def test_export_docx_bytes_appends_missing_sections_for_partial_orders():
    """If an agent emits a partial section_order, the renderer should
    append any missing sections at the end so we never silently drop
    user-supplied content."""
    artifact = _make_full_resume_artifact()
    # Drop publications + certifications from the order to simulate a
    # partial agent response.
    artifact.section_order = ["summary", "skills", "experience", "projects", "education"]

    data = export_docx_bytes(artifact)
    doc = _parse_docx(data)

    headings = [
        text
        for _, text in _docx_paragraph_pairs(doc)
        if text in {"PUBLICATIONS", "CERTIFICATIONS"}
    ]
    # Both should still appear because they have content, even though
    # the agent's order didn't list them.
    assert headings == ["PUBLICATIONS", "CERTIFICATIONS"]


def test_export_docx_bytes_preserves_unicode_in_header_and_bullets():
    """Non-Latin names + accented strings must round-trip the DOCX
    layer the same way they survive the HTML render."""
    artifact = TailoredResumeArtifact(
        title="François Müller Resume",
        filename_stem="francois",
        summary="",
        markdown="",
        plain_text="",
        theme="classic_ats",
        header=ResumeHeader(
            full_name="François Müller",
            location="Zürich, Schweiz",
            contact_lines=["francois@example.ch"],
        ),
        professional_summary="Senior engineer based in Zürich.",
        highlighted_skills=["Python", "Rust"],
        experience_entries=[
            ResumeExperienceEntry(
                title="Engineer",
                organization="Société Générale",
                start="2020",
                end="Present",
                bullets=["Refactored the légère pipeline."],
            ),
        ],
        education_entries=[EducationEntry(institution="ETH Zürich")],
    )

    data = export_docx_bytes(artifact)
    doc = _parse_docx(data)

    text_blob = "\n".join(p.text for p in doc.paragraphs)
    assert "François Müller" in text_blob
    assert "Zürich, Schweiz" in text_blob
    assert "Société Générale" in text_blob
    assert "Refactored the légère pipeline." in text_blob


def test_export_docx_bytes_renders_cover_letter_with_paragraphs_and_bullets():
    """The cover letter artifact only exposes flat markdown; verify the
    DOCX path parses it into the right sequence of paragraphs and
    bullet items."""
    artifact = CoverLetterArtifact(
        title="Leander Antony - Senior ML Engineer Cover Letter",
        filename_stem="leander-cover-letter",
        summary="",
        markdown=(
            "# Leander Antony - Senior ML Engineer Cover Letter\n"
            "\n"
            "Dear Hiring Team,\n"
            "\n"
            "I am writing to apply for the Senior ML Engineer role at Acme. "
            "Three points stand out from my background:\n"
            "\n"
            "- Built distributed Python services on AWS\n"
            "- Reduced p99 latency by 30%\n"
            "- Mentored two junior engineers\n"
            "\n"
            "Sincerely,\n"
            "\n"
            "Leander Antony"
        ),
        plain_text="placeholder",
        theme="classic_ats",
    )

    data = export_docx_bytes(artifact)
    doc = _parse_docx(data)
    pairs = _docx_paragraph_pairs(doc)
    texts = [text for _, text in pairs]

    # The H1 from the markdown is consumed as the header title; the
    # role suffix becomes the eyebrow line.
    assert "Leander Antony" in texts[0]
    # Body paragraphs surface as plain text.
    assert any("Dear Hiring Team," in t for t in texts)
    assert any("Senior ML Engineer role at Acme" in t for t in texts)
    # Three bullet items via the List Bullet style.
    bullets = [text for style, text in pairs if style == "List Bullet"]
    assert bullets == [
        "Built distributed Python services on AWS",
        "Reduced p99 latency by 30%",
        "Mentored two junior engineers",
    ]
    # Signoff lines preserved.
    assert any("Sincerely," in t for t in texts)
    assert any(t == "Leander Antony" for t in texts[-3:])


def test_export_docx_bytes_rejects_unsupported_artifact_type():
    """Defensive behavior: passing a non-artifact object should raise
    ExportError rather than blow up deep in python-docx."""
    try:
        export_docx_bytes("not an artifact")  # type: ignore[arg-type]
        raise AssertionError("Expected ExportError for non-artifact input")
    except ExportError as exc:
        assert "Unsupported artifact type" in exc.user_message


# ---------------------------------------------------------------------------
# Phase 4: theme switch (classic_ats vs professional_neutral)
# ---------------------------------------------------------------------------


def _docx_run_color_hexes(doc) -> set[str]:
    """Collect every hex color string used by any run in the document.

    python-docx exposes per-run colors via `run.font.color.rgb`
    (RGBColor or None when the run inherits). For our themed renderer
    every visible run gets an explicit color, so the union of these
    values is a fingerprint we can compare across themes."""
    colors: set[str] = set()
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            rgb = run.font.color.rgb
            if rgb is not None:
                colors.add(str(rgb))
    return colors


def _docx_run_font_names(doc) -> set[str]:
    """Same idea for font families. Returns the set of distinct font
    names used across all runs."""
    names: set[str] = set()
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            if run.font.name:
                names.add(run.font.name)
    return names


def test_export_docx_bytes_classic_ats_uses_warm_brown_palette():
    artifact = _make_full_resume_artifact()
    artifact.theme = "classic_ats"

    data = export_docx_bytes(artifact)
    doc = _parse_docx(data)

    colors = _docx_run_color_hexes(doc)
    # Warm-brown classic_ats palette: ink=221912, muted=6B5648,
    # accent=8F6845. At least the section-heading accent should land
    # in the rendered file.
    assert "8F6845" in colors, f"missing classic_ats accent in {colors}"
    assert "221912" in colors, f"missing classic_ats ink in {colors}"

    fonts = _docx_run_font_names(doc)
    # Unified typography (2026-05-21): every theme — classic_ats
    # included — uses the single Arial sans family. No serif font
    # should appear anywhere in the DOCX run set.
    assert "Arial" in fonts
    assert "Georgia" not in fonts, f"serif leaked into classic_ats: {fonts}"


def test_export_docx_bytes_professional_neutral_uses_black_palette():
    artifact = _make_full_resume_artifact()
    artifact.theme = "professional_neutral"

    data = export_docx_bytes(artifact)
    doc = _parse_docx(data)

    colors = _docx_run_color_hexes(doc)
    # Conservative palette: ink=0A0A0A, muted=555555, accent=0A0A0A
    # (collapses to ink — no warm tone), line=BFBFBF (used for
    # underline shapes, not run text, so it may or may not surface in
    # the run-color set).
    assert "0A0A0A" in colors, f"missing neutral ink in {colors}"
    assert "555555" in colors, f"missing neutral muted in {colors}"
    # The warm-brown accent must NOT appear in this theme.
    assert "8F6845" not in colors, f"classic_ats accent leaked: {colors}"

    fonts = _docx_run_font_names(doc)
    # Unified typography (2026-05-21): professional_neutral now runs
    # the shared Arial sans family throughout (it used to be all-serif
    # Georgia). No serif font should appear in run-level fonts.
    assert "Arial" in fonts
    assert "Georgia" not in fonts, f"serif leaked into neutral theme: {fonts}"


def test_export_docx_bytes_themes_produce_different_outputs():
    """Sanity check: render the same artifact under two themes and
    verify the PALETTE actually differs — a regression guard against
    accidentally short-circuiting the theme switch.

    Compares the extracted run-color set, NOT raw .docx bytes: a
    .docx is a ZIP and python-docx stamps the current mtime into
    every entry header, so two renders always have different bytes
    regardless of content — a raw-byte `!=` would pass even if the
    palette switch were broken. The color set is the real fingerprint."""
    artifact = _make_full_resume_artifact()

    artifact.theme = "classic_ats"
    classic_doc = _parse_docx(export_docx_bytes(artifact))

    artifact.theme = "professional_neutral"
    neutral_doc = _parse_docx(export_docx_bytes(artifact))

    # classic_ats ships warm-brown ink/accent; professional_neutral is
    # pure black/gray. Their run-color sets must differ.
    assert _docx_run_color_hexes(classic_doc) != _docx_run_color_hexes(neutral_doc)


def test_export_docx_bytes_unknown_theme_falls_back_to_classic_ats():
    """Unknown theme strings should resolve to classic_ats so the
    renderer never crashes on an unexpected artifact.theme value."""
    artifact = _make_full_resume_artifact()
    artifact.theme = "some_made_up_theme_we_have_not_built"
    fallback_doc = _parse_docx(export_docx_bytes(artifact))

    artifact.theme = "classic_ats"
    classic_doc = _parse_docx(export_docx_bytes(artifact))

    # The unknown-theme render must be EQUIVALENT to classic_ats —
    # same content/structure, same palette, same fonts. We compare
    # the extracted document, NOT raw .docx bytes: a .docx is a ZIP
    # and python-docx stamps the current mtime into every entry
    # header, so two renders a second apart differ byte-wise even
    # when the content is identical (this test used to be flaky for
    # exactly that reason). Compare what actually matters.
    assert _docx_paragraph_pairs(fallback_doc) == _docx_paragraph_pairs(classic_doc)
    assert _docx_run_color_hexes(fallback_doc) == _docx_run_color_hexes(classic_doc)
    assert _docx_run_font_names(fallback_doc) == _docx_run_font_names(classic_doc)


def test_export_docx_bytes_cover_letter_respects_theme():
    """Cover letter varies by color only (all themes share the one
    Arial sans family since the 2026-05-21 typography unification);
    the color set should still differ per theme."""
    classic = CoverLetterArtifact(
        title="Leander - Cover Letter",
        filename_stem="leander",
        summary="",
        markdown=(
            "# Leander - Cover Letter\n\n"
            "Dear Hiring Team,\n\n"
            "Body paragraph here.\n"
        ),
        plain_text="placeholder",
        theme="classic_ats",
    )
    classic_data = export_docx_bytes(classic)
    classic_colors = _docx_run_color_hexes(_parse_docx(classic_data))
    assert "221912" in classic_colors  # classic_ats ink

    neutral = CoverLetterArtifact(
        title=classic.title,
        filename_stem=classic.filename_stem,
        summary=classic.summary,
        markdown=classic.markdown,
        plain_text=classic.plain_text,
        theme="professional_neutral",
    )
    neutral_data = export_docx_bytes(neutral)
    neutral_colors = _docx_run_color_hexes(_parse_docx(neutral_data))
    assert "0A0A0A" in neutral_colors  # neutral ink
    # Warm-brown classic ink must not leak into the neutral cover letter.
    assert "221912" not in neutral_colors


def _headline_artifact(target_role, theme="classic_ats"):
    return TailoredResumeArtifact(
        title="Resume",
        filename_stem="r",
        summary="s",
        markdown="# Resume",
        plain_text="Resume",
        theme=theme,
        header=ResumeHeader(
            full_name="Leander Antony",
            location="Chennai, India",
            contact_lines=["leander@example.com"],
        ),
        target_role=target_role,
        professional_summary="Builds grounded AI products.",
        highlighted_skills=["Python"],
        education_entries=[],
        certifications=[],
    )


def test_resume_headline_renders_target_role_between_name_and_contact():
    # JD-tailored path: target_role set → role line shows, ordered
    # name → role → contact, in BOTH the classic and two-column header.
    for theme, role_cls in (
        ("classic_ats", "resume-classic-role"),
        ("presentation_twocol", "resume-tc-role"),
    ):
        art = _headline_artifact("AI Engineer - FDE", theme=theme)
        out = _build_resume_html(art.markdown, theme=theme, artifact=art)
        role_el = '<p class="{0}">'.format(role_cls)
        # Assert on the ELEMENT tag, not the bare class — the class
        # also appears in the <style> block.
        assert role_el + "AI Engineer - FDE</p>" in out
        assert out.index("<h1>") < out.index(role_el) < out.index(
            '<p class="resume-contact-inline'
        )


def test_resume_headline_omitted_when_no_target_role():
    # No-JD / resume-builder path: target_role "" → NO role line at all
    # (name-only header is standard; never fabricated), in both layouts.
    for theme, role_cls in (
        ("classic_ats", "resume-classic-role"),
        ("presentation_twocol", "resume-tc-role"),
    ):
        art = _headline_artifact("", theme=theme)
        out = _build_resume_html(art.markdown, theme=theme, artifact=art)
        # The ELEMENT must be absent (the .{cls} {{}} CSS rule in the
        # <style> block legitimately remains — assert on the tag).
        assert '<p class="{0}">'.format(role_cls) not in out
        assert "Leander Antony" in out and "leander@example.com" in out


def _contact_rows(html_str):
    import re

    return [
        (cls, re.findall(r'<span class="rc-item">([^<]+)</span>', body))
        for cls, body in re.findall(
            r'<p class="([^"]+)">(.*?)</p>', html_str, flags=re.DOTALL
        )
    ]


def test_looks_like_contact_link_classifies_details_vs_links():
    # Emails and plain details are NOT links even though they contain a
    # dot/domain; a URL that merely contains '@' (medium.com/@handle)
    # still IS a link (the bug this guards).
    assert _looks_like_contact_link("antony.leander@gmail.com") is False
    assert _looks_like_contact_link("Chennai, India") is False
    assert _looks_like_contact_link("+91 8610317213") is False
    assert _looks_like_contact_link("github.com/LEANDERANTONY") is True
    assert _looks_like_contact_link("linkedin.com/in/x") is True
    assert _looks_like_contact_link("leander-portfolio.framer.website") is True
    assert _looks_like_contact_link("https://x.io/a") is True
    assert _looks_like_contact_link("medium.com/@leander") is True


def test_contact_block_packs_into_at_most_two_lines_without_splitting_urls():
    details = ["Chennai, India", "+91 8610317213", "antony.leander@gmail.com"]

    # 0-1 links → ONE line (no links row).
    one = _contact_rows(_build_resume_contact_inline_html(details + ["github.com/x"]))
    assert len(one) == 1
    assert "github.com/x" in one[0][1]

    # 2 links → details on line 1, BOTH links together on line 2
    # (no lone short link stranded on its own line).
    two = _contact_rows(
        _build_resume_contact_inline_html(
            details + ["github.com/x", "leander.framer.website"]
        )
    )
    assert len(two) == 2
    assert two[0][1] == details
    assert "resume-contact-links" in two[1][0]
    assert two[1][1] == ["github.com/x", "leander.framer.website"]

    # 3 links → details + first link on line 1, the other two on line 2.
    three = _contact_rows(
        _build_resume_contact_inline_html(
            details + ["github.com/x", "linkedin.com/in/x", "p.dev"]
        )
    )
    assert len(three) == 2
    assert three[0][1] == details + ["github.com/x"]
    assert three[1][1] == ["linkedin.com/in/x", "p.dev"]

    # URL-never-splits guarantee: every value is wrapped in an
    # individually nowrap rc-item span (the CSS keeps it whole).
    for _cls, items in three:
        for value in items:
            assert (
                '<span class="rc-item">{0}</span>'.format(value)
                in _build_resume_contact_inline_html(
                    details + ["github.com/x", "linkedin.com/in/x", "p.dev"]
                )
            )

