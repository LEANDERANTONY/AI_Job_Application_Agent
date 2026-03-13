from io import BytesIO

import docx

from src.resume_parser import parse_resume


class NamedBytesIO(BytesIO):
    def __init__(self, initial_bytes=b"", name="sample.bin"):
        super().__init__(initial_bytes)
        self.name = name


def build_docx_resume():
    document = docx.Document()
    document.add_paragraph("Leander Antony")
    document.add_paragraph("Machine Learning Engineer")
    handle = NamedBytesIO(name="resume.docx")
    document.save(handle)
    handle.seek(0)
    return handle


def test_parse_resume_reads_docx():
    parsed_text, filetype = parse_resume(build_docx_resume())

    assert filetype == "DOCX"
    assert "Leander Antony" in parsed_text
    assert "Machine Learning Engineer" in parsed_text


def test_parse_resume_reads_plain_text():
    handle = NamedBytesIO(b"Leander Antony\nPython", name="resume.txt")

    parsed_text, filetype = parse_resume(handle)

    assert filetype == "TXT"
    assert parsed_text == "Leander Antony\nPython"

